#!/usr/bin/env python3
"""
创建包含中文架构图的Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# 路径
OUTPUT_DIR = "/root/.openclaw/workspace/03_输出成果/Token采集需求分析"
ARCH_DIR = os.path.join(OUTPUT_DIR, "arch_diagrams")

# 创建文档
doc = Document()

# 设置文档标题
title = doc.add_heading('Token采集系统方案文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加副标题
subtitle = doc.add_paragraph('架构设计方案')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = None

# 添加日期
date_para = doc.add_paragraph('2026年3月18日')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # 空行

# 第一节：总体架构
doc.add_heading('一、总体架构', level=1)
doc.add_paragraph(
    'Token采集系统采用四层架构设计，从下到上依次为：数据源层、数据采集层、数据处理层、数据应用层。'
)
doc.add_paragraph(
    '【数据源层】包含总部部门（23个）、专业公司（26个）、省公司（31个）、直属机构（6个），共86个数据源单位。'
)
doc.add_paragraph(
    '【数据采集层】提供SFTP接收、文件解析、文件校验、回执生成四大核心功能，支撑数据文件采集全流程。'
)
doc.add_paragraph(
    '【数据处理层】负责数据清洗、转换、入库、汇总，将原始数据转化为可分析的标准化数据。'
)
doc.add_paragraph(
    '【数据应用层】提供Token统计Dashboard、成本分析Dashboard、效果评估Dashboard、可视化报表、数据查询API等应用服务。'
)

# 添加总体架构图
if os.path.exists(os.path.join(ARCH_DIR, '总体架构图_中文版.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(ARCH_DIR, '总体架构图_中文版.png'), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph('图1：系统总体架构图')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = None

doc.add_paragraph()  # 空行

# 第二节：部署架构
doc.add_heading('二、部署架构', level=1)
doc.add_paragraph(
    '系统部署在生产环境，通过CN2网络与各数据源单位连接。核心网络信息如下：'
)
doc.add_paragraph('• SFTP服务器地址：10.141.208.176:12222')
doc.add_paragraph('• 网络类型：CN2-1124 电信专网')
doc.add_paragraph('• 文件格式：DAT.gz（GZIP压缩）')

# 添加部署架构图
if os.path.exists(os.path.join(ARCH_DIR, '部署架构图_中文版.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(ARCH_DIR, '部署架构图_中文版.png'), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph('图2：系统部署架构图')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = None

doc.add_paragraph()  # 空行

# 第三节：数据库架构
doc.add_heading('三、数据库架构', level=1)
doc.add_paragraph(
    '数据存储采用Hive+Doris混合架构。明细数据存储在Hive，汇总数据和应用查询使用Doris。'
)
doc.add_paragraph(
    '【分层设计】'
)
doc.add_paragraph('• 明细层（Hive）：存储原始Token明细数据，按data_date日分区，保留180天')
doc.add_paragraph('• 汇总层（Hive）：存储日汇总、月汇总、单位汇总数据，保留3年')
doc.add_paragraph('• 应用层（Doris）：提供高性能查询服务，支撑Dashboard和API')

# 添加数据库架构图
if os.path.exists(os.path.join(ARCH_DIR, '数据库架构图_中文版.png')):
    doc.add_paragraph()
    doc.add_picture(os.path.join(ARCH_DIR, '数据库架构图_中文版.png'), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph('图3：数据库架构图')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = None

doc.add_paragraph()  # 空行

# 第四节：关键设计要点
doc.add_heading('四、关键设计要点', level=1)

points = [
    ('数据采集方式', '采用SFTP被动接收模式，各数据源单位主动推送数据文件。'),
    ('文件格式规范', 'DAT格式，GZIP压缩，文件命名：TOKEN_YYYYMMDD_机构代码.DAT.gz'),
    ('数据分区策略', '按data_date进行日分区，明细数据保留180天，汇总数据保留3年。'),
    ('安全机制', 'CN2专网传输，TLS加密，RBAC访问控制。'),
    ('监控告警', 'Prometheus监控，AlertManager告警，Grafana可视化。'),
    ('任务调度', 'Airflow编排，支持定时采集、重试、告警。'),
]

for i, (title, desc) in enumerate(points, 1):
    para = doc.add_paragraph()
    para.add_run(f'{i}. {title}：').bold = True
    para.add_run(desc)

doc.add_paragraph()  # 空行

# 第五节：技术栈
doc.add_heading('五、技术栈', level=1)

tech_stack = [
    ('SFTP服务', 'Apache Mina SSHD'),
    ('文件处理', 'Python 3.9+ (pandas, loguru)'),
    ('数据入库', 'Apache Spark 3.x'),
    ('数据仓库', 'Apache Hive 3.x, Apache Doris 2.x'),
    ('任务调度', 'Apache Airflow 2.x'),
    ('监控告警', 'Prometheus + Grafana + AlertManager'),
    ('可视化', 'Apache Superset'),
]

for tech, desc in tech_stack:
    para = doc.add_paragraph()
    para.add_run(f'• {tech}：').bold = True
    para.add_run(desc)

doc.add_paragraph()  # 空行

# 第六节：项目工期
doc.add_heading('六、项目工期', level=1)
doc.add_paragraph('预计工期：25个工作日（约5周）')
doc.add_paragraph()
doc.add_paragraph('• 第1周：需求确认、环境准备、架构设计评审')
doc.add_paragraph('• 第2周：核心开发（SFTP服务、文件解析、文件校验）')
doc.add_paragraph('• 第3周：数据处理、入库、调度开发')
doc.add_paragraph('• 第4周：应用开发、接口联调、Dashboard开发')
doc.add_paragraph('• 第5周：系统集成测试、上线部署、文档交付')

# 保存文档
output_path = os.path.join(OUTPUT_DIR, 'Token采集系统_架构设计方案_中文版_20260318.docx')
doc.save(output_path)
print(f"文档已保存: {output_path}")
