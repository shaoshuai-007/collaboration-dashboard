#!/usr/bin/env python3
"""
Create Word document with architecture diagrams
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Paths
OUTPUT_DIR = "/root/.openclaw/workspace/03_输出成果/Token采集需求分析"
ARCH_DIR = os.path.join(OUTPUT_DIR, "arch_diagrams")

# Create document
doc = Document()

# Set document title
title = doc.add_heading('Token采集系统方案文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('架构设计方案')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = None

# Add date
date_para = doc.add_paragraph('2026年3月18日')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# Section 1: Overall Architecture
doc.add_heading('一、总体架构', level=1)
doc.add_paragraph(
    'Token采集系统采用四层架构设计，从下到上依次为：数据源层、数据采集层、数据处理层、数据应用层。'
)

# Add overall architecture diagram
if os.path.exists(os.path.join(ARCH_DIR, 'Overall_Architecture.png')):
    doc.add_paragraph('图1：系统总体架构图', style='Caption')
    doc.add_picture(os.path.join(ARCH_DIR, 'Overall_Architecture.png'), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# Section 2: Deployment Architecture
doc.add_heading('二、部署架构', level=1)
doc.add_paragraph(
    '系统部署在生产环境，通过CN2网络与各数据源单位连接。SFTP服务器地址：10.141.208.176:12222。'
)

# Add deployment diagram
if os.path.exists(os.path.join(ARCH_DIR, 'Deployment_Architecture.png')):
    doc.add_paragraph('图2：系统部署架构图', style='Caption')
    doc.add_picture(os.path.join(ARCH_DIR, 'Deployment_Architecture.png'), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# Section 3: Database Architecture
doc.add_heading('三、数据库架构', level=1)
doc.add_paragraph(
    '数据存储采用Hive+Doris混合架构。明细数据存储在Hive，汇总数据和应用查询使用Doris。'
)

# Add database diagram
if os.path.exists(os.path.join(ARCH_DIR, 'Database_Architecture.png')):
    doc.add_paragraph('图3：数据库架构图', style='Caption')
    doc.add_picture(os.path.join(ARCH_DIR, 'Database_Architecture.png'), width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# Section 4: Key Design Points
doc.add_heading('四、关键设计要点', level=1)

points = [
    ('数据采集方式', '采用SFTP被动接收模式，各数据源单位主动推送数据文件。'),
    ('文件格式规范', 'DAT格式，GZIP压缩，文件命名：TOKEN_YYYYMMDD_机构代码.DAT.gz'),
    ('数据分区策略', '按data_date进行日分区，明细数据保留180天，汇总数据保留3年。'),
    ('安全机制', 'CN2专网传输，TLS加密，RBAC访问控制。'),
    ('监控告警', 'Prometheus监控，AlertManager告警，Grafana可视化。'),
]

for i, (title, desc) in enumerate(points, 1):
    para = doc.add_paragraph()
    para.add_run(f'{i}. {title}：').bold = True
    para.add_run(desc)

doc.add_paragraph()  # Spacer

# Section 5: Project Schedule
doc.add_heading('五、项目工期', level=1)
doc.add_paragraph(
    '预计工期：25个工作日（约5周）\n'
    '- 第1周：需求确认、环境准备\n'
    '- 第2周：核心开发（SFTP服务、文件解析）\n'
    '- 第3周：数据处理、入库\n'
    '- 第4周：应用开发、接口联调\n'
    '- 第5周：测试、上线'
)

# Save document
output_path = os.path.join(OUTPUT_DIR, 'Token采集系统_架构设计方案_20260318.docx')
doc.save(output_path)
print(f"Document saved: {output_path}")
