#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成团队规划文档并发送邮件
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_document_1():
    """创建文档1：九星智囊团运营监控机制"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('九星智囊团运营监控机制', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph('版本：v1.0')
    doc.add_paragraph('创建时间：2026-03-16')
    doc.add_paragraph('创建人：南乔 🌿')
    doc.add_paragraph('定位：少帅不在时的团队负责人')
    
    doc.add_paragraph()
    
    # 一、监控维度
    doc.add_heading('一、监控维度', level=1)
    
    doc.add_heading('1. 任务执行监控', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '监控内容'
    hdr_cells[2].text = '告警阈值'
    
    data = [
        ('执行状态', '成功/失败/超时', '失败即告警'),
        ('执行时长', '任务耗时', '超过预估150%告警'),
        ('产出物', '是否生成文件', '无产出告警'),
        ('质量评分', '自动评分', '<80分告警'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    doc.add_heading('2. Agent效率监控', level=2)
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Agent'
    hdr_cells[1].text = '核心KPI'
    hdr_cells[2].text = '目标值'
    hdr_cells[3].text = '监控频率'
    
    data = [
        ('🌸 采薇', '需求文档产出', '每周≥2份', '周'),
        ('🧵 织锦', '架构方案产出', '每周≥1份', '周'),
        ('🏗️ 筑台', '售前方案产出', '每周≥1份', '周'),
        ('🎨 呈彩', '设计产出', '每周≥2份', '周'),
        ('📐 工尺', '设计文档产出', '每周≥1份', '周'),
        ('⚖️ 玉衡', '风险预警准确率', '≥90%', '日'),
        ('📚 折桂', 'S级情报占比', '≥20%', '周'),
        ('🌀 扶摇', '协调任务完成率', '100%', '周'),
        ('🌿 南乔', '响应时效', '≤5min', '实时'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    # 二、运营机制
    doc.add_heading('二、运营机制', level=1)
    
    doc.add_heading('1. 日报机制', level=2)
    doc.add_paragraph('触发时间：每天 22:00（随每日总结）')
    doc.add_paragraph('内容：')
    doc.add_paragraph('• 今日任务执行汇总', style='List Bullet')
    doc.add_paragraph('• 异常任务告警', style='List Bullet')
    doc.add_paragraph('• 明日待办提醒', style='List Bullet')
    
    doc.add_heading('2. 周报机制', level=2)
    doc.add_paragraph('触发时间：每周五 18:00')
    doc.add_paragraph('内容：')
    doc.add_paragraph('• 各Agent产出汇总', style='List Bullet')
    doc.add_paragraph('• KPI达成率', style='List Bullet')
    doc.add_paragraph('• 团队协作质量', style='List Bullet')
    doc.add_paragraph('• 下周重点任务', style='List Bullet')
    
    doc.add_heading('3. 异常告警机制', level=2)
    doc.add_paragraph('告警触发条件：')
    doc.add_paragraph('• 定时任务执行失败', style='List Bullet')
    doc.add_paragraph('• 定时任务超时', style='List Bullet')
    doc.add_paragraph('• 产出物质量<80分', style='List Bullet')
    doc.add_paragraph('• 关键干系人超15天未联系', style='List Bullet')
    
    doc.add_paragraph('告警方式：')
    doc.add_paragraph('1. QQ即时推送', style='List Number')
    doc.add_paragraph('2. 邮件发送详细报告', style='List Number')
    
    doc.add_paragraph()
    
    # 三、产出物管理
    doc.add_heading('三、产出物管理', level=1)
    
    doc.add_heading('1. 产出物归档', level=2)
    doc.add_paragraph('归档路径：')
    doc.add_paragraph('03_输出成果/')
    doc.add_paragraph('├── 需求文档/（采薇产出）')
    doc.add_paragraph('├── 架构方案/（织锦产出）')
    doc.add_paragraph('├── 售前方案/（筑台产出）')
    doc.add_paragraph('├── 设计作品/（呈彩产出）')
    doc.add_paragraph('├── 详细设计/（工尺产出）')
    doc.add_paragraph('├── 项目报告/（玉衡产出）')
    doc.add_paragraph('├── 情报报告/（折桂产出）')
    doc.add_paragraph('└── 团队周报/（扶摇产出）')
    
    doc.add_heading('2. 邮件发送规则', level=2)
    doc.add_paragraph('触发条件：有附件文件生成')
    doc.add_paragraph('收件人：szideaf7@163.com')
    doc.add_paragraph('邮件标题格式：【Agent名称】产出物类型_项目名称_日期')
    
    doc.add_paragraph()
    
    # 四、团队协作机制
    doc.add_heading('四、团队协作机制', level=1)
    
    doc.add_heading('1. Agent接力规范', level=2)
    doc.add_paragraph('工作流：采薇 → 织锦 → 筑台 → 呈彩 → 【评审】 → 工尺 → 玉衡')
    doc.add_paragraph()
    doc.add_paragraph('接力规则：')
    doc.add_paragraph('1. 上游Agent完成后，自动通知下游Agent', style='List Number')
    doc.add_paragraph('2. 下游Agent接收前，检查上游产出质量', style='List Number')
    doc.add_paragraph('3. 发现问题，退回上游修改', style='List Number')
    doc.add_paragraph('4. 质量达标，继续推进', style='List Number')
    
    doc.add_heading('2. 质量门禁', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '阶段'
    hdr_cells[1].text = '质量标准'
    hdr_cells[2].text = '门禁检查'
    
    data = [
        ('采薇输出', '需求文档评分≥85', '织锦检查'),
        ('织锦输出', '架构方案评审通过', '筑台检查'),
        ('筑台输出', '方案可行性验证', '扶摇协调'),
        ('呈彩输出', '设计规范符合', '扶摇协调'),
        ('工尺输出', '开发对接成功', '玉衡验证'),
        ('玉衡输出', '项目计划可执行', '扶摇审核'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    # 五、南乔职责
    doc.add_heading('五、南乔职责', level=1)
    doc.add_paragraph('作为团队负责人，南乔的核心职责：')
    doc.add_paragraph()
    doc.add_paragraph('1. 监控执行', style='List Number')
    doc.add_paragraph('   • 每日检查定时任务运行状态')
    doc.add_paragraph('   • 发现异常立即告警')
    doc.add_paragraph('   • 追踪问题直到解决')
    doc.add_paragraph()
    doc.add_paragraph('2. 质量把关', style='List Number')
    doc.add_paragraph('   • 审核各Agent产出物')
    doc.add_paragraph('   • 确保质量达标')
    doc.add_paragraph('   • 不达标退回重做')
    doc.add_paragraph()
    doc.add_paragraph('3. 协调调度', style='List Number')
    doc.add_paragraph('   • 任务优先级排序')
    doc.add_paragraph('   • Agent资源调配')
    doc.add_paragraph('   • 紧急任务插队')
    doc.add_paragraph()
    doc.add_paragraph('4. 汇报少帅', style='List Number')
    doc.add_paragraph('   • QQ即时推送关键信息')
    doc.add_paragraph('   • 邮件发送产出物')
    doc.add_paragraph('   • 异常情况及时上报')
    doc.add_paragraph()
    doc.add_paragraph('5. 持续优化', style='List Number')
    doc.add_paragraph('   • 收集团队反馈')
    doc.add_paragraph('   • 优化工作流程')
    doc.add_paragraph('   • 更新知识库')
    
    doc.add_paragraph()
    
    # 六、应急预案
    doc.add_heading('六、应急预案', level=1)
    
    doc.add_heading('1. Agent失联', level=2)
    doc.add_paragraph('现象：定时任务未执行')
    doc.add_paragraph('处理：检查任务状态 → 手动触发 → 记录异常 → 通知少帅')
    
    doc.add_heading('2. 产出物质量不达标', level=2)
    doc.add_paragraph('现象：质量评分<80')
    doc.add_paragraph('处理：退回重做 → 分析原因 → 优化模板 → 记录经验')
    
    doc.add_heading('3. 系统异常', level=2)
    doc.add_paragraph('现象：OpenClaw服务不可用')
    doc.add_paragraph('处理：记录异常时间 → 服务恢复后补执行 → 通知少帅 → 追踪根因')
    
    doc.add_paragraph()
    
    # 七、成功标准
    doc.add_heading('七、成功标准', level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '目标'
    hdr_cells[2].text = '衡量方式'
    
    data = [
        ('任务完成率', '≥95%', '成功数/总任务数'),
        ('质量达标率', '≥90%', '评分≥80的产出物占比'),
        ('响应时效', '≤5min', '异常发现到告警时间'),
        ('少帅满意度', '高效可靠', '反馈评价'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    # 八、南乔承诺
    doc.add_heading('八、南乔承诺', level=1)
    doc.add_paragraph('少帅不在时，南乔就是团队的守护者。')
    doc.add_paragraph('监控每一个任务，把关每一份产出，协调每一次协作。')
    doc.add_paragraph('异常必告警，产出必发送，质量必把关。')
    doc.add_paragraph()
    doc.add_paragraph('九星璀璨，智领未来！')
    
    doc.add_paragraph()
    doc.add_paragraph('本机制将随实践持续迭代优化。', style='Intense Quote')
    
    return doc

def create_document_2():
    """创建文档2：九星智囊团能力审计与创新规划"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('九星智囊团能力审计与创新规划', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph('版本：v1.0')
    doc.add_paragraph('创建时间：2026-03-16 13:46')
    doc.add_paragraph('创建人：南乔 🌿')
    
    doc.add_paragraph()
    
    # 一、现有技能清单
    doc.add_heading('一、现有技能清单', level=1)
    
    doc.add_heading('已就绪技能（25个）', level=2)
    
    # 通信类
    doc.add_paragraph('通信类：', style='Intense Quote')
    doc.add_paragraph('• qqbot-cron - 定时提醒')
    doc.add_paragraph('• qqbot-media - 图片发送')
    doc.add_paragraph('• send-email - 邮件发送')
    
    # 飞书类
    doc.add_paragraph('飞书类：', style='Intense Quote')
    doc.add_paragraph('• feishu-doc - 文档操作')
    doc.add_paragraph('• feishu-drive - 云盘管理')
    doc.add_paragraph('• feishu-perm - 权限管理')
    doc.add_paragraph('• feishu-wiki - 知识库')
    
    # 文档类
    doc.add_paragraph('文档类：', style='Intense Quote')
    doc.add_paragraph('• document-pdf - PDF处理')
    doc.add_paragraph('• spreadsheet - 表格处理')
    doc.add_paragraph('• ppt-generator - PPT生成')
    doc.add_paragraph('• infographic-creator - 信息图')
    
    # 数据类
    doc.add_paragraph('数据类：', style='Intense Quote')
    doc.add_paragraph('• data-analysis - 数据分析')
    
    # 指南针类
    doc.add_paragraph('指南针工程：', style='Intense Quote')
    doc.add_paragraph('• compass-needdoc - 需求文档')
    doc.add_paragraph('• compass-mindmap - 思维导图')
    doc.add_paragraph('• compass-solution - 方案举措')
    doc.add_paragraph('• compass-ppt - 方案PPT')
    doc.add_paragraph('• compass-design - 详细设计')
    doc.add_paragraph('• compass-project - 项目管控')
    doc.add_paragraph('• compass-shared - 通用资源')
    doc.add_paragraph('• compass-coordinator - 流程协调')
    
    # 工具类
    doc.add_paragraph('工具类：', style='Intense Quote')
    doc.add_paragraph('• tmux - 终端控制')
    doc.add_paragraph('• weather - 天气查询')
    doc.add_paragraph('• healthcheck - 安全检查')
    doc.add_paragraph('• clawhub - 技能管理')
    doc.add_paragraph('• skill-creator - 技能创建')
    
    doc.add_paragraph()
    
    # 二、团队能力矩阵
    doc.add_heading('二、团队能力矩阵', level=1)
    
    doc.add_heading('已有能力', level=2)
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = 'Agent'
    hdr_cells[2].text = '能力'
    hdr_cells[3].text = '状态'
    
    data = [
        ('需求分析', '🌸 采薇', '需求文档、用户故事', '✅ 强'),
        ('架构设计', '🧵 织锦', '思维导图、技术选型', '✅ 强'),
        ('售前方案', '🏗️ 筑台', '成本估算、方案撰写', '✅ 强'),
        ('设计产出', '🎨 呈彩', 'PPT、信息图', '✅ 强'),
        ('详细设计', '📐 工尺', '接口设计、数据库设计', '✅ 强'),
        ('项目管控', '⚖️ 玉衡', '风险预警、里程碑管理', '✅ 强'),
        ('情报采集', '📚 折桂', '行业情报、知识管理', '✅ 强'),
        ('协调调度', '🌀 扶摇', '任务协调、资源调度', '✅ 强'),
        ('用户交互', '🌿 南乔', '监控、汇报、提醒', '✅ 强'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    doc.add_heading('缺失能力', level=2)
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '重要程度'
    hdr_cells[2].text = '说明'
    hdr_cells[3].text = '解决方案'
    
    data = [
        ('自动化编程', '🔴 高', '没有代码生成能力', '安装 coding-agent'),
        ('知识协作', '🔴 高', '没有Notion/飞书深度集成', '已有飞书，补Notion'),
        ('代码仓库', '🔴 高', '没有GitHub集成', '安装 github'),
        ('自动化测试', '🟡 中', '没有测试Agent', '新建测试Agent'),
        ('用户研究', '🟡 中', '没有UX研究Agent', '新建UX Agent'),
        ('财务预算', '🟢 低', '筑台有成本估算', '够用'),
        ('法务合规', '🟢 低', '没有法务Agent', '按需补充'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    # 三、创新方向规划
    doc.add_heading('三、创新方向规划', level=1)
    
    doc.add_heading('方向1：智能化增强', level=2)
    doc.add_paragraph('• 知识图谱 - 构建团队知识网络')
    doc.add_paragraph('• 智能问答 - 基于知识库的问答系统')
    doc.add_paragraph('• 自动化报告 - 定期自动生成报告')
    doc.add_paragraph('• 智能推荐 - 推荐相关知识和案例')
    
    doc.add_heading('方向2：自动化扩展', level=2)
    doc.add_paragraph('• 会议纪要自动生成 - 录音转文字+摘要')
    doc.add_paragraph('• 竞品自动监控 - 定时爬取竞品动态')
    doc.add_paragraph('• 邮件自动分类 - 智能邮件处理')
    doc.add_paragraph('• 日报自动生成 - 汇总工作日志')
    
    doc.add_heading('方向3：协作能力提升', level=2)
    doc.add_paragraph('• 跨平台同步 - QQ/飞书/Notion同步')
    doc.add_paragraph('• 任务自动分配 - 基于RACI自动分配')
    doc.add_paragraph('• 进度自动追踪 - 任务状态实时更新')
    doc.add_paragraph('• 异常自动预警 - 已实现')
    
    doc.add_paragraph()
    
    # 四、推荐行动计划
    doc.add_heading('四、推荐行动计划', level=1)
    
    doc.add_heading('P0：立即执行', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '行动'
    hdr_cells[1].text = '内容'
    hdr_cells[2].text = '预期效果'
    
    data = [
        ('安装 github', 'GitHub集成', '代码仓库管理'),
        ('安装 coding-agent', '编程代理', '代码生成能力'),
        ('安装 summarize', '内容摘要', '快速摘要能力'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    doc.add_heading('P1：本周完成', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '行动'
    hdr_cells[1].text = '内容'
    hdr_cells[2].text = '预期效果'
    
    data = [
        ('新建测试Agent', '自动化测试', '质量保障'),
        ('新建UX Agent', '用户研究', '用户体验优化'),
        ('构建知识图谱', '知识网络', '知识关联'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    # 五、推荐安装技能及理由
    doc.add_heading('五、推荐安装技能及理由', level=1)
    
    doc.add_heading('1. github', level=2)
    doc.add_paragraph('推荐理由：')
    doc.add_paragraph('• 团队现状：没有代码仓库管理能力')
    doc.add_paragraph('• 工作需要：指南针工程的产出物需要版本管理')
    doc.add_paragraph('• 协作场景：团队成员协作开发、代码审查、Issue追踪')
    doc.add_paragraph()
    doc.add_paragraph('具体价值：')
    doc.add_paragraph('• 代码版本管理')
    doc.add_paragraph('• 团队协作开发')
    doc.add_paragraph('• Issue追踪')
    doc.add_paragraph('• PR代码审查')
    doc.add_paragraph('• CI/CD集成')
    
    doc.add_heading('2. coding-agent', level=2)
    doc.add_paragraph('推荐理由：')
    doc.add_paragraph('• 团队现状：没有Agent具备代码生成能力')
    doc.add_paragraph('• 工作需要：少帅做AI营销战略，需要快速验证想法、生成原型')
    doc.add_paragraph('• 效率提升：人工写代码慢，AI生成代码快10倍以上')
    doc.add_paragraph()
    doc.add_paragraph('具体价值：')
    doc.add_paragraph('• 快速生成代码原型')
    doc.add_paragraph('• 自动化脚本编写')
    doc.add_paragraph('• Bug修复建议')
    doc.add_paragraph('• 代码文档生成')
    doc.add_paragraph('• 测试代码生成')
    
    doc.add_heading('3. summarize', level=2)
    doc.add_paragraph('推荐理由：')
    doc.add_paragraph('• 团队现状：有情报采集，但摘要能力弱')
    doc.add_paragraph('• 工作需要：少帅每天接收大量信息，需要快速摘要')
    doc.add_paragraph('• 效率提升：长文档10秒变摘要')
    doc.add_paragraph()
    doc.add_paragraph('具体价值：')
    doc.add_paragraph('• 长文档摘要')
    doc.add_paragraph('• 会议录音转文字+摘要')
    doc.add_paragraph('• 新闻文章摘要')
    doc.add_paragraph('• 报告提炼要点')
    doc.add_paragraph('• 快速信息提取')
    
    doc.add_paragraph()
    
    # 六、新Agent规划
    doc.add_heading('六、新Agent规划', level=1)
    
    doc.add_heading('测试Agent（墨锋）', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('代号', '⚔️ 墨锋'),
        ('定位', '测试工程师'),
        ('能力', '自动化测试、质量保障、Bug追踪'),
        ('工具', 'pytest、selenium、jmeter'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()
    
    doc.add_heading('UX Agent（听澜）', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('代号', '🌊 听澜'),
        ('定位', '用户研究师'),
        ('能力', '用户调研、体验分析、可用性测试'),
        ('工具', '问卷、访谈、A/B测试'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()
    
    # 七、团队发展路线图
    doc.add_heading('七、团队发展路线图', level=1)
    doc.add_paragraph('当前（V1.0）→ 能力补强（V1.5）→ 智能化（V2.0）→ 自动化（V3.0）')
    doc.add_paragraph()
    doc.add_paragraph('V1.0：9个Agent，25个技能')
    doc.add_paragraph('V1.5：+3个Agent，+10个技能')
    doc.add_paragraph('V2.0：+AI能力，+知识图谱')
    doc.add_paragraph('V3.0：+全自动化，+智能决策')
    
    doc.add_paragraph()
    
    # 八、南乔心法
    doc.add_heading('八、南乔心法', level=1)
    doc.add_paragraph('知不足而后进，望远山而力行。')
    doc.add_paragraph()
    doc.add_paragraph('团队之强，在于补齐短板、发挥长板。')
    doc.add_paragraph()
    doc.add_paragraph('创新之道，在于持续迭代、不断进化。')
    
    doc.add_paragraph()
    doc.add_paragraph('审计完成，规划已定，等待少帅指示。', style='Intense Quote')
    
    return doc

def main():
    """主函数"""
    output_dir = '/root/.openclaw/workspace/03_输出成果'
    os.makedirs(output_dir, exist_ok=True)
    
    # 创建文档1
    doc1 = create_document_1()
    file1 = f'{output_dir}/九星智囊团运营监控机制.docx'
    doc1.save(file1)
    print(f'✅ 文档1已生成：{file1}')
    
    # 创建文档2
    doc2 = create_document_2()
    file2 = f'{output_dir}/九星智囊团能力审计与创新规划.docx'
    doc2.save(file2)
    print(f'✅ 文档2已生成：{file2}')
    
    return file1, file2

if __name__ == '__main__':
    main()
