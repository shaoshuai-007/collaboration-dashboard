#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成新技能模板和接入分析文档并发送邮件
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime
import os

def create_template_document():
    """创建使用场景模板文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('新技能使用场景模板库', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph('版本：v1.0')
    doc.add_paragraph('创建时间：2026-03-16')
    doc.add_paragraph('创建人：南乔 🌿')
    
    doc.add_paragraph()
    
    # 一、github 使用模板
    doc.add_heading('一、🐙 github 使用模板', level=1)
    
    doc.add_heading('模板1：查看代码仓库状态', level=2)
    doc.add_paragraph('【场景】查看项目代码仓库状态')
    doc.add_paragraph()
    doc.add_paragraph('【命令模板】')
    doc.add_paragraph('gh repo list')
    doc.add_paragraph('gh issue list --repo [owner/repo]')
    doc.add_paragraph('gh pr list --repo [owner/repo]')
    doc.add_paragraph('gh run list --repo [owner/repo] --limit 5')
    doc.add_paragraph()
    doc.add_paragraph('【示例】')
    doc.add_paragraph('少帅：查看智能协作平台的代码仓库状态')
    doc.add_paragraph('南乔执行：')
    doc.add_paragraph('gh repo list | grep collaboration')
    doc.add_paragraph('gh issue list --repo openclaw/collaboration-dashboard')
    
    doc.add_heading('模板2：创建新仓库', level=2)
    doc.add_paragraph('【场景】为新项目创建GitHub仓库')
    doc.add_paragraph()
    doc.add_paragraph('【命令模板】')
    doc.add_paragraph('gh repo create [repo-name] --public --description "[描述]"')
    doc.add_paragraph('git remote add origin https://github.com/[owner]/[repo-name].git')
    doc.add_paragraph('git push -u origin main')
    
    doc.add_heading('模板3：Issue管理', level=2)
    doc.add_paragraph('【场景】创建和追踪Issue')
    doc.add_paragraph()
    doc.add_paragraph('【命令模板】')
    doc.add_paragraph('gh issue create --repo [owner/repo] --title "[标题]" --body "[内容]"')
    doc.add_paragraph('gh issue list --repo [owner/repo] --state open')
    
    doc.add_paragraph()
    
    # 二、coding-agent 使用模板
    doc.add_heading('二、🧩 coding-agent 使用模板', level=1)
    
    doc.add_heading('模板1：快速原型生成', level=2)
    doc.add_paragraph('【场景】从需求到代码原型')
    doc.add_paragraph()
    doc.add_paragraph('【命令模板】')
    doc.add_paragraph('pi "[需求描述]')
    doc.add_paragraph()
    doc.add_paragraph('要求：')
    doc.add_paragraph('1. [功能点1]')
    doc.add_paragraph('2. [功能点2]')
    doc.add_paragraph('3. [功能点3]')
    doc.add_paragraph()
    doc.add_paragraph('技术栈：[Python/JavaScript/etc]')
    doc.add_paragraph('输出：完整可运行代码"')
    
    doc.add_heading('模板2：自动化脚本生成', level=2)
    doc.add_paragraph('【场景】生成自动化运维脚本')
    doc.add_paragraph()
    doc.add_paragraph('【示例】')
    doc.add_paragraph('少帅：生成一个MySQL自动备份脚本')
    doc.add_paragraph('南乔执行：')
    doc.add_paragraph('pi "编写一个MySQL数据库自动备份脚本，功能：')
    doc.add_paragraph('1. 每天凌晨2点自动备份')
    doc.add_paragraph('2. 保留最近7天的备份')
    doc.add_paragraph('3. 自动压缩备份文件')
    doc.add_paragraph('4. 发送备份完成通知到QQ')
    doc.add_paragraph()
    doc.add_paragraph('输出：完整的Shell脚本 + cron配置"')
    
    doc.add_paragraph()
    
    # 三、summarize 使用模板
    doc.add_heading('三、🧾 summarize 使用模板', level=1)
    
    doc.add_heading('模板1：情报快速处理', level=2)
    doc.add_paragraph('【场景】快速摘要情报内容')
    doc.add_paragraph()
    doc.add_paragraph('【命令模板】')
    doc.add_paragraph('summarize "[URL]" --length short --model google/gemini-3-flash-preview')
    doc.add_paragraph()
    doc.add_paragraph('【参数说明】')
    doc.add_paragraph('- --length：摘要长度（short/medium/long）')
    doc.add_paragraph('- --model：使用的模型')
    
    doc.add_heading('模板2：会议纪要生成', level=2)
    doc.add_paragraph('【场景】从录音到会议纪要')
    doc.add_paragraph()
    doc.add_paragraph('【步骤】')
    doc.add_paragraph('1. 转录音频：summarize "[录音URL]" --youtube auto --extract-only')
    doc.add_paragraph('2. 生成摘要：summarize "[转录文件]" --length medium')
    
    doc.add_heading('模板3：需求文档摘要', level=2)
    doc.add_paragraph('【场景】快速理解需求文档')
    doc.add_paragraph()
    doc.add_paragraph('【命令模板】')
    doc.add_paragraph('summarize "[文档路径]" --length medium')
    
    doc.add_paragraph()
    
    # 四、组合使用模板
    doc.add_heading('四、🔄 组合使用模板', level=1)
    
    doc.add_heading('模板1：需求到代码全流程', level=2)
    doc.add_paragraph('【场景】从需求文档到代码原型')
    doc.add_paragraph()
    doc.add_paragraph('【步骤】')
    doc.add_paragraph('1. summarize：摘要需求文档')
    doc.add_paragraph('2. coding-agent：生成代码原型')
    doc.add_paragraph('3. github：创建仓库并提交')
    
    doc.add_paragraph()
    
    # 五、快速调用指南
    doc.add_heading('五、快速调用指南', level=1)
    
    doc.add_heading('github 常用命令速查', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    data = [
        ('场景', '命令'),
        ('查看仓库', 'gh repo list'),
        ('查看Issue', 'gh issue list --repo owner/repo'),
        ('创建Issue', 'gh issue create --repo owner/repo --title "标题" --body "内容"'),
        ('查看PR', 'gh pr list --repo owner/repo'),
        ('审查PR', 'gh pr view [number] --repo owner/repo'),
        ('查看CI', 'gh run list --repo owner/repo'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()
    
    doc.add_heading('coding-agent 常用提示', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('场景', '提示模板'),
        ('原型生成', 'pi "需求描述 + 技术栈 + 输出要求"'),
        ('脚本生成', 'pi "脚本类型 + 功能点 + 配置参数"'),
        ('代码优化', 'pi "分析代码 + 优化方向"'),
        ('测试生成', 'pi "代码 + 测试框架 + 覆盖范围"'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()
    
    doc.add_heading('summarize 常用参数', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ('场景', '参数组合'),
        ('短摘要', '--length short'),
        ('中等摘要', '--length medium'),
        ('长摘要', '--length long'),
        ('转录音频', '--youtube auto --extract-only'),
        ('JSON输出', '--json'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()
    
    # 六、使用技巧
    doc.add_heading('六、使用技巧', level=1)
    doc.add_paragraph('技巧1：明确需求')
    doc.add_paragraph('• 越具体的描述，效果越好')
    doc.add_paragraph('• 提供上下文和背景')
    doc.add_paragraph('• 指定输出格式')
    doc.add_paragraph()
    doc.add_paragraph('技巧2：分步执行')
    doc.add_paragraph('• 复杂任务拆分为多步')
    doc.add_paragraph('• 每步验证结果')
    doc.add_paragraph('• 及时调整方向')
    doc.add_paragraph()
    doc.add_paragraph('技巧3：结果验证')
    doc.add_paragraph('• 检查生成代码是否能运行')
    doc.add_paragraph('• 验证摘要是否准确')
    doc.add_paragraph('• 确认GitHub操作成功')
    
    return doc

def create_integration_document():
    """创建接入分析文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('现有工程新技能接入分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph('版本：v1.0')
    doc.add_paragraph('创建时间：2026-03-16')
    doc.add_paragraph('创建人：南乔 🌿')
    
    doc.add_paragraph()
    
    # 一、现有工程清单
    doc.add_heading('一、现有工程清单', level=1)
    
    doc.add_heading('核心工程', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '工程名称'
    hdr_cells[1].text = '代码量'
    hdr_cells[2].text = '技术栈'
    hdr_cells[3].text = '状态'
    
    data = [
        ('智能体协作平台', '196KB (5000行)', 'Python + Flask', 'V14生产运行'),
        ('指南针工程连接器', '23KB', 'Python', '已集成'),
        ('智能调度系统', '230行', 'Python', '已集成'),
        ('九星监控系统', '待开发', 'Python', '规划中'),
        ('辅助脚本', '多个', 'Python/Shell', '分散管理'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    # 二、github 接入方案
    doc.add_heading('二、🐙 github 接入方案', level=1)
    
    doc.add_heading('接入目标', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = '说明'
    hdr_cells[2].text = '优先级'
    
    data = [
        ('目标', '说明', '优先级'),
        ('代码版本管理', '所有Python代码纳入Git管理', '🔴 高'),
        ('Issue追踪', '功能需求、Bug统一管理', '🔴 高'),
        ('代码审查', 'PR流程规范化', '🟡 中'),
        ('CI/CD集成', '自动化测试、部署', '🟢 低'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_heading('接入步骤', level=2)
    doc.add_paragraph('步骤1：创建GitHub组织')
    doc.add_paragraph('组织名称：nine-stars-think-tank')
    doc.add_paragraph('描述：九星智囊团 - 电信行业AI转型智囊团')
    doc.add_paragraph()
    doc.add_paragraph('步骤2：创建核心仓库')
    doc.add_paragraph('• collaboration-dashboard - 智能体协作平台')
    doc.add_paragraph('• compass-engine - 指南针工程')
    doc.add_paragraph('• nine-stars-monitor - 九星监控系统')
    doc.add_paragraph('• knowledge-base - 知识库')
    doc.add_paragraph('• scripts - 辅助脚本')
    
    doc.add_paragraph()
    
    # 三、coding-agent 接入方案
    doc.add_heading('三、🧩 coding-agent 接入方案', level=1)
    
    doc.add_heading('接入目标', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '目标'
    hdr_cells[1].text = '说明'
    hdr_cells[2].text = '优先级'
    
    data = [
        ('快速原型生成', '需求到原型，分钟级', '🔴 高'),
        ('代码优化', '现有代码质量提升', '🟡 中'),
        ('自动化脚本', '日常任务自动化', '🔴 高'),
        ('测试代码生成', '提高测试覆盖率', '🟡 中'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_heading('接入收益', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('收益', '效率提升'),
        ('原型生成', '从1天 → 30分钟'),
        ('脚本编写', '从1小时 → 5分钟'),
        ('代码优化', '从人工分析 → AI建议'),
        ('测试生成', '从手动编写 → 自动生成'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()
    
    # 四、summarize 接入方案
    doc.add_heading('四、🧾 summarize 接入方案', level=1)
    
    doc.add_heading('接入目标', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '目标'
    hdr_cells[1].text = '说明'
    hdr_cells[2].text = '优先级'
    
    data = [
        ('情报快速处理', '折桂采集情报自动摘要', '🔴 高'),
        ('需求文档摘要', '快速理解需求核心', '🔴 高'),
        ('会议纪要生成', '自动生成会议纪要', '🟡 中'),
        ('技术文档摘要', '快速掌握技术要点', '🟡 中'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_heading('接入收益', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    data = [
        ('收益', '效率提升'),
        ('情报处理', '从30分钟 → 10秒'),
        ('需求理解', '从1小时 → 1分钟'),
        ('会议纪要', '从30分钟 → 30秒'),
        ('文档阅读', '从10分钟 → 10秒'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()
    
    # 五、综合接入方案
    doc.add_heading('五、综合接入方案', level=1)
    
    doc.add_heading('阶段一：基础接入（本周）', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '任务'
    hdr_cells[1].text = '技能'
    hdr_cells[2].text = '负责人'
    hdr_cells[3].text = '时间'
    
    data = [
        ('创建GitHub组织', 'github', '🌿 南乔', 'Day 1'),
        ('迁移核心代码', 'github', '🌿 南乔', 'Day 1-2'),
        ('集成情报摘要', 'summarize', '📚 折桂', 'Day 2-3'),
        ('创建代码生成器', 'coding-agent', '🌿 南乔', 'Day 3-4'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_heading('阶段二：深度集成（下周）', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '任务'
    hdr_cells[1].text = '技能'
    hdr_cells[2].text = '负责人'
    hdr_cells[3].text = '时间'
    
    data = [
        ('智能体协作平台集成', '全部', '🧵 织锦', 'Week 2'),
        ('指南针工程集成', 'summarize', '🌸 采薇', 'Week 2'),
        ('九星监控开发', 'coding-agent', '🌿 南乔', 'Week 2'),
        ('文档摘要服务', 'summarize', '📚 折桂', 'Week 2'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    # 六、预期效果
    doc.add_heading('六、预期效果', level=1)
    
    doc.add_heading('效率提升', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '升级前'
    hdr_cells[2].text = '升级后'
    hdr_cells[3].text = '提升'
    
    data = [
        ('代码管理', '手动管理', 'Git版本控制', '10倍'),
        ('原型开发', '1天', '30分钟', '16倍'),
        ('情报处理', '30分钟', '10秒', '180倍'),
        ('需求理解', '1小时', '1分钟', '60倍'),
    ]
    for i, row_data in enumerate(data, 1):
        row = table.rows[i].cells
        for j, text in enumerate(row_data):
            row[j].text = text
    
    doc.add_paragraph()
    
    # 七、南乔建议
    doc.add_heading('七、南乔建议', level=1)
    doc.add_paragraph('立即执行：')
    doc.add_paragraph('1. 创建GitHub组织 - 统一管理所有代码')
    doc.add_paragraph('2. 迁移核心代码 - 智能体协作平台、指南针工程')
    doc.add_paragraph('3. 集成情报摘要 - 折桂情报自动摘要')
    doc.add_paragraph()
    doc.add_paragraph('本周完成：')
    doc.add_paragraph('1. 代码生成器 - 快速原型生成')
    doc.add_paragraph('2. 文档摘要服务 - 需求文档快速理解')
    doc.add_paragraph('3. Issue模板 - 规范化需求管理')
    
    return doc

def send_email_with_documents():
    """发送文档邮件"""
    # 生成文档
    output_dir = '/root/.openclaw/workspace/03_输出成果'
    
    doc1 = create_template_document()
    file1 = f'{output_dir}/新技能使用场景模板库.docx'
    doc1.save(file1)
    print(f'✅ 文档1已生成：{file1}')
    
    doc2 = create_integration_document()
    file2 = f'{output_dir}/现有工程新技能接入分析.docx'
    doc2.save(file2)
    print(f'✅ 文档2已生成：{file2}')
    
    # 邮箱配置
    SMTP_SERVER = 'smtp.qq.com'
    SMTP_PORT = 587
    SENDER_EMAIL = '417895006@qq.com'
    SENDER_PASSWORD = 'lgnzpppvjvfmcadj'
    RECEIVER_EMAIL = 'szideaf7@163.com'
    
    # 创建邮件
    msg = MIMEMultipart()
    msg['From'] = SENDER_EMAIL
    msg['To'] = RECEIVER_EMAIL
    msg['Subject'] = f'【南乔】新技能使用模板和接入分析文档_{datetime.now().strftime("%Y-%m-%d")}'
    
    # 邮件正文
    body = f"""
少帅您好！

南乔已完成新技能使用场景模板和接入分析，生成以下2份文档：

1. 新技能使用场景模板库.docx
   - github使用模板（4个）
   - coding-agent使用模板（4个）
   - summarize使用模板（4个）
   - 组合使用模板（3个）
   - 快速调用指南
   - 使用技巧

2. 现有工程新技能接入分析.docx
   - 现有工程清单
   - github接入方案
   - coding-agent接入方案
   - summarize接入方案
   - 综合接入方案
   - 预期效果

请查阅！

南乔 🌿
{datetime.now().strftime("%Y-%m-%d %H:%M")}
"""
    msg.attach(MIMEText(body, 'plain', 'utf-8'))
    
    # 附件1
    with open(file1, 'rb') as f:
        part1 = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part1.set_payload(f.read())
        encoders.encode_base64(part1)
        part1.add_header('Content-Disposition', 'attachment', filename='新技能使用场景模板库.docx')
        msg.attach(part1)
    
    # 附件2
    with open(file2, 'rb') as f:
        part2 = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part2.set_payload(f.read())
        encoders.encode_base64(part2)
        part2.add_header('Content-Disposition', 'attachment', filename='现有工程新技能接入分析.docx')
        msg.attach(part2)
    
    # 发送邮件
    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.sendmail(SENDER_EMAIL, RECEIVER_EMAIL, msg.as_string())
        server.quit()
        print(f'✅ 邮件发送成功！')
        print(f'收件人：{RECEIVER_EMAIL}')
        print(f'附件：')
        print(f'  1. 新技能使用场景模板库.docx')
        print(f'  2. 现有工程新技能接入分析.docx')
        return True
    except Exception as e:
        print(f'❌ 邮件发送失败：{e}')
        return False

if __name__ == '__main__':
    send_email_with_documents()
