#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
方案导出模块
- 支持Word文档导出
- 支持Excel表格导出
- 支持Markdown导出

Author: 南乔
Date: 2026-03-14
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from datetime import datetime
from typing import List, Dict, Optional
from dataclasses import dataclass
import json
import os

# ==================== 数据结构 ====================
@dataclass
class DiscussionTurn:
    """讨论轮次"""
    turn_id: int
    speaker: str
    speaker_name: str
    speaker_role: str
    content: str
    timestamp: str
    is_challenging: bool = False
    reply_to: str = ""


@dataclass
class DiscussionSummary:
    """讨论摘要"""
    task: str
    start_time: str
    end_time: str
    total_turns: int
    consensus_level: int
    participants: List[str]
    key_points: List[str]
    risks: List[str]
    decisions: List[str]


# ==================== Word文档导出 ====================
class WordExporter:
    """Word文档导出器"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
    
    def export_discussion(self, turns: List[DiscussionTurn], summary: DiscussionSummary, output_path: str):
        """
        导出讨论记录为Word文档
        
        Args:
            turns: 讨论轮次列表
            summary: 讨论摘要
            output_path: 输出文件路径
        """
        # 封面
        self._add_cover(summary)
        
        # 摘要
        self._add_summary(summary)
        
        # 参与者
        self._add_participants(turns)
        
        # 讨论记录
        self._add_discussion(turns)
        
        # 关键结论
        self._add_conclusions(summary)
        
        # 风险清单
        self._add_risks(summary)
        
        # 保存
        self.doc.save(output_path)
    
    def _add_cover(self, summary: DiscussionSummary):
        """添加封面"""
        # 标题
        title = self.doc.add_heading('智能体协作讨论报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"任务：{summary.task}")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x16, 0x5D, 0xFF)
        
        # 时间
        time_para = self.doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_para.add_run(f"生成时间：{summary.end_time}")
        
        self.doc.add_page_break()
    
    def _add_summary(self, summary: DiscussionSummary):
        """添加摘要"""
        self.doc.add_heading('一、讨论摘要', level=1)
        
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        data = [
            ('任务主题', summary.task),
            ('讨论时间', f"{summary.start_time} ~ {summary.end_time}"),
            ('讨论轮次', str(summary.total_turns)),
            ('共识度', f"{summary.consensus_level}%"),
            ('参与角色', ', '.join(summary.participants))
        ]
        
        for i, (key, value) in enumerate(data):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = value
        
        self.doc.add_paragraph()
    
    def _add_participants(self, turns: List[DiscussionTurn]):
        """添加参与者列表"""
        self.doc.add_heading('二、参与角色', level=1)
        
        # 统计每个角色的发言次数
        role_stats = {}
        for turn in turns:
            if turn.speaker not in ['user', 'nanqiao']:
                key = (turn.speaker, turn.speaker_name, turn.speaker_role)
                role_stats[key] = role_stats.get(key, 0) + 1
        
        table = self.doc.add_table(rows=len(role_stats) + 1, cols=3)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['角色名称', '角色职责', '发言次数']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
        # 数据
        for i, ((agent_id, name, role), count) in enumerate(role_stats.items(), 1):
            table.rows[i].cells[0].text = name
            table.rows[i].cells[1].text = role
            table.rows[i].cells[2].text = str(count)
        
        self.doc.add_paragraph()
    
    def _add_discussion(self, turns: List[DiscussionTurn]):
        """添加讨论记录"""
        self.doc.add_heading('三、讨论记录', level=1)
        
        for turn in turns:
            # 发言者
            para = self.doc.add_paragraph()
            run = para.add_run(f"【{turn.speaker_name}】")
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x16, 0x5D, 0xFF)
            
            # 时间戳
            para.add_run(f" {turn.timestamp}")
            
            # 内容
            content_para = self.doc.add_paragraph()
            if turn.is_challenging:
                run = content_para.add_run("⚠️ ")
                run.font.color.rgb = RGBColor(0xF5, 0x6C, 0x6C)
            content_para.add_run(turn.content)
            
            # 回复标记
            if turn.reply_to:
                reply_para = self.doc.add_paragraph()
                reply_para.add_run(f"↩️ 回复 {turn.reply_to}").font.size = Pt(9)
            
            self.doc.add_paragraph()
    
    def _add_conclusions(self, summary: DiscussionSummary):
        """添加关键结论"""
        self.doc.add_heading('四、关键结论', level=1)
        
        for i, point in enumerate(summary.key_points, 1):
            para = self.doc.add_paragraph()
            para.add_run(f"{i}. {point}")
    
    def _add_risks(self, summary: DiscussionSummary):
        """添加风险清单"""
        self.doc.add_heading('五、风险清单', level=1)
        
        if not summary.risks:
            self.doc.add_paragraph("本次讨论未识别出重大风险。")
            return
        
        table = self.doc.add_table(rows=len(summary.risks) + 1, cols=2)
        table.style = 'Table Grid'
        
        # 表头
        table.rows[0].cells[0].text = '序号'
        table.rows[0].cells[1].text = '风险描述'
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        
        # 数据
        for i, risk in enumerate(summary.risks, 1):
            table.rows[i].cells[0].text = str(i)
            table.rows[i].cells[1].text = risk


# ==================== Excel表格导出 ====================
class ExcelExporter:
    """Excel表格导出器"""
    
    def __init__(self):
        self.wb = Workbook()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置样式"""
        self.header_fill = PatternFill(
            start_color='165DFF',
            end_color='165DFF',
            fill_type='solid'
        )
        self.header_font = Font(bold=True, color='FFFFFF', size=11)
        self.border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
    
    def export_discussion(self, turns: List[DiscussionTurn], summary: DiscussionSummary, output_path: str):
        """
        导出讨论记录为Excel表格
        
        Args:
            turns: 讨论轮次列表
            summary: 讨论摘要
            output_path: 输出文件路径
        """
        # 摘要表
        ws_summary = self.wb.active
        ws_summary.title = '讨论摘要'
        self._add_summary_sheet(ws_summary, summary)
        
        # 讨论记录表
        ws_discussion = self.wb.create_sheet('讨论记录')
        self._add_discussion_sheet(ws_discussion, turns)
        
        # 风险清单表
        ws_risks = self.wb.create_sheet('风险清单')
        self._add_risks_sheet(ws_risks, summary)
        
        # 保存
        self.wb.save(output_path)
    
    def _add_summary_sheet(self, ws, summary: DiscussionSummary):
        """添加摘要表"""
        # 标题
        ws.merge_cells('A1:B1')
        ws['A1'] = '智能体协作讨论摘要'
        ws['A1'].font = Font(bold=True, size=16)
        ws['A1'].alignment = Alignment(horizontal='center')
        
        # 数据
        data = [
            ('任务主题', summary.task),
            ('开始时间', summary.start_time),
            ('结束时间', summary.end_time),
            ('讨论轮次', summary.total_turns),
            ('共识度', f"{summary.consensus_level}%"),
            ('参与角色', ', '.join(summary.participants))
        ]
        
        for i, (key, value) in enumerate(data, 3):
            ws[f'A{i}'] = key
            ws[f'B{i}'] = value
            ws[f'A{i}'].font = Font(bold=True)
        
        # 调整列宽
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 50
    
    def _add_discussion_sheet(self, ws, turns: List[DiscussionTurn]):
        """添加讨论记录表"""
        # 表头
        headers = ['序号', '发言者', '角色', '内容', '时间', '是否质疑', '回复对象']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.fill = self.header_fill
            cell.font = self.header_font
            cell.alignment = Alignment(horizontal='center')
            cell.border = self.border
        
        # 数据
        for row, turn in enumerate(turns, 2):
            ws.cell(row=row, column=1, value=turn.turn_id).border = self.border
            ws.cell(row=row, column=2, value=turn.speaker_name).border = self.border
            ws.cell(row=row, column=3, value=turn.speaker_role).border = self.border
            ws.cell(row=row, column=4, value=turn.content).border = self.border
            ws.cell(row=row, column=5, value=turn.timestamp).border = self.border
            ws.cell(row=row, column=6, value='是' if turn.is_challenging else '否').border = self.border
            ws.cell(row=row, column=7, value=turn.reply_to).border = self.border
        
        # 调整列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 60
        ws.column_dimensions['E'].width = 20
        ws.column_dimensions['F'].width = 10
        ws.column_dimensions['G'].width = 12
    
    def _add_risks_sheet(self, ws, summary: DiscussionSummary):
        """添加风险清单表"""
        # 表头
        headers = ['序号', '风险描述', '风险等级', '应对建议']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.fill = self.header_fill
            cell.font = self.header_font
            cell.alignment = Alignment(horizontal='center')
            cell.border = self.border
        
        # 数据
        for row, risk in enumerate(summary.risks, 2):
            ws.cell(row=row, column=1, value=row - 1).border = self.border
            ws.cell(row=row, column=2, value=risk).border = self.border
            ws.cell(row=row, column=3, value='中').border = self.border
            ws.cell(row=row, column=4, value='待补充').border = self.border
        
        # 调整列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 30


# ==================== Markdown导出 ====================
class MarkdownExporter:
    """Markdown文档导出器"""
    
    def export_discussion(self, turns: List[DiscussionTurn], summary: DiscussionSummary) -> str:
        """
        导出讨论记录为Markdown文本
        
        Returns:
            Markdown格式文本
        """
        lines = []
        
        # 标题
        lines.append(f"# 智能体协作讨论报告")
        lines.append("")
        lines.append(f"**任务：{summary.task}**")
        lines.append("")
        
        # 摘要
        lines.append("## 一、讨论摘要")
        lines.append("")
        lines.append(f"| 项目 | 内容 |")
        lines.append(f"|------|------|")
        lines.append(f"| 任务主题 | {summary.task} |")
        lines.append(f"| 讨论时间 | {summary.start_time} ~ {summary.end_time} |")
        lines.append(f"| 讨论轮次 | {summary.total_turns} |")
        lines.append(f"| 共识度 | {summary.consensus_level}% |")
        lines.append(f"| 参与角色 | {', '.join(summary.participants)} |")
        lines.append("")
        
        # 讨论记录
        lines.append("## 二、讨论记录")
        lines.append("")
        for turn in turns:
            challenge_mark = "⚠️ " if turn.is_challenging else ""
            lines.append(f"### 【{turn.speaker_name}】")
            lines.append("")
            lines.append(f"{challenge_mark}{turn.content}")
            lines.append("")
            if turn.reply_to:
                lines.append(f"> 回复 @{turn.reply_to}")
                lines.append("")
        
        # 关键结论
        if summary.key_points:
            lines.append("## 三、关键结论")
            lines.append("")
            for i, point in enumerate(summary.key_points, 1):
                lines.append(f"{i}. {point}")
            lines.append("")
        
        # 风险清单
        if summary.risks:
            lines.append("## 四、风险清单")
            lines.append("")
            for i, risk in enumerate(summary.risks, 1):
                lines.append(f"{i}. {risk}")
            lines.append("")
        
        return "\n".join(lines)


# ==================== 导出API ====================
class ExportAPI:
    """导出API接口"""
    
    def __init__(self, output_dir: str = "/root/.openclaw/workspace/03_输出成果/导出文档"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        self.word_exporter = WordExporter()
        self.excel_exporter = ExcelExporter()
        self.markdown_exporter = MarkdownExporter()
    
    def export_all(self, turns: List[DiscussionTurn], summary: DiscussionSummary) -> Dict[str, str]:
        """
        导出所有格式
        
        Returns:
            文件路径字典
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"讨论报告_{timestamp}"
        
        # Word
        word_path = os.path.join(self.output_dir, f"{base_name}.docx")
        self.word_exporter.export_discussion(turns, summary, word_path)
        
        # Excel
        excel_path = os.path.join(self.output_dir, f"{base_name}.xlsx")
        self.excel_exporter.export_discussion(turns, summary, excel_path)
        
        # Markdown
        md_content = self.markdown_exporter.export_discussion(turns, summary)
        md_path = os.path.join(self.output_dir, f"{base_name}.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return {
            'word': word_path,
            'excel': excel_path,
            'markdown': md_path
        }
    
    def export_word(self, turns: List[DiscussionTurn], summary: DiscussionSummary) -> str:
        """导出Word文档"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_path = os.path.join(self.output_dir, f"讨论报告_{timestamp}.docx")
        self.word_exporter.export_discussion(turns, summary, word_path)
        return word_path
    
    def export_excel(self, turns: List[DiscussionTurn], summary: DiscussionSummary) -> str:
        """导出Excel表格"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_path = os.path.join(self.output_dir, f"讨论报告_{timestamp}.xlsx")
        self.excel_exporter.export_discussion(turns, summary, excel_path)
        return excel_path
    
    def export_markdown(self, turns: List[DiscussionTurn], summary: DiscussionSummary) -> str:
        """导出Markdown文档"""
        return self.markdown_exporter.export_discussion(turns, summary)


# ==================== 测试 ====================
if __name__ == '__main__':
    # 创建测试数据
    test_turns = [
        DiscussionTurn(
            turn_id=1, speaker='caiwei', speaker_name='采薇', speaker_role='需求分析专家',
            content='【需求分析】我梳理了核心需求。主要功能点已识别，建议进行技术可行性评估。',
            timestamp='2026-03-14 10:00:00', is_challenging=False
        ),
        DiscussionTurn(
            turn_id=2, speaker='zhijin', speaker_name='织锦', speaker_role='架构设计师',
            content='【架构设计】基于需求，推荐微服务架构。技术栈：Spring Cloud + K8s + PostgreSQL。',
            timestamp='2026-03-14 10:01:00', is_challenging=False
        ),
        DiscussionTurn(
            turn_id=3, speaker='zhutai', speaker_name='筑台', speaker_role='售前工程师',
            content='【成本质疑】微服务架构成本较高，建议评估必要性。可采用单体架构降低初期投入。',
            timestamp='2026-03-14 10:02:00', is_challenging=True, reply_to='织锦'
        ),
    ]
    
    test_summary = DiscussionSummary(
        task='智能客服系统需求分析',
        start_time='2026-03-14 10:00:00',
        end_time='2026-03-14 10:10:00',
        total_turns=3,
        consensus_level=75,
        participants=['采薇', '织锦', '筑台'],
        key_points=[
            '核心需求已识别，包含智能问答、多轮对话、知识库管理',
            '推荐微服务架构，但需评估成本',
            '建议敏捷迭代，首期交付核心功能'
        ],
        risks=[
            '微服务架构成本较高',
            '开发周期可能延长'
        ],
        decisions=[
            '采纳微服务架构方案',
            '首期交付核心功能'
        ]
    )
    
    # 测试导出
    api = ExportAPI()
    paths = api.export_all(test_turns, test_summary)
    
    print("✅ 导出完成！")
    print(f"Word文档: {paths['word']}")
    print(f"Excel表格: {paths['excel']}")
    print(f"Markdown: {paths['markdown']}")
