#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多Agent协作框架 V15.2 - 产物生成集成版
- V14基础功能
- V15质疑增强机制
- V15.1 并行调度优化（工尺+天工设计）
  * 多Agent并行调用API
  * 智能缓存机制
  * 超时熔断+降级策略
  * 保留串行模式作为回退
- V15.2 产物生成集成（天工设计）
  * 产物生成引擎集成
  * 质量审核引擎集成
  * 模板管理API集成
  * 讨论结束后一键生成产物

Author: 南乔 + 工尺 + 天工
Date: 2026-03-19
"""

from flask import Flask, render_template_string, jsonify, request, send_file, Response
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
from enum import Enum
import threading
import time
import os
import requests
import json
import glob
import sys

# ==================== V15.2模块导入 ====================
sys.path.insert(0, '/root/.openclaw/workspace/03_输出成果/天工产出/V15.2')
try:
    from api_routes import register_api_routes
    V15_2_LOADED = True
    print("[V15.2] 模块加载成功：产物生成、质量审核、模板管理API")
except ImportError as e:
    V15_2_LOADED = False
    print(f"[WARN] V15.2模块导入失败: {e}")

# 导入任务调度模块
TASK_API_LOADED = False
task_bp = None
try:
    from task_api import task_bp
    TASK_API_LOADED = True
    print("[V15.2] 任务调度模块加载成功")
except ImportError as e:
    print(f"[WARN] 任务调度模块导入失败: {e}")

# 导入导出模块
from export_module import ExportAPI, DiscussionTurn, DiscussionSummary

# 导入智能调度模块
from intent_scheduler import IntelligentScheduler, TASK_TYPES

# 导入质疑增强模块
from challenge_enhancement import ChallengeEnhancementManager, ChallengeIntensity, ChallengeCategory

# V15模块延迟导入（避免循环导入）
V15_MODULES_LOADED = False
MultiRoundDiscussion = None
NanqiaoIntentAnalyzer = None
MeetingMinutesGenerator = None

# ==================== 质疑增强配置 ====================
CHALLENGE_ENHANCEMENT_ENABLED = True  # 质疑增强开关（可配置）
challenge_manager = None  # 全局质疑增强管理器

# ==================== 并行调度配置 V15.1 ====================
PARALLEL_MODE_ENABLED = True  # 并行模式开关（可配置）
PARALLEL_TIMEOUT = 30  # 并行超时时间（秒）
PARALLEL_MAX_WORKERS = 5  # 最大并行数
PARALLEL_CACHE_ENABLED = True  # 缓存开关
PARALLEL_FALLBACK_ENABLED = True  # 降级开关

def get_challenge_manager():
    """获取质疑增强管理器（单例）"""
    global challenge_manager
    if challenge_manager is None:
        challenge_manager = ChallengeEnhancementManager(enable_enhancement=CHALLENGE_ENHANCEMENT_ENABLED)
    return challenge_manager

def load_v15_modules():
    """延迟加载V15模块"""
    global V15_MODULES_LOADED, NanqiaoIntentAnalyzer, MeetingMinutesGenerator, MultiRoundDiscussion
    if V15_MODULES_LOADED:
        return True
    try:
        from nanqiao_intent_analyzer import NanqiaoIntentAnalyzer as NIA
        from meeting_minutes_generator import MeetingMinutesGenerator as MMG
        from multi_round_discussion import MultiRoundDiscussion as MRD
        NanqiaoIntentAnalyzer = NIA
        MeetingMinutesGenerator = MMG
        MultiRoundDiscussion = MRD
        V15_MODULES_LOADED = True
        print("[V15] 模块加载成功：意图分析、多轮讨论、会议纪要")
        return True
    except ImportError as e:
        print(f"[WARN] V15模块导入失败: {e}")
        import traceback
        traceback.print_exc()
        return False

# ==================== 知识库加载 ====================
class KnowledgeBase:
    """知识库管理器"""

    def __init__(self):
        self.base_path = '/root/.openclaw/skills/compass-shared/knowledge'
        self.cache = {}

    def load_file(self, relative_path: str) -> str:
        """加载知识文件"""
        if relative_path in self.cache:
            return self.cache[relative_path]

        full_path = os.path.join(self.base_path, relative_path)
        if os.path.exists(full_path):
            with open(full_path, 'r', encoding='utf-8') as f:
                content = f.read()
                self.cache[relative_path] = content
                return content
        return ""

    def get_telecom_business(self) -> str:
        """获取电信业务知识"""
        return self.load_file('industries/telecom/business/business-overview.md')

    def get_telecom_systems(self) -> str:
        """获取电信系统知识"""
        return self.load_file('industries/telecom/systems/systems-overview.md')

    def get_telecom_scenarios(self) -> str:
        """获取电信场景知识"""
        return self.load_file('industries/telecom/scenarios/scenarios-overview.md')

    def get_ai_capabilities(self) -> str:
        """获取AI能力知识"""
        return self.load_file('industries/telecom/ai-capabilities/ai-parameters.md')

    def get_business_rules(self) -> str:
        """获取业务规则"""
        return self.load_file('industries/telecom/business/business-rules.md')

    def get_interface_spec(self) -> str:
        """获取接口规范"""
        return self.load_file('industries/telecom/systems/interface-specification.md')

    def get_knowledge_for_task(self, task_type: str) -> Dict[str, str]:
        """根据任务类型获取相关知识"""
        knowledge = {}

        # 需求类任务
        if task_type.startswith('REQ'):
            knowledge['business'] = self.get_telecom_business()[:2000]
            knowledge['scenarios'] = self.get_telecom_scenarios()[:2000]

        # 设计类任务
        elif task_type.startswith('DES'):
            knowledge['systems'] = self.get_telecom_systems()[:2000]
            knowledge['ai'] = self.get_ai_capabilities()[:2000]

        # 项目管理类任务
        elif task_type.startswith('PM'):
            knowledge['business_rules'] = self.get_business_rules()[:2000]

        # 技术类任务
        elif task_type.startswith('DEV'):
            knowledge['systems'] = self.get_telecom_systems()[:2000]
            knowledge['interface'] = self.get_interface_spec()[:2000]

        return knowledge

    def list_available_knowledge(self) -> List[str]:
        """列出所有可用知识"""
        files = []
        for root, dirs, filenames in os.walk(self.base_path):
            for f in filenames:
                if f.endswith('.md'):
                    rel_path = os.path.relpath(os.path.join(root, f), self.base_path)
                    files.append(rel_path)
        return sorted(files)

# 初始化知识库
knowledge_base = KnowledgeBase()

# 导入场景模板库
from scene_templates import SceneTemplateAPI
scene_api = SceneTemplateAPI()

# 导入成本估算模型
from cost_estimator import CostEstimator
cost_estimator = CostEstimator()

app = Flask(__name__)

# ==================== 注册V15.2 API路由 ====================
if V15_2_LOADED:
    try:
        register_api_routes(app)
        print("[V15.2] API路由注册成功")
    except Exception as e:
        print(f"[WARN] V15.2 API路由注册失败: {e}")

# 注册任务调度API
if TASK_API_LOADED and task_bp:
    try:
        app.register_blueprint(task_bp)
        print("[V15.2] 任务调度API注册成功")
    except Exception as e:
        print(f"[WARN] 任务调度API注册失败: {e}")

# 导出API
export_api = ExportAPI()

# 智能调度器
intelligent_scheduler = IntelligentScheduler()

# Agent ID → 中文名称映射
AGENT_NAMES = {
    'caiwei': '采薇',
    'zhijin': '织锦',
    'zhutai': '筑台',
    'chengcai': '呈彩',
    'gongchi': '工尺',
    'yuheng': '玉衡',
    'fuyao': '扶摇',
    'nanqiao': '南乔',
}

# ==================== 百度千帆API配置 ====================
QIANFAN_API_KEY = os.environ.get('QIANFAN_API_KEY', 'bce-v3/ALTAKSP-14YyizFlbkiA0cKHpR4ya/b2b8db94725048693a15c4479c980c848a6a4c19')

def call_qianfan(system_prompt: str, user_message: str, temperature: float = 0.7) -> Optional[str]:
    global total_tokens_used
    if not QIANFAN_API_KEY:
        return None
    try:
        url = "https://qianfan.baidubce.com/v2/coding/chat/completions"
        headers = {
            "Authorization": f"Bearer {QIANFAN_API_KEY}",
            "Content-Type": "application/json"
        }
        combined_message = f"{system_prompt}\n\n---\n\n{user_message}"
        payload = {
            "model": "qianfan-code-latest",
            "messages": [{"role": "user", "content": combined_message}],
            "temperature": temperature,
            "top_p": 0.9
        }
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        result = response.json()
        
        # 调试日志
        print(f"[API] 响应状态: {response.status_code}", flush=True)
        if 'error' in result:
            print(f"[API] 错误: {result['error']}", flush=True)
        if 'model' in result:
            print(f"[API] 模型: {result['model']}", flush=True)
        
        # 统计token使用量
        if 'usage' in result:
            usage = result['usage']
            total_tokens_used += usage.get('total_tokens', 0)
            print(f"[TOKEN] 本次: {usage.get('total_tokens', 0)}, 累计: {total_tokens_used}", flush=True)
        
        if 'choices' in result and len(result['choices']) > 0:
            return result['choices'][0]['message']['content']
        return None
    except Exception as e:
        print(f"千帆API调用失败: {e}", flush=True)
        return None


def call_qianfan_stream(system_prompt: str, user_message: str, temperature: float = 0.7):
    """流式调用千帆API"""
    if not QIANFAN_API_KEY:
        yield "error: API Key未配置"
        return

    try:
        url = "https://qianfan.baidubce.com/v2/coding/chat/completions"
        headers = {
            "Authorization": f"Bearer {QIANFAN_API_KEY}",
            "Content-Type": "application/json"
        }
        combined_message = f"{system_prompt}\n\n---\n\n{user_message}"
        payload = {
            "model": "qianfan-code-latest",
            "messages": [{"role": "user", "content": combined_message}],
            "temperature": temperature,
            "top_p": 0.9,
            "stream": True  # 启用流式输出
        }

        response = requests.post(url, headers=headers, json=payload, timeout=60, stream=True)

        for line in response.iter_lines():
            if line:
                line_text = line.decode('utf-8')
                if line_text.startswith('data: '):
                    data_str = line_text[6:]  # 去掉 'data: ' 前缀
                    if data_str.strip() and data_str != '[DONE]':
                        try:
                            data = json.loads(data_str)
                            if 'choices' in data and len(data['choices']) > 0:
                                delta = data['choices'][0].get('delta', {})
                                if 'content' in delta:
                                    yield delta['content']
                        except json.JSONDecodeError:
                            continue
    except Exception as e:
        print(f"千帆API流式调用失败: {e}")
        yield f"error: {str(e)}"


# ==================== 并行调度模块 V15.1 ====================
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock

# 缓存目录
PARALLEL_CACHE_DIR = "/root/.openclaw/workspace/03_输出成果/agent_cache"
os.makedirs(PARALLEL_CACHE_DIR, exist_ok=True)

# Token统计锁
parallel_token_lock = Lock()
parallel_total_tokens = 0


def get_parallel_cache_key(system_prompt: str, user_message: str) -> str:
    """生成缓存键"""
    content = f"{system_prompt}|{user_message}"
    return hashlib.md5(content.encode()).hexdigest()


def get_parallel_cached_response(cache_key: str) -> Optional[str]:
    """获取缓存响应"""
    cache_file = os.path.join(PARALLEL_CACHE_DIR, f"{cache_key}.json")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                # 缓存有效期24小时
                if time.time() - data.get("timestamp", 0) < 86400:
                    return data.get("response")
        except:
            pass
    return None


def save_parallel_cached_response(cache_key: str, response: str):
    """保存缓存响应"""
    cache_file = os.path.join(PARALLEL_CACHE_DIR, f"{cache_key}.json")
    try:
        with open(cache_file, "w", encoding="utf-8") as f:
            json.dump({
                "response": response,
                "timestamp": time.time()
            }, f, ensure_ascii=False)
    except:
        pass


def call_agent_parallel(agent_id: str, system_prompt: str, user_message: str, 
                        temperature: float = 0.7) -> Tuple[str, str, bool]:
    """
    并行调用单个Agent
    
    Returns:
        Tuple[response, agent_id, from_cache]
    """
    global parallel_total_tokens
    
    # 检查缓存
    if PARALLEL_CACHE_ENABLED:
        cache_key = get_parallel_cache_key(system_prompt, user_message)
        cached = get_parallel_cached_response(cache_key)
        if cached:
            print(f"[{agent_id}] 命中缓存", flush=True)
            return cached, agent_id, True
    
    if not QIANFAN_API_KEY:
        return None, agent_id, False
    
    try:
        headers = {
            "Authorization": f"Bearer {QIANFAN_API_KEY}",
            "Content-Type": "application/json"
        }
        combined_message = f"{system_prompt}\n\n---\n\n{user_message}"
        payload = {
            "model": "qianfan-code-latest",
            "messages": [{"role": "user", "content": combined_message}],
            "temperature": temperature,
            "top_p": 0.9
        }
        
        start_time = time.time()
        response = requests.post(
            "https://qianfan.baidubce.com/v2/coding/chat/completions", 
            headers=headers, json=payload, timeout=PARALLEL_TIMEOUT
        )
        elapsed = time.time() - start_time
        
        print(f"[{agent_id}] API耗时: {elapsed:.1f}秒", flush=True)
        
        result = response.json()
        
        # 统计token
        if 'usage' in result:
            with parallel_token_lock:
                parallel_total_tokens += result['usage'].get('total_tokens', 0)
        
        if 'choices' in result and len(result['choices']) > 0:
            content = result['choices'][0]['message']['content']
            # 保存缓存
            if PARALLEL_CACHE_ENABLED:
                save_parallel_cached_response(cache_key, content)
            return content, agent_id, False
            
        return None, agent_id, False
        
    except requests.exceptions.Timeout:
        print(f"[{agent_id}] API超时（{PARALLEL_TIMEOUT}秒）", flush=True)
        return None, agent_id, False
    except Exception as e:
        print(f"[{agent_id}] API调用失败: {e}", flush=True)
        return None, agent_id, False


def dispatch_agents_parallel(agents: List[str], task: str, 
                             memory_context: str = "") -> Dict[str, Tuple[str, bool]]:
    """
    并行调度多个Agent
    
    Args:
        agents: Agent ID列表
        task: 任务描述
        memory_context: 对话上下文
    
    Returns:
        Dict[agent_id, Tuple[response, from_cache]]
    """
    if not PARALLEL_MODE_ENABLED:
        return {}
    
    results = {}
    
    # 讨论维度提示
    discussion_prompts = {
        'caiwei': '请从需求拆解、用户故事、验收标准角度分析',
        'zhijin': '请从技术选型、架构方案、扩展性角度分析',
        'zhutai': '请从成本估算、资源需求角度分析',
        'chengcai': '请从UI设计、用户体验角度分析',
        'gongchi': '请从接口设计、模块划分角度分析',
        'yuheng': '请从进度计划、风险管控、里程碑角度分析',
        'fuyao': '请从整体平衡、决策建议、优先级排序角度分析',
        'zhegui': '请从人员配置、技术资源、知识库支持角度分析',
        'nanqiao': '请从整体协调角度分析'
    }
    
    # 构建并行任务
    tasks = []
    for agent_id in agents:
        if agent_id not in AGENTS:
            continue
        agent = AGENTS[agent_id]
        role_prompt = discussion_prompts.get(agent_id, '请发表专业意见')
        
        user_message = f"""当前任务：{task}

已有讨论：
{memory_context}

{role_prompt}。如果有不同意见请明确提出质疑。"""
        
        tasks.append({
            "agent_id": agent_id,
            "system_prompt": agent.get_system_prompt(),
            "user_message": user_message
        })
    
    # 并行执行
    start_time = time.time()
    print(f"[并行调度] 启动 {len(tasks)} 个Agent并行响应...", flush=True)
    
    with ThreadPoolExecutor(max_workers=min(len(tasks), PARALLEL_MAX_WORKERS)) as executor:
        futures = {
            executor.submit(
                call_agent_parallel, 
                t['agent_id'],
                t['system_prompt'], 
                t['user_message']
            ): t['agent_id'] 
            for t in tasks
        }
        
        for future in as_completed(futures):
            try:
                response, agent_id, from_cache = future.result()
                if response:
                    results[agent_id] = (response, from_cache)
            except Exception as e:
                print(f"[并行调度] 任务异常: {e}", flush=True)
    
    elapsed = time.time() - start_time
    print(f"[并行调度] 完成，总耗时: {elapsed:.1f}秒", flush=True)
    
    return results


# ==================== Agent角色定义（完整9个）====================
class Stance(Enum):
    COST_FIRST = "成本优先"
    QUALITY_FIRST = "质量优先"
    SCHEDULE_FIRST = "进度优先"
    BALANCED = "平衡折中"

@dataclass
class AgentPersona:
    agent_id: str
    name: str
    role: str
    emoji: str
    stance: Stance
    expertise: List[str]
    concern_points: List[str]
    color: str

    def get_system_prompt(self) -> str:
        stance_desc = {
            Stance.COST_FIRST: "你非常关注成本控制，反对不必要的投入，倾向于选择性价比高的方案。",
            Stance.QUALITY_FIRST: "你非常关注技术质量和用户体验，愿意为更好的方案投入更多资源。",
            Stance.SCHEDULE_FIRST: "你非常关注项目进度和风险控制，反对可能导致延期的方案。",
            Stance.BALANCED: "你综合考虑各方面因素，追求平衡的方案。"
        }
        return f"""你是{self.name}，角色是{self.role}。

你的立场：{stance_desc[self.stance]}

你的专业领域：{', '.join(self.expertise)}
你最关注的问题：{', '.join(self.concern_points)}

重要规则：
1. 必须基于你的立场和专业领域发言
2. 如果看到其他人的方案与你的立场冲突，要明确提出质疑
3. 回复要专业、简洁，控制在80字以内
4. 使用【角色名】开头"""


# 定义所有9个Agent
AGENTS = {
    'nanqiao': AgentPersona(
        agent_id='nanqiao', name='南乔', role='主控Agent', emoji='🌿',
        stance=Stance.BALANCED,
        expertise=['需求整理', '任务协调', '进度跟踪'],
        concern_points=['需求范围', '进度风险'],
        color='#9C27B0'
    ),
    'caiwei': AgentPersona(
        agent_id='caiwei', name='采薇', role='需求分析专家', emoji='🌸',
        stance=Stance.QUALITY_FIRST,
        expertise=['需求调研', '用户故事', '验收标准'],
        concern_points=['需求完整性', '用户体验'],
        color='#409EFF'
    ),
    'zhijin': AgentPersona(
        agent_id='zhijin', name='织锦', role='架构设计师', emoji='🧵',
        stance=Stance.QUALITY_FIRST,
        expertise=['架构设计', '技术选型', '系统设计'],
        concern_points=['技术可行性', '扩展性', '性能'],
        color='#67C23A'
    ),
    'zhutai': AgentPersona(
        agent_id='zhutai', name='筑台', role='售前工程师', emoji='🏗️',
        stance=Stance.COST_FIRST,
        expertise=['成本评估', '方案报价', 'ROI计算'],
        concern_points=['成本过高', '预算控制'],
        color='#E6A23C'
    ),
    'gongchi': AgentPersona(
        agent_id='gongchi', name='工尺', role='详细设计师', emoji='📐',
        stance=Stance.QUALITY_FIRST,
        expertise=['详细设计', '接口设计', '数据库设计'],
        concern_points=['实现难度', '代码质量'],
        color='#607D8B'
    ),
    'yuheng': AgentPersona(
        agent_id='yuheng', name='玉衡', role='项目经理', emoji='⚖️',
        stance=Stance.SCHEDULE_FIRST,
        expertise=['项目计划', '风险管理', '进度跟踪'],
        concern_points=['进度风险', '资源冲突'],
        color='#F56C6C'
    ),
    'fuyao': AgentPersona(
        agent_id='fuyao', name='扶摇', role='总指挥', emoji='🌀',
        stance=Stance.BALANCED,
        expertise=['决策判断', '资源调配', '团队协调'],
        concern_points=['决策风险', '团队共识'],
        color='#165DFF'
    ),
    'chengcai': AgentPersona(
        agent_id='chengcai', name='呈彩', role='方案设计师', emoji='🎨',
        stance=Stance.QUALITY_FIRST,
        expertise=['方案设计', 'PPT制作', '视觉呈现'],
        concern_points=['方案完整性', '展示效果'],
        color='#FF9800'
    ),
    'zhegui': AgentPersona(
        agent_id='zhegui', name='折桂', role='资源管家', emoji='📚',
        stance=Stance.BALANCED,
        expertise=['资源管理', '知识库维护', '文档归档'],
        concern_points=['资源不足', '知识沉淀'],
        color='#00BCD4'
    )
}


# ==================== 对话记忆 ====================
@dataclass
class ConversationTurn:
    turn_id: int
    speaker: str
    speaker_name: str
    content: str
    timestamp: str
    msg_type: str = "answer"
    is_challenging: bool = False
    reply_to: str = ""

class ConversationMemory:
    def __init__(self):
        self.history: List[ConversationTurn] = []
        self.turn_count: int = 0
        self.current_task: str = ""
        self.start_time: datetime = None

    def add_turn(self, speaker: str, speaker_name: str, content: str,
                 msg_type: str = "answer", is_challenging: bool = False, reply_to: str = ""):
        self.turn_count += 1
        turn = ConversationTurn(
            turn_id=self.turn_count,
            speaker=speaker,
            speaker_name=speaker_name,
            content=content,
            timestamp=datetime.now().isoformat(),
            msg_type=msg_type,
            is_challenging=is_challenging,
            reply_to=reply_to
        )
        self.history.append(turn)
        return turn

    def get_context(self, limit: int = 8) -> str:
        if not self.history:
            return ""
        parts = []
        for turn in self.history[-limit:]:
            prefix = "【质疑】" if turn.is_challenging else ""
            parts.append(f"{prefix}{turn.speaker_name}: {turn.content}")
        return "\n".join(parts)

    def get_consensus_level(self) -> int:
        if not self.history:
            return 0
        challenges = sum(1 for t in self.history if t.is_challenging)
        total = len([t for t in self.history if t.speaker != 'user' and t.speaker != 'nanqiao'])
        if total == 0:
            return 0
        # 基础共识度70%，每条质疑扣5分，最低30%
        consensus = max(30, 70 - (challenges * 5))
        return min(100, consensus)

    def clear(self):
        self.history.clear()
        self.turn_count = 0
        self.current_task = ""
        self.start_time = None


# ==================== 智能响应生成 ====================
class IntelligentResponder:
    def __init__(self):
        self.api_configured = bool(QIANFAN_API_KEY)

    def generate_system_message(self, msg_type: str, context: dict) -> str:
        """
        动态生成系统消息
        
        Args:
            msg_type: 消息类型 (schedule|task_start|risk_debate|task_done|all_done|consensus)
            context: 上下文信息
        
        Returns:
            动态生成的系统消息
        """
        nanqiao_prompt = """你是南乔，九星智囊团的主控Agent。你的风格是：温暖、贴心、有点俏皮。

你的任务是生成简洁、专业的系统提示消息。

规则：
1. 消息要简短有力，控制在50字以内
2. 使用🌿开头
3. 不要使用【】符号
4. 语气要友好、专业"""

        if msg_type == 'schedule':
            if context.get('multi'):
                task_list = context['task_list']
                user_message = f"用户提交了多个任务：{task_list}。请生成一句话，说明已识别这些任务，将依次引导团队讨论。"
            else:
                task_name = context.get('task_name', '未知任务')
                confidence = context.get('confidence', 0.9)
                lead_agent = context.get('lead_agent', '采薇')
                participants = context.get('participants', [])
                est_time = context.get('est_time', '3-5天')
                user_message = f"用户任务：{task_name}，置信度{confidence:.0%}。主导：{lead_agent}，参与：{', '.join(participants)}。预计耗时：{est_time}。请生成一句话，说明已识别任务并开始调度。"
        
        elif msg_type == 'task_start':
            current = context.get('current', 1)
            total = context.get('total', 1)
            task_name = context.get('task_name', '')
            est_time = context.get('est_time', '3-5天')
            user_message = f"开始第{current}/{total}个任务：{task_name}，预计耗时{est_time}。请生成一句话，说明任务开始。"
        
        elif msg_type == 'risk_debate':
            user_message = "进入风险辩论环节，请各位专家质疑方案风险。请生成一句话，引导专家发言。"
        
        elif msg_type == 'task_done':
            current = context.get('current', 1)
            task_name = context.get('task_name', '')
            consensus = context.get('consensus', 0)
            next_task = context.get('next_task', '')
            if next_task:
                user_message = f"任务「{task_name}」讨论完成，共识度{consensus}%。下一个任务：{next_task}。请生成过渡提示。"
            else:
                user_message = f"任务「{task_name}」讨论完成，共识度{consensus}%。请生成完成提示。"
        
        elif msg_type == 'all_done':
            total = context.get('total', 1)
            consensus = context.get('consensus', 0)
            outputs = context.get('outputs', [])
            discussion_summary = context.get('discussion_summary', '')
            key_decisions = context.get('key_decisions', [])
            key_challenges = context.get('key_challenges', [])
            pending_challenges = context.get('pending_challenges', 0)
            
            # 构建关键决策文本
            decisions_text = ""
            if key_decisions:
                decisions_text = "\n\n📌 **关键决策：**\n"
                for i, d in enumerate(key_decisions[:3], 1):
                    decisions_text += f"{i}. {d['agent']}：{d['decision'][:50]}...\n"
            
            # 构建质疑文本
            challenges_text = ""
            if key_challenges:
                challenges_text = "\n\n⚠️ **主要质疑：**\n"
                for i, c in enumerate(key_challenges[:3], 1):
                    challenges_text += f"{i}. {c['challenger']}：{c['content'][:50]}...\n"
            
            # 待回应提示
            pending_text = ""
            pending_count = len(pending_challenges) if isinstance(pending_challenges, list) else 0
            if pending_count > 0:
                pending_text = f"\n\n⏳ **待回应质疑：{pending_count}条**\n请关注未解决的质疑点。"
            
            user_message = f"""全部{total}个任务讨论完成，整体共识度{consensus}%。

决策方案产出：{', '.join(outputs)}

讨论摘要：
{discussion_summary}{decisions_text}{challenges_text}{pending_text}

请基于以上实际讨论内容，生成多维度总结：

【需求维度】根据讨论，明确哪些核心功能？验收标准是什么？
【技术维度】技术路线如何？关键技术点有哪些？
【成本维度】成本估算多少？需要哪些资源？
【进度维度】工期预估多久？里程碑节点是什么？
【风险维度】识别了哪些风险？应对策略是什么？

要求：
- 每个维度必须有具体内容，基于讨论摘要生成
- 引用讨论中的关键观点和数据
- 如果有质疑，说明质疑处理情况
- 语气专业有力，使用🌿开头"""
        
        elif msg_type == 'consensus':
            consensus = context.get('consensus', 0)
            output = context.get('output', '')
            discussion_summary = context.get('discussion_summary', '')
            
            user_message = f"""讨论完成，共识度{consensus}%。

决策方案产出：{output}

讨论摘要：
{discussion_summary}

请基于以上讨论内容，生成具体的多维度总结：

【需求维度】根据讨论，明确哪些核心功能？验收标准是什么？
【技术维度】技术路线如何？关键技术点有哪些？
【成本维度】成本估算多少？需要哪些资源？
【进度维度】工期预估多久？里程碑节点是什么？
【风险维度】识别了哪些风险？应对策略是什么？

要求：
- 每个维度必须有具体内容，不能泛泛而谈
- 引用讨论中的关键观点和数据
- 语气专业有力，使用🌿开头"""
        
        else:
            return "🌿 系统处理中..."
        
        response = call_qianfan(nanqiao_prompt, user_message, temperature=0.8)
        
        if response:
            # 确保有🌿开头
            if not response.startswith('🌿'):
                response = f'🌿 {response}'
            return response.strip()
        
        # 备用固定话术
        return self._fallback_system_message(msg_type, context)
    
    def _fallback_system_message(self, msg_type: str, context: dict) -> str:
        """备用固定话术（API失败时使用）"""
        if msg_type == 'schedule':
            if context.get('multi'):
                return f"🌿 已识别{len(context.get('tasks', []))}个任务，正在调度团队..."
            return f"🌿 任务已识别，正在调度{context.get('lead_agent', '采薇')}主导分析..."
        elif msg_type == 'task_start':
            return f"🌿 开始任务：{context.get('task_name', '')}，预计{context.get('est_time', '3-5天')}..."
        elif msg_type == 'risk_debate':
            return "🌿 进入风险辩论，请各位专家提出质疑..."
        elif msg_type == 'task_done':
            return f"🌿 任务完成，共识度{context.get('consensus', 0)}%..."
        elif msg_type == 'all_done':
            total = context.get('total', 1)
            consensus = context.get('consensus', 0)
            outputs = context.get('outputs', [])
            pending = context.get('pending_challenges', 0)
            
            base_msg = f"""🌿 多维度论证完成！共识度{consensus}%

【需求维度】已明确核心功能和验收标准
【技术维度】已确定技术路线和关键方案
【成本维度】已完成成本估算和资源规划
【进度维度】已制定工期计划和里程碑
【风险维度】已识别风险并制定应对策略

产出{len(outputs)}份结构化方案，降低返工风险。"""
            
            if pending > 0:
                base_msg += f"\n\n⏳ 待回应质疑：{pending}条，请关注。"
            
            return base_msg
        elif msg_type == 'consensus':
            return f"""🌿 多维度决策完成！共识度{context.get('consensus', 0)}%

【需求维度】核心功能已明确
【技术维度】技术路线已确定
【成本维度】成本估算已完成
【进度维度】工期预估已确定
【风险维度】风险已识别

产出{context.get('output', '')}，降低返工风险。"""
        return "🌿 处理中..."

    def generate_estimate_message(self, agent_name: str, task_name: str, est_time: str, complexity: str) -> str:
        """
        动态生成工期评估话术
        
        Args:
            agent_name: Agent名称
            task_name: 任务名称
            est_time: 预估时间
            complexity: 复杂度
        
        Returns:
            动态生成的工期评估话术
        """
        prompt = f"""你是{agent_name}，正在评估一个{task_name}任务的工期。

任务复杂度：{complexity}
预估工期：{est_time}

请生成一句话，说明工期评估结果。

规则：
1. 使用⏱️开头
2. 简洁专业，控制在30字以内
3. 提及复杂度因素（如有）
4. 保持友好、专业的语气
5. 不要使用【】符号"""

        response = call_qianfan(prompt, f"请为{task_name}任务（{complexity}复杂度，{est_time}工期）生成评估话术。", temperature=0.8)
        
        if response:
            response = response.strip()
            # 确保有⏱️开头
            if not response.startswith('⏱️'):
                response = f'⏱️ {response}'
            return response
        
        # 备用话术（根据复杂度）
        if complexity == 'high':
            return f"⏱️ 任务复杂度高，需{est_time}，建议预留缓冲时间。"
        elif complexity == 'medium':
            return f"⏱️ 任务规模适中，预计{est_time}完成。"
        else:
            return f"⏱️ 任务明确，{est_time}可交付，可协商加速。"

    def generate(self, agent: AgentPersona, task: str, memory: ConversationMemory) -> Tuple[str, bool, str]:
        context = memory.get_context()
        system_prompt = agent.get_system_prompt()

        # 根据角色立场引导讨论维度
        discussion_prompts = {
            '需求分析专家': '请从需求拆解、用户故事、验收标准、需求完整性角度分析',
            '项目经理': '请从进度计划、风险管控、资源协调、里程碑节点角度分析',
            '架构设计师': '请从技术选型、架构方案、关键技术点、扩展性角度分析',
            '售前工程师': '请从成本估算、资源需求、预算控制角度分析',
            '详细设计师': '请从接口设计、模块划分、数据库设计角度分析',
            '方案设计师': '请从UI设计、用户体验、视觉呈现角度分析',
            '资源管家': '请从人员配置、技术资源、知识库支持角度分析',
            '总指挥': '请从整体平衡、决策建议、优先级排序角度分析'
        }
        
        role_prompt = discussion_prompts.get(agent.role, '请发表专业意见')

        if context:
            user_message = f"""当前任务：{task}

已有讨论：
{context}

{role_prompt}。如果有不同意见请明确提出质疑。"""
        else:
            user_message = f"""当前任务：{task}

{role_prompt}。

讨论要求：
1. 不要只讨论工期，要深入讨论专业内容
2. 提出具体的分析、建议或质疑
3. 如果发现问题，明确指出"""

        response = call_qianfan(system_prompt, user_message)

        if response:
            response = response.strip()
            # 判断内容类型
            is_challenge = any(kw in response for kw in ['质疑', '反对', '不同意', '不可行', '成本过高', '周期太长', '挑战'])
            is_decision = any(kw in response for kw in ['决策', '结论', '最终方案', '共识', '采纳', '确定'])

            reply_to = ""
            if is_challenge and memory.history:
                for t in reversed(memory.history):
                    if t.speaker != agent.agent_id:
                        reply_to = t.speaker_name
                        break

            # 根据内容类型添加emoji表情
            response = self._add_emoji(response, agent, is_challenge, is_decision)

            return response, is_challenge, reply_to

        return self._fallback(agent, task, context)

    def _add_emoji(self, text: str, agent: AgentPersona, is_challenge: bool, is_decision: bool) -> str:
        """根据内容类型添加emoji表情"""

        # 质疑/反驳类型
        if is_challenge:
            challenge_emojis = ['🤔', '💭', '❓', '⚠️', '🔍', '🧐', '🙅', '🚧']
            import random
            emoji = random.choice(challenge_emojis)
            return f"{emoji} {text}"

        # 决策类型
        if is_decision:
            decision_emojis = ['⚖️', '🎯', '📌', '✨', '🏆', '💯', '🚀']
            import random
            emoji = random.choice(decision_emojis)
            return f"{emoji} {text}"

        # 支持/普通类型 - 根据Agent角色添加特色emoji
        agent_emojis = {
            'caiwei': ['🌸', '📝', '📋', '✏️'],      # 需求分析
            'zhijin': ['🧵', '🏗️', '📐', '🔧'],      # 架构设计
            'zhutai': ['🏗️', '💰', '📊', '📈'],      # 成本分析
            'chengcai': ['🎨', '🖼️', '✨', '💎'],    # 方案设计
            'gongchi': ['📐', '📏', '📊', '🗂️'],     # 详细设计
            'yuheng': ['⚖️', '📅', '⏰', '📋'],      # 项目管理
            'fuyao': ['🌀', '🎯', '⚡', '🔮'],        # 总指挥
            'nanqiao': ['🌿', '💬', '✨', '🌱'],      # 主控
        }

        import random
        emojis = agent_emojis.get(agent.agent_id, ['💬', '✅', '💪'])
        emoji = random.choice(emojis)
        return f"{emoji} {text}"

    def _fallback(self, agent: AgentPersona, task: str, context: str) -> Tuple[str, bool, str]:
        responses = {
            'caiwei': ("🌸 【需求分析】我梳理了核心需求。主要功能点已识别，建议进行技术可行性评估。", False, ""),
            'zhijin': ("🧵 【架构设计】基于需求，推荐微服务架构。技术栈：Spring Cloud + K8s + PostgreSQL。", False, ""),
            'zhutai': ("⚠️ 【成本质疑】微服务架构成本较高，建议评估必要性。可采用单体架构降低初期投入。", True, "织锦"),
            'yuheng': ("⏰ 【进度提醒】复杂方案会增加开发周期。建议敏捷迭代，首期交付核心功能。", True, "织锦"),
            'gongchi': ("📐 【详细设计】数据库设计约30张表，核心接口50个。开发工时评估完成。", False, ""),
            'fuyao': ("🎯 【决策】综合讨论，采纳分阶段方案。平衡成本、质量和进度。", False, ""),
            'nanqiao': ("🌿 【主控】收到任务，协调团队分析中。", False, ""),
            'chengcai': ("🎨 【方案设计】建议采用分层架构展示，突出核心价值点。", False, ""),
            'zhegui': ("📚 【资源管理】已整理相关技术文档和案例库，可随时调用。", False, "")
        }
        return responses.get(agent.agent_id, ("💬 处理中...", False, ""))


# ==================== 全局状态 ====================
memory = ConversationMemory()
responder = IntelligentResponder()
agent_status: Dict[str, str] = {}
is_processing = False
discussion_completed = False  # 讨论完成标志

# Token统计
total_tokens_used = 0  # 总消耗token数
TOKEN_BUDGET = 1000000  # Token预算（100万）

def calculate_dynamic_turn_limit(complexity: str, num_agents: int, num_tasks: int = 1) -> int:
    """
    动态计算讨论轮次上限
    
    Args:
        complexity: 任务复杂度 (low/medium/high)
        num_agents: 参与Agent数量
        num_tasks: 任务数量
    
    Returns:
        动态计算的轮次上限
    """
    # 基础轮次（根据复杂度）
    base_turns = {
        'low': 10,
        'medium': 15,
        'high': 20
    }
    
    base = base_turns.get(complexity, 12)
    
    # Agent数量调整（每多1个Agent增加1轮）
    agent_bonus = max(0, num_agents - 3)  # 3个Agent为基准
    
    # 多任务调整（每个任务增加3轮）
    task_bonus = max(0, (num_tasks - 1) * 3)
    
    # 计算总轮次
    total_turns = base + agent_bonus + task_bonus
    
    # 限制在合理范围内
    return min(max(total_turns, 8), 30)  # 最少8轮，最多30轮


# ==================== HTML模板 ====================
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>⭐ 九星智囊团 - 智能协作平台 V14 + V15深度融合版</title>
    <style>
        :root {
            --primary: #165DFF;
            --primary-light: #E8F3FF;
            --bg-main: #F5F7FA;
            --bg-card: #FFFFFF;
            --text-primary: #303133;
            --text-secondary: #606266;
            --text-muted: #909399;
            --border: #E4E7ED;
            --success: #67C23A;
            --warning: #E6A23C;
            --danger: #F56C6C;
        }

        * { margin: 0; padding: 0; box-sizing: border-box; }

        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif;
            background: var(--bg-main);
            color: var(--text-primary);
            height: 100vh;
            display: flex;
            flex-direction: column;
            font-size: 14px;
        }

        /* ========== 顶部导航 ========== */
        .header {
            background: var(--bg-card);
            border-bottom: 1px solid var(--border);
            padding: 0 24px;
            height: 56px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            box-shadow: 0 1px 4px rgba(0,0,0,0.04);
        }
        .header-left {
            display: flex;
            align-items: center;
            gap: 12px;
        }
        .logo {
            font-size: 20px;
            font-weight: 600;
            color: var(--primary);
        }
        .logo-sub {
            font-size: 12px;
            color: var(--text-muted);
            padding: 2px 8px;
            background: var(--primary-light);
            border-radius: 4px;
        }
        .header-right {
            display: flex;
            align-items: center;
            gap: 16px;
        }
        .status-badge {
            display: flex;
            align-items: center;
            gap: 6px;
            font-size: 12px;
            color: var(--success);
        }
        .status-dot {
            width: 8px;
            height: 8px;
            background: var(--success);
            border-radius: 50%;
        }

        /* ========== 主区域 ========== */
        .main {
            flex: 1;
            display: flex;
            overflow: hidden;
        }

        /* ========== 左侧边栏 ========== */
        .sidebar {
            width: 300px;
            background: var(--bg-card);
            border-right: 1px solid var(--border);
            display: flex;
            flex-direction: column;
        }
        .sidebar-header {
            padding: 16px;
            border-bottom: 1px solid var(--border);
            font-size: 13px;
            font-weight: 600;
            color: var(--text-secondary);
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .sidebar-content {
            flex: 1;
            overflow-y: auto;
            padding: 12px;
            min-height: 200px;
            background: var(--bg-card);
        }

        /* Agent卡片 */
        .agent-card {
            display: flex;
            align-items: center;
            padding: 10px 12px;
            margin-bottom: 6px;
            background: var(--bg-main);
            border-radius: 8px;
            border-left: 3px solid transparent;
            transition: all 0.2s ease;
            cursor: pointer;
        }
        .agent-card:hover {
            background: var(--primary-light);
        }
        .agent-card.active {
            border-left-color: var(--agent-color, var(--primary));
            background: var(--primary-light);
        }
        .agent-card.speaking {
            border-left-color: var(--agent-color, var(--primary));
            box-shadow: 0 2px 8px rgba(22, 93, 255, 0.15);
        }
        .agent-card.challenge {
            border-left-color: var(--warning);
            background: #FFF7E6;
        }

        .agent-avatar {
            width: 36px;
            height: 36px;
            border-radius: 8px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 16px;
            margin-right: 10px;
            color: white;
        }
        .agent-info {
            flex: 1;
        }
        .agent-name {
            font-size: 13px;
            font-weight: 600;
            color: var(--text-primary);
        }
        .agent-role {
            font-size: 11px;
            color: var(--text-muted);
            margin-top: 1px;
        }
        .agent-stance {
            display: inline-block;
            font-size: 9px;
            padding: 1px 5px;
            border-radius: 3px;
            margin-top: 3px;
            background: var(--primary-light);
            color: var(--primary);
        }

        /* 情绪状态 */
        .agent-mood {
            font-size: 10px;
            margin-top: 4px;
            display: flex;
            align-items: center;
            gap: 4px;
        }
        .mood-icon {
            font-size: 12px;
        }
        .mood-text {
            color: var(--text-muted);
        }
        .mood-thinking { color: #409EFF; }
        .mood-worried { color: #E6A23C; }
        .mood-excited { color: #67C23A; }
        .mood-satisfied { color: #909399; }
        .mood-challenging { color: #F56C6C; }

        /* ========== 对话区域 ========== */
        .chat-area {
            flex: 1;
            display: flex;
            flex-direction: column;
            min-width: 0;
        }

        /* 筛选栏 */
        .filter-bar {
            padding: 10px 16px;
            background: var(--bg-card);
            border-bottom: 1px solid var(--border);
            display: flex;
            gap: 8px;
            align-items: center;
        }
        .filter-btn {
            padding: 5px 12px;
            border: 1px solid var(--border);
            border-radius: 6px;
            background: var(--bg-card);
            font-size: 12px;
            color: var(--text-secondary);
            cursor: pointer;
            transition: all 0.2s;
        }
        .filter-btn:hover, .filter-btn.active {
            border-color: var(--primary);
            color: var(--primary);
            background: var(--primary-light);
        }

        /* 消息区 */
        .chat-messages {
            flex: 1;
            overflow-y: auto;
            padding: 16px;
            background: var(--bg-main);
        }

        /* 消息气泡 */
        .message {
            margin-bottom: 16px;
            position: relative;
            padding-left: 48px;
        }

        .message.user {
            text-align: right;
            padding-left: 0;
            padding-right: 0;
        }

        .message-avatar {
            position: absolute;
            left: 0;
            top: 0;
            width: 36px;
            height: 36px;
            border-radius: 8px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 16px;
            color: white;
            z-index: 1;
        }

        .message-content {
            display: inline-block;
            max-width: 75%;
            background: var(--bg-card);
            border-radius: 10px;
            padding: 12px 14px;
            box-shadow: 0 2px 6px rgba(0,0,0,0.04);
            position: relative;
            border-left: 3px solid transparent;
            transition: all 0.3s ease;
        }

        /* 气泡样式差异化 */
        .message[data-type="challenge"] .message-content {
            border-left: 3px solid #F56C6C;
            background: linear-gradient(135deg, #FFF5F5 0%, #FFFFFF 100%);
        }
        .message[data-type="support"] .message-content {
            border-left: 3px solid #67C23A;
            background: linear-gradient(135deg, #F0F9EB 0%, #FFFFFF 100%);
        }
        .message[data-type="decision"] .message-content {
            border-left: 3px solid #E6A23C;
            background: linear-gradient(135deg, #FDF6EC 0%, #FFFFFF 100%);
            box-shadow: 0 4px 12px rgba(230, 162, 60, 0.2);
        }

        .message.user .message-content {
            background: var(--primary);
            color: white;
        }

        .message-header {
            display: flex;
            align-items: center;
            gap: 6px;
            margin-bottom: 6px;
        }
        .message-speaker {
            font-size: 12px;
            font-weight: 600;
        }
        .message-time {
            font-size: 10px;
            color: var(--text-muted);
        }
        .message-badge {
            font-size: 9px;
            padding: 2px 6px;
            border-radius: 4px;
            background: #FFF7E6;
            color: var(--warning);
            font-weight: 500;
        }

        /* 正在输入提示 */
        .typing-indicator {
            display: flex;
            align-items: center;
            gap: 8px;
            padding: 10px 16px;
            background: var(--primary-light);
            border-radius: 20px;
            margin: 8px 0;
            animation: fadeIn 0.3s ease;
        }
        .typing-indicator .dots {
            display: flex;
            gap: 4px;
        }
        .typing-indicator .dot {
            width: 6px;
            height: 6px;
            background: var(--primary);
            border-radius: 50%;
            animation: bounce 1.4s infinite ease-in-out;
        }
        .typing-indicator .dot:nth-child(1) { animation-delay: 0s; }
        .typing-indicator .dot:nth-child(2) { animation-delay: 0.2s; }
        .typing-indicator .dot:nth-child(3) { animation-delay: 0.4s; }

        @keyframes bounce {
            0%, 80%, 100% { transform: translateY(0); }
            40% { transform: translateY(-6px); }
        }

        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }

        /* 头像脉冲动画 */
        .message-avatar.speaking {
            animation: pulse 1.5s infinite;
        }
        @keyframes pulse {
            0% { box-shadow: 0 0 0 0 rgba(22, 93, 255, 0.4); }
            70% { box-shadow: 0 0 0 10px rgba(22, 93, 255, 0); }
            100% { box-shadow: 0 0 0 0 rgba(22, 93, 255, 0); }
        }

        /* 消息操作按钮 */
        .message-actions {
            display: flex;
            gap: 8px;
            margin-top: 8px;
            opacity: 0.6;
            transition: opacity 0.2s;
        }
        .message:hover .message-actions {
            opacity: 1;
        }
        .action-btn {
            font-size: 12px;
            padding: 4px 10px;
            border-radius: 4px;
            cursor: pointer;
            background: rgba(255,255,255,0.9);
            color: #606266;
            border: 1px solid #DCDFE6;
            transition: all 0.2s;
        }
        .action-btn:hover {
            background: #409EFF;
            color: white;
            border-color: #409EFF;
        }
        .action-btn.liked {
            background: #67C23A;
            color: white;
            border-color: #67C23A;
        }
        .action-btn.disliked {
            background: #F56C6C;
            color: white;
            border-color: #F56C6C;
        }

        .message-text {
            font-size: 13px;
            line-height: 1.6;
            white-space: pre-wrap;
        }
        .typing-cursor {
            color: var(--primary);
            animation: blink 0.8s infinite;
            font-weight: bold;
        }
        @keyframes blink {
            0%, 50% { opacity: 1; }
            51%, 100% { opacity: 0; }
        }
        .message-footer {
            display: flex;
            gap: 10px;
            margin-top: 6px;
        }
        .message-action {
            font-size: 10px;
            color: var(--text-muted);
            cursor: pointer;
        }
        .message-action:hover {
            color: var(--primary);
        }

        /* 回复箭头 */
        .reply-arrow {
            font-size: 11px;
            color: var(--text-muted);
            margin: 6px 0;
            padding-left: 20px;
        }

        /* 输入区 */
        .input-area {
            padding: 14px 16px;
            background: var(--bg-card);
            border-top: 1px solid var(--border);
        }
        .input-row {
            display: flex;
            gap: 10px;
        }
        .input-field {
            flex: 1;
            padding: 10px 14px;
            border: 1px solid var(--border);
            border-radius: 8px;
            font-size: 14px;
            outline: none;
            transition: border-color 0.2s;
        }
        .input-field:focus {
            border-color: var(--primary);
        }
        .send-btn {
            padding: 10px 24px;
            background: var(--primary);
            color: white;
            border: none;
            border-radius: 8px;
            font-size: 14px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.2s;
        }
        .send-btn:hover:not(:disabled) {
            background: #0D47A1;
        }
        .send-btn:disabled {
            opacity: 0.5;
            cursor: not-allowed;
        }

        .quick-actions {
            display: flex;
            gap: 6px;
            margin-top: 10px;
            flex-wrap: wrap;
        }
        .quick-btn {
            padding: 5px 12px;
            border: 1px solid var(--border);
            border-radius: 6px;
            background: var(--bg-card);
            font-size: 11px;
            color: var(--text-secondary);
            cursor: pointer;
            transition: all 0.2s;
        }
        .quick-btn:hover {
            border-color: var(--primary);
            color: var(--primary);
        }

        /* ========== 右侧数据面板 ========== */
        .data-panel {
            width: 280px;
            background: var(--bg-card);
            border-left: 1px solid var(--border);
            padding: 16px;
            overflow-y: auto;
        }
        .panel-title {
            font-size: 13px;
            font-weight: 600;
            color: var(--text-primary);
            margin-bottom: 12px;
        }

        /* 共识度 */
        .consensus-card {
            background: var(--bg-main);
            border-radius: 10px;
            padding: 14px;
            margin-bottom: 16px;
        }
        .consensus-label {
            font-size: 11px;
            color: var(--text-secondary);
            margin-bottom: 6px;
        }
        .consensus-bar {
            height: 6px;
            background: var(--border);
            border-radius: 3px;
            overflow: hidden;
        }
        .consensus-fill {
            height: 100%;
            border-radius: 3px;
            transition: width 0.5s ease, background 0.5s ease;
        }
        .consensus-value {
            font-size: 22px;
            font-weight: 700;
            color: var(--text-primary);
            margin-top: 6px;
        }

        /* 统计卡片 */
        .stats-card {
            background: var(--bg-main);
            border-radius: 10px;
            padding: 12px;
            margin-bottom: 12px;
        }
        .stat-item {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 6px 0;
            border-bottom: 1px solid var(--border);
        }
        .stat-item:last-child {
            border-bottom: none;
        }
        .stat-label {
            font-size: 11px;
            color: var(--text-secondary);
        }
        .stat-value {
            font-size: 13px;
            font-weight: 600;
            color: var(--text-primary);
        }

        /* 发言统计图 */
        .speaker-stats {
            margin-top: 12px;
        }
        .speaker-bar {
            display: flex;
            align-items: center;
            margin-bottom: 6px;
        }
        .speaker-name {
            width: 50px;
            font-size: 10px;
            color: var(--text-secondary);
        }
        .speaker-progress {
            flex: 1;
            height: 5px;
            background: var(--border);
            border-radius: 3px;
            overflow: hidden;
            margin: 0 6px;
        }
        .speaker-fill {
            height: 100%;
            border-radius: 3px;
        }
        .speaker-count {
            width: 24px;
            font-size: 10px;
            color: var(--text-muted);
            text-align: right;
        }

        /* ========== 空状态 ========== */
        .empty-state {
            text-align: center;
            padding: 60px 20px;
            color: var(--text-muted);
        }
        .empty-icon {
            font-size: 56px;
            margin-bottom: 12px;
            opacity: 0.5;
        }
        .empty-title {
            font-size: 15px;
            color: var(--text-secondary);
            margin-bottom: 6px;
        }
        .empty-desc {
            font-size: 12px;
        }

        /* ========== 加载状态 ========== */
        .agent-loading {
            display: flex;
            align-items: center;
            gap: 12px;
            padding: 12px 16px;
            background: var(--bg-main);
            border-radius: 12px;
            margin-bottom: 16px;
        }
        .agent-loading-avatar {
            width: 36px;
            height: 36px;
            border-radius: 50%;
            background: var(--primary);
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 18px;
            animation: pulse 2s infinite;
        }
        @keyframes pulse {
            0%, 100% { opacity: 1; transform: scale(1); }
            50% { opacity: 0.8; transform: scale(1.05); }
        }
        .agent-loading-text {
            flex: 1;
            color: var(--text-secondary);
        }
        .agent-loading-name {
            font-weight: 600;
            color: var(--text-primary);
        }
        .typing-dots {
            display: inline-flex;
            gap: 3px;
        }
        .typing-dots span {
            width: 6px;
            height: 6px;
            border-radius: 50%;
            background: var(--primary);
            animation: typing 1.4s infinite;
        }
        .typing-dots span:nth-child(2) { animation-delay: 0.2s; }
        .typing-dots span:nth-child(3) { animation-delay: 0.4s; }
        @keyframes typing {
            0%, 100% { opacity: 0.3; }
            50% { opacity: 1; }
        }

        /* ========== 消息交互 ========== */
        .message-actions {
            display: flex;
            gap: 8px;
            margin-top: 8px;
            opacity: 0;
            transition: opacity 0.2s;
        }
        .message-wrapper:hover .message-actions {
            opacity: 1;
        }
        .msg-action-btn {
            background: none;
            border: none;
            color: var(--text-muted);
            cursor: pointer;
            padding: 4px 8px;
            border-radius: 4px;
            font-size: 11px;
            transition: all 0.2s;
        }
        .msg-action-btn:hover {
            background: var(--bg-main);
            color: var(--text-primary);
        }
        .msg-action-btn.liked {
            color: #f43f5e;
        }

        /* ========== 响应式 ========== */
        @media (max-width: 1024px) {
            .data-panel { display: none; }
        }
        @media (max-width: 768px) {
            .sidebar {
                position: absolute;
                left: -300px;
                z-index: 100;
                height: 100%;
            }
            .sidebar.open { left: 0; }
        }
    </style>
</head>
<body>
    <div class="header">
        <div class="header-left">
            <div class="logo">🧭 指南针工程</div>
            <div class="logo-sub">智能协作平台 V14+V15.2产物生成 | 九星智囊团 <span style="color:#ff0">[V15.2:2026-03-19 产物生成集成]</span></div>
            <div style="font-size:11px;color:#909399;margin-top:2px;">九星汇聚，智胜千里 | 以智为针，以信为盘</div>
        </div>
        <div class="header-right">
            <div class="status-badge">
                <div class="status-dot"></div>
                <span id="apiStatus">千帆API已连接</span>
            </div>
            <button class="filter-btn" onclick="window.location.href='/dispatch'" style="background:#C93832; color:white;">📋 调度看板</button>
            <button class="filter-btn" id="generateProductBtn" style="display:none; background:#67C23A; color:white; font-weight:bold;" onclick="generateProduct()">🚀 生成产物</button>
            <button class="filter-btn" onclick="exportWord()">📥 Word</button>
            <button class="filter-btn" onclick="exportExcel()">📊 Excel</button>
            <button class="filter-btn" onclick="exportMarkdown()">📝 MD</button>
            <button class="filter-btn" onclick="previewMinutes()">📋 会议纪要</button>
            <button class="filter-btn" onclick="downloadMinutes()">⬇️ 下载纪要</button>
        </div>
    </div>

    <div class="main">
        <!-- 左侧Agent面板 -->
        <div class="sidebar">
            <div class="sidebar-header">
                <span>⭐ 九星智囊团</span>
            </div>
            <div class="sidebar-content" id="agentList">
                <!-- 静态后备内容 - 如果JavaScript渲染失败则显示 -->
                <div class="agent-card" style="border-left-color: #165DFF;"><div class="agent-avatar" style="background: #165DFF">🌀</div><div class="agent-info"><div class="agent-name">扶摇</div><div class="agent-role">总指挥</div></div></div>
                <div class="agent-card" style="border-left-color: #9C27B0;"><div class="agent-avatar" style="background: #9C27B0">🌿</div><div class="agent-info"><div class="agent-name">南乔</div><div class="agent-role">主控Agent</div></div></div>
                <div class="agent-card" style="border-left-color: #F56C6C;"><div class="agent-avatar" style="background: #F56C6C">⚖️</div><div class="agent-info"><div class="agent-name">玉衡</div><div class="agent-role">项目经理</div></div></div>
                <div class="agent-card" style="border-left-color: #409EFF;"><div class="agent-avatar" style="background: #409EFF">🌸</div><div class="agent-info"><div class="agent-name">采薇</div><div class="agent-role">需求分析专家</div></div></div>
                <div class="agent-card" style="border-left-color: #67C23A;"><div class="agent-avatar" style="background: #67C23A">🧵</div><div class="agent-info"><div class="agent-name">织锦</div><div class="agent-role">架构设计师</div></div></div>
                <div class="agent-card" style="border-left-color: #E6A23C;"><div class="agent-avatar" style="background: #E6A23C">🏗️</div><div class="agent-info"><div class="agent-name">筑台</div><div class="agent-role">售前工程师</div></div></div>
                <div class="agent-card" style="border-left-color: #607D8B;"><div class="agent-avatar" style="background: #607D8B">📐</div><div class="agent-info"><div class="agent-name">工尺</div><div class="agent-role">详细设计师</div></div></div>
                <div class="agent-card" style="border-left-color: #FF9800;"><div class="agent-avatar" style="background: #FF9800">🎨</div><div class="agent-info"><div class="agent-name">呈彩</div><div class="agent-role">方案设计师</div></div></div>
                <div class="agent-card" style="border-left-color: #00BCD4;"><div class="agent-avatar" style="background: #00BCD4">📚</div><div class="agent-info"><div class="agent-name">折桂</div><div class="agent-role">资源管家</div></div></div>
            </div>
        </div>

        <!-- 中间对话区 -->
        <div class="chat-area">
            <div class="filter-bar">
                <button class="filter-btn active" onclick="filterMessages('all')">全部</button>
                <button class="filter-btn" onclick="filterMessages('challenge')">质疑</button>
                <button class="filter-btn" onclick="filterMessages('support')">支持</button>
                <button class="filter-btn" onclick="filterMessages('decision')">决策</button>
            </div>

            <div class="chat-messages" id="chatMessages">
                <div class="empty-state">
                    <div class="empty-icon">💬</div>
                    <div class="empty-title">开始智能协作</div>
                    <div class="empty-desc">输入任务主题，Agent团队将进行专业讨论和辩论</div>
                </div>
            </div>

            <div class="input-area">
                <div class="input-row">
                    <input type="text" class="input-field" id="taskInput"
                           placeholder="输入任务主题，按 Enter 发送..."
                           onkeypress="handleKeyPress(event)">
                    <button class="send-btn" id="sendBtn" onclick="submitTask()">发送</button>
                </div>
                <div class="quick-actions">
                    <button class="quick-btn" onclick="showScenePanel()">🎭 场景模板</button>
                    <button class="quick-btn" onclick="showCostPanel()">💰 成本估算</button>
                    <label class="quick-btn" style="cursor: pointer;">
                        📎 上传文档
                        <input type="file" id="docUpload" accept=".pdf,.doc,.docx,.txt" style="display:none" onchange="uploadDocument(this)">
                    </label>
                    <button class="quick-btn" onclick="clearChat()">🗑️ 清空</button>
                </div>
            </div>
        </div>

        <!-- 右侧数据面板 -->
        <div class="data-panel">
            <div class="panel-title">📊 协作数据</div>

            <div class="consensus-card">
                <div class="consensus-label">共识达成度</div>
                <div class="consensus-bar">
                    <div class="consensus-fill" id="consensusFill" style="width: 0%; background: #F56C6C;"></div>
                </div>
                <div class="consensus-value" id="consensusValue">0%</div>
            </div>

            <div class="stats-card">
                <div class="stat-item">
                    <span class="stat-label">当前轮次</span>
                    <span class="stat-value" id="roundCount">0 / 15</span>
                </div>
                <div class="stat-item">
                    <span class="stat-label">辩论耗时</span>
                    <span class="stat-value" id="elapsedTime">00:00</span>
                </div>
                <div class="stat-item">
                    <span class="stat-label">质疑次数</span>
                    <span class="stat-value" id="challengeCount">0</span>
                </div>
            </div>

            <!-- 质疑增强面板 -->
            <div class="panel-title" style="margin-top: 16px;">🎯 质疑中心</div>
            <div class="challenge-panel" id="challengePanel" style="background: #f8f9fa; border-radius: 8px; padding: 12px;">
                <div class="challenge-stats" style="display: flex; gap: 8px; margin-bottom: 10px;">
                    <div style="flex: 1; text-align: center; padding: 6px; background: white; border-radius: 4px;">
                        <div style="font-size: 12px; color: #909399;">待回应</div>
                        <div id="pendingCount" style="font-size: 18px; font-weight: bold; color: #E6A23C;">0</div>
                    </div>
                    <div style="flex: 1; text-align: center; padding: 6px; background: white; border-radius: 4px;">
                        <div style="font-size: 12px; color: #909399;">已接受</div>
                        <div id="acceptedCount" style="font-size: 18px; font-weight: bold; color: #67C23A;">0</div>
                    </div>
                    <div style="flex: 1; text-align: center; padding: 6px; background: white; border-radius: 4px;">
                        <div style="font-size: 12px; color: #909399;">已拒绝</div>
                        <div id="rejectedCount" style="font-size: 18px; font-weight: bold; color: #F56C6C;">0</div>
                    </div>
                </div>
                <div id="challengeList" style="max-height: 150px; overflow-y: auto;">
                    <div style="color: #909399; font-size: 12px; text-align: center; padding: 10px;">
                        暂无质疑记录
                    </div>
                </div>
            </div>

            <div class="panel-title" style="margin-top: 16px;">🎤 发言统计</div>
            <div class="speaker-stats" id="speakerStats"></div>
        </div>
    </div>

    <script>
        // ========== 立即测试 ==========
        console.log('=== JavaScript开始执行 ===');
        console.log('页面加载时间:', new Date().toLocaleString());
        
        // 完整9个Agent配置
        const AGENTS = {
            'fuyao': { name: '扶摇', role: '总指挥', emoji: '🌀', stance: '平衡', color: '#165DFF' },
            'nanqiao': { name: '南乔', role: '主控Agent', emoji: '🌿', stance: '平衡', color: '#9C27B0' },
            'yuheng': { name: '玉衡', role: '项目经理', emoji: '⚖️', stance: '进度优先', color: '#F56C6C' },
            'caiwei': { name: '采薇', role: '需求分析专家', emoji: '🌸', stance: '质量优先', color: '#409EFF' },
            'zhijin': { name: '织锦', role: '架构设计师', emoji: '🧵', stance: '质量优先', color: '#67C23A' },
            'zhutai': { name: '筑台', role: '售前工程师', emoji: '🏗️', stance: '成本优先', color: '#E6A23C' },
            'gongchi': { name: '工尺', role: '详细设计师', emoji: '📐', stance: '质量优先', color: '#607D8B' },
            'chengcai': { name: '呈彩', role: '方案设计师', emoji: '🎨', stance: '质量优先', color: '#FF9800' },
            'zhegui': { name: '折桂', role: '资源管家', emoji: '📚', stance: '平衡', color: '#00BCD4' }
        };
        console.log('AGENTS定义完成:', Object.keys(AGENTS).length, '个');
        
        // ========== 立即渲染Agent列表 ==========
        document.addEventListener('DOMContentLoaded', function() {
            console.log('=== DOMContentLoaded触发 ===');
            renderAgentListNow();
        });
        
        // 如果DOMContentLoaded已经触发，立即执行
        if (document.readyState === 'complete' || document.readyState === 'interactive') {
            console.log('=== 文档已加载，立即渲染 ===');
            setTimeout(function() {
                renderAgentListNow();
            }, 100);
        }
        
        function renderAgentListNow() {
            const container = document.getElementById('agentList');
            console.log('agentList容器:', container);
            if (!container) {
                console.error('找不到agentList容器！');
                return;
            }
            
            container.innerHTML = '';
            const order = ['fuyao', 'nanqiao', 'yuheng', 'caiwei', 'zhijin', 'zhutai', 'gongchi', 'chengcai', 'zhegui'];
            
            order.forEach(id => {
                const agent = AGENTS[id];
                if (!agent) return;
                
                const card = document.createElement('div');
                card.className = 'agent-card';
                card.style.borderLeftColor = agent.color;
                card.innerHTML = `
                    <div class="agent-avatar" style="background: ${agent.color}">${agent.emoji}</div>
                    <div class="agent-info">
                        <div class="agent-name">${agent.name}</div>
                        <div class="agent-role">${agent.role}</div>
                        <div class="agent-stance">${agent.stance}</div>
                    </div>
                `;
                container.appendChild(card);
            });
            
            console.log('渲染完成:', order.length, '个Agent');
            document.getElementById('agentCount').textContent = order.length + '位专家';
        }

        // 常量定义
        const TOKEN_BUDGET = 1000000;  // Token预算
        const DISCUSSION_TURN_LIMIT = 15;  // 轮次限制（动态计算，前端仅作展示用）

        let conversations = [];
        let agentStatus = {};
        let agentMoods = {};  // 情绪状态
        let isProcessing = false;
        let startTime = null;
        let timerInterval = null;
        let currentFilter = 'all';
        let messageLikes = {};  // 点赞状态

        // 彩蛋关键词
        const EASTER_EGGS = {
            '加班': '😭 为什么要加班！我们强烈建议工作生活平衡！',
            '996': '😤 996是不合理的！我们坚决反对！',
            '甲方': '🤔 甲方需求永远在变，这是我们的宿命...',
            'bug': '🐛 Bug是程序员的宿命，但我们会尽力避免！',
            '加班吗': '🙅 今天不加班！效率第一！',
        };

        // 初始化
        function init() {
            console.log('=== init() 开始执行 ===');
            console.log('AGENTS对象:', AGENTS);
            console.log('agentList容器:', document.getElementById('agentList'));
            
            // 首次进入清空历史对话（不加载localStorage）
            localStorage.removeItem('chat_history');
            conversations = [];
            lastRenderedIndex = -1;
            consensus = 70;
            
            // 初始化渲染Agent列表
            renderAgentList();
            
            // 不再自动加载历史会话
            // if (loadFromHistory()) {
            //     renderMessages(true);
            //     renderAgentList();
            // }

            fetchStatus();
            setInterval(fetchStatus, 800);

            // 定期保存会话
            setInterval(saveToHistory, 30000);  // 每30秒保存一次
            
            console.log('=== init() 执行完成 ===');
        }

        // 渲染Agent列表（完整9个）- 带情绪状态
        function renderAgentList() {
            console.log('renderAgentList called'); // 调试日志
            const container = document.getElementById('agentList');
            if (!container) {
                console.error('agentList container not found!');
                return;
            }
            
            container.innerHTML = '';

            const order = ['fuyao', 'nanqiao', 'yuheng', 'caiwei', 'zhijin', 'zhutai', 'gongchi', 'chengcai', 'zhegui'];

            order.forEach(id => {
                const agent = AGENTS[id];
                if (!agent) return;

                const status = agentStatus[id] || 'idle';
                const mood = agentMoods[id] || { icon: '😊', text: '待命', class: 'mood-satisfied' };

                const card = document.createElement('div');
                card.className = `agent-card ${status === 'speaking' ? 'speaking' : ''} ${status === 'challenge' ? 'challenge' : ''}`;
                card.style.setProperty('--agent-color', agent.color);
                card.id = `agent-${id}`;
                card.onclick = () => filterByAgent(id);

                card.innerHTML = `
                    <div class="agent-avatar ${status === 'speaking' ? 'speaking' : ''}" style="background: ${agent.color}">${agent.emoji}</div>
                    <div class="agent-info">
                        <div class="agent-name">${agent.name}</div>
                        <div class="agent-role">${agent.role}</div>
                    </div>
                `;
                container.appendChild(card);
            });
        }

        // 获取状态
        async function fetchStatus() {
            try {
                const response = await fetch('/api/conversation');
                const data = await response.json();
                conversations = data.conversations || [];
                agentStatus = data.agentStatus || {};

                // 检测讨论完成
                if (data.isCompleted) {
                    console.log('[DEBUG] 检测到讨论完成，isProcessing:', isProcessing);
                    if (timerInterval) {
                        clearInterval(timerInterval);
                    }
                    isProcessing = false;
                    const sendBtn = document.getElementById('sendBtn');
                    if (sendBtn) {
                        sendBtn.disabled = false;
                    }
                    
                    // 显示"生成产物"按钮
                    const productBtn = document.getElementById('generateProductBtn');
                    if (productBtn) {
                        productBtn.style.display = 'inline-block';
                        console.log('[DEBUG] 已显示"生成产物"按钮');
                    }
                    
                    console.log('[DEBUG] ✅ 讨论完成，已解锁输入，isProcessing:', isProcessing);
                }

                // 如果正在打字，跳过渲染（避免打断打字效果）
                if (isTyping) {
                    console.log('[DEBUG] 正在打字，跳过渲染');
                    updateAgentCards();
                    updateDataPanel();
                    return;
                }

                // 如果正在处理中（讨论进行中），则使用打字效果
                // 如果讨论已完成，则直接显示所有消息
                renderMessages(!isProcessing);
                updateAgentCards();
                updateDataPanel();
            } catch (e) { console.error('[ERROR] fetchStatus error:', e); }
        }

        // 打字机效果 - 逐字显示（更明显的打字效果）
        let typingQueue = [];
        let isTyping = false;
        let typingSpeed = 30; // 每个字符的延迟（毫秒）

        function typeWriter(element, text, callback) {
            if (!text || text.length === 0) {
                console.log('[打字机] 文本为空，跳过');
                if (callback) callback();
                return;
            }
            
            isTyping = true;
            let i = 0;
            element.innerHTML = '';
            element.style.borderRight = '3px solid #165DFF'; // 添加光标效果
            element.style.paddingRight = '3px';
            console.log('[打字机] 开始打字，长度:', text.length);
            
            function type() {
                if (i < text.length) {
                    // 每次添加一个字符
                    element.innerHTML += text.charAt(i);
                    i++;
                    // 每添加一个字符滚动到底部
                    const container = document.getElementById('chatMessages');
                    container.scrollTop = container.scrollHeight;
                    setTimeout(type, typingSpeed);
                } else {
                    // 打字完成，移除光标
                    element.style.borderRight = 'none';
                    element.style.paddingRight = '0';
                    console.log('[打字机] 打字完成');
                    // 延迟一点再回调，让用户看到完成效果
                    setTimeout(() => {
                        isTyping = false;
                        if (callback) callback();
                    }, 200);
                }
            }
            type();
        }

        // 渲染消息（支持打字机效果 + 依次显示）
        let lastRenderedIndex = -1;

        function renderMessages(skipTyping = false) {
            const container = document.getElementById('chatMessages');

            if (conversations.length === 0) {
                container.innerHTML = `
                    <div class="empty-state">
                        <div class="empty-icon">💬</div>
                        <div class="empty-title">开始智能协作</div>
                        <div class="empty-desc">输入任务主题，Agent团队将进行专业讨论和辩论</div>
                    </div>
                `;
                lastRenderedIndex = -1;
                hideAgentLoading();
                return;
            }

            // 有新消息时隐藏加载状态
            if (conversations.length > 0) {
                hideAgentLoading();
            }

            // 只渲染新消息（一次一条，依次显示）
            const newMessages = conversations.slice(lastRenderedIndex + 1);

            if (newMessages.length === 0) return;

            // 如果使用打字效果，一次只渲染一条消息
            // 渲染完一条后，再渲染下一条
            if (!skipTyping && newMessages.length > 0 && !isTyping) {
                renderOneMessage(newMessages[0], lastRenderedIndex + 1, () => {
                    lastRenderedIndex++;
                    // 当前消息处理完后，继续渲染下一条
                    if (lastRenderedIndex < conversations.length - 1) {
                        setTimeout(() => renderMessages(skipTyping), 100);
                    }
                });
            } else if (skipTyping) {
                // 讨论完成，直接显示所有消息
                newMessages.forEach((conv, i) => {
                    renderOneMessage(conv, lastRenderedIndex + 1 + i, null, true);
                });
                lastRenderedIndex = conversations.length - 1;
            }

            container.scrollTop = container.scrollHeight;

            // 保存会话到历史
            if (conversations.length > 0) {
                saveToHistory();
            }
        }

        // 渲染单条消息
        function renderOneMessage(conv, idx, callback, immediate = false) {
            const container = document.getElementById('chatMessages');
            
            // 如果是第一条消息，清除空状态
            if (idx === 0 || container.querySelector('.empty-state')) {
                container.innerHTML = '';
            }
            
            const agent = AGENTS[conv.speaker] || {
                name: conv.speaker_name || conv.speaker,
                emoji: '🤖',
                color: '#165DFF',
                role: ''
            };
            
            const isUser = conv.speaker === 'user';
            
            // 时间格式化
            let time = '';
            if (conv.timestamp) {
                const date = new Date(conv.timestamp);
                time = date.toLocaleTimeString('zh-CN', { hour: '2-digit', minute: '2-digit' });
            }
            
            // 操作按钮（认可/存疑）
            let actionsHtml = '';
            if (!isUser && conv.msg_type !== 'system') {
                actionsHtml = `
                    <div class="message-actions">
                        <button class="action-btn approve" data-idx="${idx}" data-type="approve">👍 认可</button>
                        <button class="action-btn challenge" data-idx="${idx}" data-type="challenge">🤔 存疑</button>
                    </div>
                `;
            }

            // 用户消息直接显示，Agent消息根据immediate参数决定
            const textContent = (isUser || immediate) ? conv.content : '';

            const msgDiv = document.createElement('div');
            msgDiv.className = `message ${isUser ? 'user' : ''}`;
            msgDiv.style.opacity = '0';
            msgDiv.style.transform = 'translateY(20px)';
            msgDiv.style.transition = 'opacity 0.3s, transform 0.3s';

            msgDiv.innerHTML = `
                ${!isUser ? `<div class="message-avatar ${agentStatus[conv.speaker] === 'speaking' ? 'speaking' : ''}" style="background: ${agent.color}">${agent.emoji}</div>` : ''}
                <div class="message-content">
                    <div class="message-header">
                        <span class="message-speaker" style="color: ${isUser ? 'white' : agent.color}">${isUser ? '👤 少帅' : agent.emoji + ' ' + agent.name}${!isUser && agent.role ? ' · ' + agent.role : ''}</span>
                        ${conv.is_challenging ? '<span class="message-badge">⚠️ 质疑</span>' : ''}
                        ${conv.msg_type === 'conclusion' ? '<span class="message-badge" style="background:#E8F5E9;color:#67C23A">✓ 决策</span>' : ''}
                        <span class="message-time">${time}</span>
                    </div>
                    <div class="message-text" id="msg-text-${idx}">${textContent}</div>
                    ${actionsHtml}
                </div>
            `;

            container.appendChild(msgDiv);
            
            // 触发淡入动画
            requestAnimationFrame(() => {
                requestAnimationFrame(() => {
                    msgDiv.style.opacity = '1';
                    msgDiv.style.transform = 'translateY(0)';
                });
            });
            
            // Agent消息且需要打字效果
            if (!isUser && !immediate && conv.content) {
                typeWriter(document.getElementById(`msg-text-${idx}`), conv.content, callback);
            } else if (callback) {
                // 用户消息或直接显示，立即回调
                callback();
            }
        }

        // 更新Agent情绪状态
        function updateAgentMood(agentId, isChallenging, msgType) {
            if (msgType === 'conclusion') {
                agentMoods[agentId] = { icon: '😊', text: '满意', class: 'mood-satisfied' };
            } else if (isChallenging) {
                agentMoods[agentId] = { icon: '🤔', text: '质疑中', class: 'mood-challenging' };
            } else if (agentStatus[agentId] === 'speaking') {
                agentMoods[agentId] = { icon: '💭', text: '思考中', class: 'mood-thinking' };
            } else {
                agentMoods[agentId] = { icon: '👀', text: '关注', class: 'mood-worried' };
            }
            renderAgentList();
        }

        // 显示正在输入提示
        function showTypingIndicator(agentId) {
            const agent = AGENTS[agentId];
            if (!agent) return;

            const container = document.getElementById('chatMessages');
            const indicator = document.createElement('div');
            indicator.className = 'typing-indicator';
            indicator.id = 'typing-indicator';
            indicator.innerHTML = `
                <span style="color: ${agent.color}; font-weight: 600;">${agent.emoji} ${agent.name}</span>
                <span style="color: var(--text-muted);">正在输入</span>
                <div class="dots">
                    <div class="dot"></div>
                    <div class="dot"></div>
                    <div class="dot"></div>
                </div>
            `;
            container.appendChild(indicator);
            container.scrollTop = container.scrollHeight;
        }

        // 隐藏正在输入提示
        function hideTypingIndicator() {
            const indicator = document.getElementById('typing-indicator');
            if (indicator) indicator.remove();
        }

        // 点赞/踩
        function likeMessage(msgIdx, type) {
            messageLikes[msgIdx] = type === 'like' ? 'liked' : 'disliked';
            // 重新渲染该消息的按钮状态
            const msgDiv = document.querySelector(`[data-index="${msgIdx}"]`);
            if (msgDiv) {
                const likeBtn = msgDiv.querySelector('.action-btn:first-child');
                const dislikeBtn = msgDiv.querySelector('.action-btn:nth-child(2)');
                if (likeBtn) likeBtn.className = `action-btn ${type === 'like' ? 'liked' : ''}`;
                if (dislikeBtn) dislikeBtn.className = `action-btn ${type === 'dislike' ? 'disliked' : ''}`;
            }
        }

        // 追问
        function askMore(speaker, msgIdx) {
            const agent = AGENTS[speaker];
            if (!agent) return;
            const input = document.getElementById('taskInput');
            input.value = `@${agent.name} 请详细解释一下 `;
            input.focus();
        }

        // 处理打字机队列 - 带思考延迟
        function processTypingQueue() {
            if (isTyping || typingQueue.length === 0) return;

            isTyping = true;
            const task = typingQueue[0]; // 先不shift，等思考完成

            // 显示正在输入提示
            showTypingIndicator(task.speaker);

            // 思考延迟：根据内容长度动态调整
            const thinkTime = Math.min(Math.max(task.text.length * 10, 500), 3000);

            setTimeout(() => {
                hideTypingIndicator();
                typingQueue.shift(); // 现在才shift

                const element = document.getElementById(task.elementId);
                if (element) {
                    typeWriter(element, task.text, task.speed, () => {
                        isTyping = false;
                        setTimeout(processTypingQueue, 100);
                    });
                } else {
                    isTyping = false;
                    processTypingQueue();
                }
            }, thinkTime);
        }

        // 更新Agent卡片
        function updateAgentCards() {
            Object.keys(agentStatus).forEach(id => {
                const card = document.getElementById(`agent-${id}`);
                if (card) {
                    card.className = `agent-card ${agentStatus[id] === 'speaking' ? 'speaking' : ''} ${agentStatus[id] === 'challenge' ? 'challenge' : ''}`;
                }
            });
        }

        // 更新数据面板
        function updateDataPanel() {
            const challengeCount = conversations.filter(c => c.is_challenging).length;
            const totalAgentMsgs = conversations.filter(c => c.speaker !== 'user').length;
            // 基础共识度70%，每条质疑扣5分，最低30%
            let consensus = 70;
            if (totalAgentMsgs > 0) {
                consensus = Math.max(30, 70 - (challengeCount * 5));
                if (conversations.some(c => c.msg_type === 'conclusion')) {
                    consensus = Math.max(50, consensus); // 讨论完成至少50%
                }
            } else {
                consensus = 0;
            }

            const fill = document.getElementById('consensusFill');
            const value = document.getElementById('consensusValue');
            fill.style.width = consensus + '%';
            if (consensus < 30) fill.style.background = '#F56C6C';
            else if (consensus < 70) fill.style.background = '#E6A23C';
            else fill.style.background = '#67C23A';
            value.textContent = consensus + '%';

            document.getElementById('roundCount').textContent = `${Math.ceil(totalAgentMsgs / 2)} / 15`;
            document.getElementById('challengeCount').textContent = challengeCount;

            // 更新质疑增强面板
            updateChallengePanel();

            // 发言统计
            const stats = {};
            conversations.forEach(c => {
                if (c.speaker !== 'user') {
                    if (!stats[c.speaker]) stats[c.speaker] = 0;
                    stats[c.speaker]++;
                }
            });

            let statsHtml = '';
            const maxCount = Math.max(...Object.values(stats), 1);
            Object.entries(stats).forEach(([id, count]) => {
                const agent = AGENTS[id];
                if (agent) {
                    const percent = (count / maxCount) * 100;
                    statsHtml += `
                        <div class="speaker-bar">
                            <span class="speaker-name">${agent.name}</span>
                            <div class="speaker-progress">
                                <div class="speaker-fill" style="width: ${percent}%; background: ${agent.color}"></div>
                            </div>
                            <span class="speaker-count">${count}</span>
                        </div>
                    `;
                }
            });
            document.getElementById('speakerStats').innerHTML = statsHtml || '<div style="color:#909399;font-size:12px;text-align:center;">暂无发言</div>';
        }

        // 更新质疑增强面板
        async function updateChallengePanel() {
            try {
                const response = await fetch('/api/challenges/list');
                const data = await response.json();
                
                if (!data.success || !data.enabled) {
                    // 增强功能未启用，使用基础显示
                    return;
                }
                
                // 更新统计
                const challenges = data.challenges || [];
                const pendingCount = challenges.filter(c => c.status === 'pending').length;
                const acceptedCount = challenges.filter(c => c.status === 'accepted').length;
                const rejectedCount = challenges.filter(c => c.status === 'rejected').length;
                
                document.getElementById('pendingCount').textContent = pendingCount;
                document.getElementById('acceptedCount').textContent = acceptedCount;
                document.getElementById('rejectedCount').textContent = rejectedCount;
                
                // 更新质疑列表
                const listContainer = document.getElementById('challengeList');
                if (challenges.length === 0) {
                    listContainer.innerHTML = '<div style="color: #909399; font-size: 12px; text-align: center; padding: 10px;">暂无质疑记录</div>';
                    return;
                }
                
                let listHtml = '';
                challenges.slice(-5).forEach(c => {
                    const statusIcon = c.status === 'pending' ? '⏳' : 
                                       c.status === 'accepted' ? '✅' : 
                                       c.status === 'rejected' ? '❌' : '✓';
                    const intensityEmoji = c.intensity_emoji || '⚠️';
                    const bgColor = c.intensity_color || '#E6A23C';
                    
                    listHtml += `
                        <div style="background: white; border-radius: 4px; padding: 8px; margin-bottom: 6px; border-left: 3px solid ${bgColor};">
                            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 4px;">
                                <span style="font-size: 12px; font-weight: 500;">${c.challenger_name} → ${c.target_agent_name}</span>
                                <span style="font-size: 11px; background: ${bgColor}; color: white; padding: 1px 4px; border-radius: 2px;">${intensityEmoji} ${c.intensity_label}</span>
                            </div>
                            <div style="font-size: 11px; color: #606266; overflow: hidden; text-overflow: ellipsis; white-space: nowrap;">
                                ${c.reason || ''}
                            </div>
                            <div style="font-size: 10px; color: #909399; margin-top: 2px;">
                                ${statusIcon} ${c.category || '技术可行性'}
                            </div>
                        </div>
                    `;
                });
                
                listContainer.innerHTML = listHtml;
                
                // 如果有新的共识度信息，更新主面板
                if (data.consensus && data.consensus.total !== undefined) {
                    const fill = document.getElementById('consensusFill');
                    const value = document.getElementById('consensusValue');
                    const consensus = data.consensus.total;
                    
                    fill.style.width = consensus + '%';
                    if (consensus < 30) fill.style.background = '#F56C6C';
                    else if (consensus < 70) fill.style.background = '#E6A23C';
                    else fill.style.background = '#67C23A';
                    value.textContent = consensus + '%';
                }
                
            } catch (e) {
                console.error('[质疑面板更新失败]', e);
            }
        }

        // 处理键盘事件
        function handleKeyPress(event) {
            if (event.key === 'Enter' && !event.shiftKey) {
                event.preventDefault();
                submitTask();
            }
        }

        // 发送任务（修复版）
        async function submitTask() {
            console.log('[DEBUG] submitTask called, isProcessing:', isProcessing);
            
            if (isProcessing) {
                console.log('[WARN] 正在处理中，请稍候...');
                alert('正在处理中，请稍候...');
                return;
            }

            const input = document.getElementById('taskInput');
            const content = input.value.trim();
            if (!content) {
                console.log('[WARN] 任务内容为空');
                return;
            }

            // 彩蛋关键词检测
            for (const [keyword, response] of Object.entries(EASTER_EGGS)) {
                if (content.includes(keyword)) {
                    // 显示彩蛋反应
                    showEasterEgg(keyword, response);
                    input.value = '';
                    return;
                }
            }

            console.log('[DEBUG] 提交任务:', content);
            isProcessing = true;
            const sendBtn = document.getElementById('sendBtn');
            sendBtn.disabled = true;
            input.value = '';

            // 清空前端会话记录，开启新会话
            conversations = [];
            lastRenderedIndex = -1;
            renderMessages();  // 显示空状态
            console.log('[DEBUG] 已清空会话，开启新任务');

            // 显示加载状态
            showAgentLoading('扶摇', '🌀', '正在分析任务，调度Agent团队...');

            // 开始计时
            startTime = Date.now();
            if (timerInterval) clearInterval(timerInterval);
            timerInterval = setInterval(updateTimer, 1000);

            try {
                const response = await fetch('/api/task', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({task: content})
                });
                const result = await response.json();
                console.log('[DEBUG] 任务提交结果:', result);

                // 更新加载状态
                if (result.agents && result.agents.length > 0) {
                    updateAgentLoading(result.agents[0]);
                }
            } catch (e) {
                console.error('[ERROR] 提交失败:', e);
                hideAgentLoading();
                // 提交失败时立即解锁
                isProcessing = false;
                sendBtn.disabled = false;
            }

            // 不再使用60秒自动解锁，改用fetchStatus检测isCompleted
        }

        // ========== 加载状态管理 ==========
        function showAgentLoading(agentName, emoji, message) {
            const container = document.getElementById('chatMessages');

            // 移除旧的加载状态
            hideAgentLoading();

            const loadingDiv = document.createElement('div');
            loadingDiv.id = 'agentLoading';
            loadingDiv.className = 'agent-loading';
            loadingDiv.innerHTML = `
                <div class="agent-loading-avatar">${emoji}</div>
                <div class="agent-loading-text">
                    <div class="agent-loading-name">${agentName}</div>
                    <div>${message}</div>
                </div>
                <div class="typing-dots">
                    <span></span>
                    <span></span>
                    <span></span>
                </div>
            `;
            container.appendChild(loadingDiv);
            container.scrollTop = container.scrollHeight;
        }

        function updateAgentLoading(agentName) {
            const loadingDiv = document.getElementById('agentLoading');
            if (loadingDiv) {
                const agent = AGENTS[agentName] || {name: agentName, emoji: '🤖'};
                const nameDiv = loadingDiv.querySelector('.agent-loading-name');
                const textDiv = loadingDiv.querySelector('.agent-loading-text > div:last-child');
                const avatarDiv = loadingDiv.querySelector('.agent-loading-avatar');

                if (nameDiv) nameDiv.textContent = agent.name;
                if (avatarDiv) avatarDiv.textContent = agent.emoji;
                if (textDiv) textDiv.textContent = '正在分析发言...';
            }
        }

        function hideAgentLoading() {
            const loadingDiv = document.getElementById('agentLoading');
            if (loadingDiv) loadingDiv.remove();
        }

        // ========== 会话历史保存 ==========
        function saveToHistory() {
            const historyData = {
                conversations: conversations,
                consensus: consensus,
                timestamp: Date.now()
            };
            localStorage.setItem('chat_history', JSON.stringify(historyData));
            console.log('会话已保存');
        }

        function loadFromHistory() {
            const saved = localStorage.getItem('chat_history');
            if (saved) {
                try {
                    const data = JSON.parse(saved);
                    // 只加载24小时内的历史
                    if (Date.now() - data.timestamp < 24 * 60 * 60 * 1000) {
                        conversations = data.conversations || [];
                        consensus = data.consensus || 70;
                        console.log('已恢复历史会话:', conversations.length, '条');
                        return true;
                    }
                } catch (e) {
                    console.error('加载历史失败:', e);
                }
            }
            return false;
        }

        function clearHistory() {
            localStorage.removeItem('chat_history');
            conversations = [];
            lastRenderedIndex = -1;
            messageLikes = {};
            currentFilter = 'all';
            
            // 重置显示
            document.querySelectorAll('.filter-btn').forEach(btn => btn.classList.remove('active'));
            document.querySelector('.filter-btn').classList.add('active');
            
            // 清空聊天区域
            renderMessages();
            
            // 重置统计
            document.getElementById('consensusValue').textContent = '0%';
            document.getElementById('roundCount').textContent = '0 / 15';
            document.getElementById('challengeCount').textContent = '0';
            document.getElementById('speakerStats').innerHTML = '';
            
            // 重置Agent状态
            Object.keys(AGENTS).forEach(id => {
                agentStatus[id] = 'idle';
            });
            updateAgentCards();
            
            console.log('历史已清除，页面已重置');
        }

        // 更新计时器
        function updateTimer() {
            if (!startTime) return;
            const elapsed = Math.floor((Date.now() - startTime) / 1000);
            const mins = Math.floor(elapsed / 60).toString().padStart(2, '0');
            const secs = (elapsed % 60).toString().padStart(2, '0');
            document.getElementById('elapsedTime').textContent = `${mins}:${secs}`;
        }

        // 快捷任务
        function quickTask(task) {
            document.getElementById('taskInput').value = task;
            submitTask();
        }

        // ========== V15.2 产物生成功能 ==========
        let productGenerationInProgress = false;
        
        async function generateProduct() {
            if (productGenerationInProgress) {
                alert('产物正在生成中，请稍候...');
                return;
            }
            
            // 获取讨论内容作为输入
            const discussionContent = conversations.map(c => 
                `【${c.speaker_name || c.speaker}】${c.content}`
            ).join('\n\n');
            
            // 获取项目名称
            const projectName = memory.current_task || '未命名项目';
            
            // 显示生成中提示
            const chatMessages = document.getElementById('chatMessages');
            const loadingMsg = document.createElement('div');
            loadingMsg.className = 'message';
            loadingMsg.id = 'product-loading';
            loadingMsg.innerHTML = `
                <div class="message-avatar" style="background: #9C27B0;">🌿</div>
                <div class="message-content">
                    <div class="message-header">
                        <span class="message-speaker" style="color: #9C27B0;">🌿 南乔</span>
                        <span class="message-time">产物生成</span>
                    </div>
                    <div class="message-text">
                        🚀 正在调用V15.2产物生成引擎...<br>
                        <span style="color: var(--text-muted);">这可能需要几秒钟</span>
                    </div>
                </div>
            `;
            chatMessages.appendChild(loadingMsg);
            chatMessages.scrollTop = chatMessages.scrollHeight;
            
            productGenerationInProgress = true;
            
            try {
                // 调用V15.2产物生成API
                const response = await fetch('/api/v15.2/product/generate', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({
                        type: 'document',  // 产物类型：code/document/model/pipeline
                        name: projectName + ' - 讨论产物',
                        description: '基于多Agent协作讨论生成的产物文档',
                        input_data: {
                            requirements: [discussionContent],
                            project_name: projectName,
                            consensus_level: memory.get_consensus_level ? memory.get_consensus_level() : 70
                        },
                        tags: ['协作讨论', '自动生成', projectName]
                    })
                });
                
                const result = await response.json();
                
                if (result.success) {
                    const taskId = result.data.task_id;
                    
                    // 更新提示
                    loadingMsg.querySelector('.message-text').innerHTML = `
                        ✅ 产物生成任务已提交<br>
                        <span style="color: var(--text-muted);">任务ID: ${taskId}</span><br>
                        <span style="color: var(--text-muted);">正在生成中，请稍候...</span>
                    `;
                    
                    // 轮询任务状态
                    const statusUrl = result.data.status_url;
                    let completed = false;
                    let attempts = 0;
                    const maxAttempts = 30;  // 最多等待30次（约30秒）
                    
                    while (!completed && attempts < maxAttempts) {
                        await new Promise(r => setTimeout(r, 1000));
                        
                        const statusResponse = await fetch(statusUrl);
                        const statusData = await statusResponse.json();
                        
                        if (statusData.success && statusData.data) {
                            const status = statusData.data.status;
                            const progress = statusData.data.progress || 0;
                            
                            loadingMsg.querySelector('.message-text').innerHTML = `
                                ⏳ 产物生成中...<br>
                                <span style="color: var(--text-muted);">进度: ${progress}% - ${statusData.data.stage || '处理中'}</span>
                            `;
                            
                            if (status === 'completed') {
                                completed = true;
                                
                                // 显示完成信息
                                loadingMsg.querySelector('.message-text').innerHTML = `
                                    🎉 产物生成完成！<br>
                                    <span style="color: var(--success);">质量评分: ${statusData.data.result ? '良好' : '待评估'}</span><br><br>
                                    <button class="quick-btn" onclick="downloadProduct('${taskId}')" style="background: var(--primary); color: white;">
                                        📥 下载产物
                                    </button>
                                    <button class="quick-btn" onclick="checkProductQuality('${taskId}')">
                                        🔍 质量检查
                                    </button>
                                `;
                            } else if (status === 'failed') {
                                completed = true;
                                loadingMsg.querySelector('.message-text').innerHTML = `
                                    ❌ 产物生成失败<br>
                                    <span style="color: var(--danger);">${statusData.data.error || '未知错误'}</span>
                                `;
                            }
                        }
                        
                        attempts++;
                    }
                    
                    if (!completed) {
                        loadingMsg.querySelector('.message-text').innerHTML = `
                            ⚠️ 产物生成超时<br>
                            <span style="color: var(--text-muted);">请稍后手动刷新查看结果</span>
                        `;
                    }
                    
                } else {
                    loadingMsg.querySelector('.message-text').innerHTML = `
                        ❌ 产物生成失败<br>
                        <span style="color: var(--danger);">${result.message || '服务异常'}</span>
                    `;
                }
                
            } catch (e) {
                console.error('[产物生成错误]', e);
                loadingMsg.querySelector('.message-text').innerHTML = `
                    ❌ 产物生成异常<br>
                    <span style="color: var(--danger);">${e.message || '网络错误'}</span>
                `;
            } finally {
                productGenerationInProgress = false;
            }
        }
        
        // 下载产物
        async function downloadProduct(taskId) {
            window.open(`/api/v15.2/product/download/${taskId}?format=zip`, '_blank');
        }
        
        // 质量检查
        async function checkProductQuality(taskId) {
            const chatMessages = document.getElementById('chatMessages');
            const qualityMsg = document.createElement('div');
            qualityMsg.className = 'message';
            qualityMsg.innerHTML = `
                <div class="message-avatar" style="background: #E6A23C;">🔍</div>
                <div class="message-content">
                    <div class="message-header">
                        <span class="message-speaker" style="color: #E6A23C;">🔍 质量检查</span>
                    </div>
                    <div class="message-text">正在执行质量检查...</div>
                </div>
            `;
            chatMessages.appendChild(qualityMsg);
            chatMessages.scrollTop = chatMessages.scrollHeight;
            
            try {
                const response = await fetch('/api/v15.2/quality/check', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({
                        target_type: 'product',
                        target_id: taskId,
                        check_types: ['code_quality', 'security', 'performance']
                    })
                });
                
                const result = await response.json();
                
                if (result.success) {
                    const checkId = result.data.check_id;
                    
                    // 等待检查完成
                    await new Promise(r => setTimeout(r, 3000));
                    
                    // 获取报告
                    const reportResponse = await fetch(`/api/v15.2/quality/report/${checkId}`);
                    const reportData = await reportResponse.json();
                    
                    if (reportData.success) {
                        const summary = reportData.data.summary;
                        qualityMsg.querySelector('.message-text').innerHTML = `
                            ✅ 质量检查完成<br>
                            <span style="color: var(--text-muted);">评分: ${summary.score}分</span><br>
                            <span style="color: ${summary.critical > 0 ? 'var(--danger)' : 'var(--success)'};">
                                严重问题: ${summary.critical}个
                            </span><br>
                            <span style="color: var(--warning);">警告: ${summary.warning}个</span><br>
                            <span style="color: var(--info);">提示: ${summary.info}个</span>
                        `;
                    }
                } else {
                    qualityMsg.querySelector('.message-text').innerHTML = `
                        ⚠️ 质量检查失败<br>
                        <span style="color: var(--text-muted);">${result.message || '服务异常'}</span>
                    `;
                }
            } catch (e) {
                qualityMsg.querySelector('.message-text').innerHTML = `
                    ❌ 质量检查异常<br>
                    <span style="color: var(--danger);">${e.message}</span>
                `;
            }
        }

        // 清空
        async function clearChat() {
            await fetch('/api/clear', {method: 'POST'});
            if (timerInterval) clearInterval(timerInterval);
            document.getElementById('elapsedTime').textContent = '00:00';
            startTime = null;
            // 重置打字机状态
            typingQueue = [];
            isTyping = false;
            lastRenderedIndex = -1;
            // 清除历史
            clearHistory();
            // 隐藏"生成产物"按钮
            const productBtn = document.getElementById('generateProductBtn');
            if (productBtn) {
                productBtn.style.display = 'none';
            }
        }

        // 上传文档
        async function uploadDocument(input) {
            const file = input.files[0];
            if (!file) return;

            // 显示上传提示
            const chatMessages = document.getElementById('chatMessages');
            const uploadMsg = document.createElement('div');
            uploadMsg.className = 'message system';
            uploadMsg.innerHTML = `<div class="message-content">📎 正在解析文档: ${file.name}...</div>`;
            chatMessages.appendChild(uploadMsg);
            chatMessages.scrollTop = chatMessages.scrollHeight;

            // 上传文件
            const formData = new FormData();
            formData.append('document', file);

            try {
                const response = await fetch('/api/parse/document', {
                    method: 'POST',
                    body: formData
                });

                const result = await response.json();

                if (result.status === 'ok') {
                    // 更新上传提示
                    uploadMsg.innerHTML = `
                        <div class="message-content">
                            <div style="font-weight: 600; margin-bottom: 8px;">📄 文档解析完成</div>
                            <div style="font-size: 12px; color: var(--text-secondary); margin-bottom: 8px;">
                                文件: ${result.title || file.name}<br>
                                类型: ${result.file_type}<br>
                                字数: ${result.word_count}<br>
                                章节: ${result.sections ? result.sections.length : 0}个
                            </div>
                            <div style="background: var(--bg-main); padding: 12px; border-radius: 8px; margin-bottom: 8px; max-height: 200px; overflow-y: auto; font-size: 13px;">
                                ${result.content.substring(0, 500)}${result.content.length > 500 ? '...' : ''}
                            </div>
                            <button class="quick-btn" onclick="useParsedContent('${encodeURIComponent(result.content.substring(0, 1000))}')">
                                📝 使用此内容
                            </button>
                        </div>
                    `;

                    // 存储解析结果
                    window.parsedDocument = result;
                } else {
                    uploadMsg.innerHTML = `<div class="message-content" style="color: var(--error);">❌ 解析失败: ${result.message}</div>`;
                }
            } catch (error) {
                uploadMsg.innerHTML = `<div class="message-content" style="color: var(--error);">❌ 上传失败: ${error.message}</div>`;
            }

            // 清空文件选择
            input.value = '';
        }

        // 使用解析后的内容
        function useParsedContent(encodedContent) {
            const content = decodeURIComponent(encodedContent);
            document.getElementById('taskInput').value = `基于以下文档内容进行需求分析：
` + content.substring(0, 200);
            document.getElementById('taskInput').focus();
        }

        // 筛选
        function filterMessages(type) {
            currentFilter = type;
            document.querySelectorAll('.filter-btn').forEach(btn => btn.classList.remove('active'));
            event.target.classList.add('active');
            
            // 停止打字机效果
            typingQueue = [];
            isTyping = false;
            
            // 重置渲染索引，重新渲染所有消息
            lastRenderedIndex = -1;
            const container = document.getElementById('chatMessages');
            container.innerHTML = '';
            
            // 直接渲染所有消息（不用打字机效果）
            conversations.forEach((conv, idx) => {
                const agent = AGENTS[conv.speaker] || { name: conv.speaker, emoji: '🤖', color: '#909399', role: '' };
                const isUser = conv.speaker === 'user';
                const time = conv.timestamp.split('T')[1].split('.')[0].substring(0, 5);

                // 判断消息类型
                let msgType = 'support';
                if (conv.is_challenging) msgType = 'challenge';
                if (conv.msg_type === 'conclusion') msgType = 'decision';

                // 筛选
                if (currentFilter !== 'all') {
                    if (currentFilter === 'challenge' && !conv.is_challenging) return;
                    if (currentFilter === 'support' && conv.is_challenging) return;
                    if (currentFilter === 'decision' && conv.msg_type !== 'conclusion') return;
                }

                // 创建消息元素
                const msgDiv = document.createElement('div');
                msgDiv.className = `message ${isUser ? 'user' : ''}`;
                msgDiv.setAttribute('data-speaker', conv.speaker);
                msgDiv.setAttribute('data-type', msgType);
                msgDiv.setAttribute('data-index', idx);

                const actionsHtml = !isUser ? `
                    <div class="message-actions">
                        <button class="action-btn" data-idx="${idx}" data-type="like">👍 认可</button>
                        <button class="action-btn" data-idx="${idx}" data-type="dislike">👎 存疑</button>
                        <button class="action-btn" data-speaker="${conv.speaker}" data-idx="${idx}" data-action="ask">💬 追问</button>
                        <span class="message-action" data-text="${encodeURIComponent(conv.content)}">📋 复制</span>
                    </div>
                ` : '';

                msgDiv.innerHTML = `
                    ${!isUser ? `<div class="message-avatar" style="background: ${agent.color}">${agent.emoji}</div>` : ''}
                    <div class="message-content">
                        <div class="message-header">
                            <span class="message-speaker" style="color: ${isUser ? 'white' : agent.color}">${isUser ? '👤 少帅' : agent.emoji + ' ' + agent.name}${!isUser && agent.role ? ' · ' + agent.role : ''}</span>
                            ${conv.is_challenging ? '<span class="message-badge">⚠️ 质疑</span>' : ''}
                            ${conv.msg_type === 'conclusion' ? '<span class="message-badge" style="background:#E8F5E9;color:#67C23A">✓ 决策</span>' : ''}
                            <span class="message-time">${time}</span>
                        </div>
                        <div class="message-text">${conv.content}</div>
                        ${actionsHtml}
                    </div>
                `;

                container.appendChild(msgDiv);
            });
            
            lastRenderedIndex = conversations.length - 1;
            
            // 滚动到底部
            container.scrollTop = container.scrollHeight;
        }

        function filterByAgent(agentId) {
            currentFilter = 'all';
            document.querySelectorAll('.filter-btn').forEach(btn => btn.classList.remove('active'));
            document.querySelector('.filter-btn').classList.add('active');
            renderMessages();

            const msgs = document.querySelectorAll(`[data-speaker="${agentId}"]`);
            if (msgs.length > 0) {
                msgs[0].scrollIntoView({behavior: 'smooth', block: 'center'});
            }
        }

        // 复制消息
        function copyMessage(element) {
            const text = decodeURIComponent(element.getAttribute('data-text'));
            navigator.clipboard.writeText(text).then(() => {
                element.textContent = '✓ 已复制';
                setTimeout(() => {
                    element.textContent = '📋 复制';
                }, 1500);
            });
        }

        // ========== 场景模板面板 ==========
        let scenePanelVisible = false;

        async function showScenePanel() {
            if (scenePanelVisible) {
                hideScenePanel();
                return;
            }

            // 获取场景模板列表
            const response = await fetch('/api/scenes/list');
            const result = await response.json();

            if (result.status !== 'ok') {
                alert('获取场景模板失败');
                return;
            }

            // 创建面板
            const panel = document.createElement('div');
            panel.id = 'scenePanel';
            panel.style.cssText = `
                position: fixed;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                background: var(--bg-card);
                border-radius: 12px;
                padding: 24px;
                box-shadow: 0 8px 32px rgba(0,0,0,0.2);
                z-index: 1000;
                width: 500px;
                max-height: 70vh;
                overflow-y: auto;
            `;

            let html = `
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 16px;">
                    <h3 style="margin: 0;">🎭 场景模板库</h3>
                    <button onclick="hideScenePanel()" style="background: none; border: none; font-size: 20px; cursor: pointer;">✕</button>
                </div>
                <div style="color: var(--text-secondary); font-size: 13px; margin-bottom: 16px;">
                    选择预设场景，一键启动专业讨论
                </div>
            `;

            for (const template of result.templates) {
                html += `
                    <div class="scene-card" onclick="selectScene('${template.id}')"
                         style="padding: 12px; margin-bottom: 8px; background: var(--bg-main); border-radius: 8px; cursor: pointer; border-left: 3px solid var(--primary);">
                        <div style="font-weight: 600; margin-bottom: 4px;">${template.name}</div>
                        <div style="font-size: 12px; color: var(--text-secondary);">${template.description}</div>
                        <div style="font-size: 11px; color: var(--text-muted); margin-top: 6px;">预计 ${template.estimated_rounds} 轮</div>
                    </div>
                `;
            }

            panel.innerHTML = html;
            document.body.appendChild(panel);

            // 添加遮罩
            const overlay = document.createElement('div');
            overlay.id = 'sceneOverlay';
            overlay.style.cssText = `
                position: fixed;
                top: 0;
                left: 0;
                right: 0;
                bottom: 0;
                background: rgba(0,0,0,0.5);
                z-index: 999;
            `;
            overlay.onclick = hideScenePanel;
            document.body.appendChild(overlay);

            scenePanelVisible = true;
        }

        function hideScenePanel() {
            const panel = document.getElementById('scenePanel');
            const overlay = document.getElementById('sceneOverlay');
            if (panel) panel.remove();
            if (overlay) overlay.remove();
            scenePanelVisible = false;
        }

        async function selectScene(sceneId) {
            const response = await fetch(`/api/scenes/${sceneId}`);
            const result = await response.json();

            if (result.status !== 'ok') {
                alert('获取模板失败');
                return;
            }

            const template = result.template;

            // 填充参数（使用默认值）
            const params = {};
            for (const param of template.params) {
                params[param.name] = param.default;
            }

            // 填充模板
            const fillResponse = await fetch('/api/scenes/fill', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({scene_id: sceneId, params: params})
            });

            const fillResult = await fillResponse.json();

            if (fillResult.status === 'ok') {
                // 填充到输入框
                document.getElementById('taskInput').value = fillResult.result.task;
                hideScenePanel();
            }
        }

        // ========== 成本估算面板 ==========
        let costPanelVisible = false;

        function showCostPanel() {
            if (costPanelVisible) {
                hideCostPanel();
                return;
            }

            // 创建面板
            const panel = document.createElement('div');
            panel.id = 'costPanel';
            panel.style.cssText = `
                position: fixed;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                background: var(--bg-card);
                border-radius: 12px;
                padding: 24px;
                box-shadow: 0 8px 32px rgba(0,0,0,0.2);
                z-index: 1000;
                width: 450px;
            `;

            panel.innerHTML = `
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 16px;">
                    <h3 style="margin: 0;">💰 成本估算</h3>
                    <button onclick="hideCostPanel()" style="background: none; border: none; font-size: 20px; cursor: pointer;">✕</button>
                </div>

                <div style="margin-bottom: 12px;">
                    <label style="font-size: 13px; color: var(--text-secondary);">项目名称</label>
                    <input type="text" id="costProjectName" value="智能客服系统"
                           style="width: 100%; padding: 8px 12px; border: 1px solid var(--border); border-radius: 6px; margin-top: 4px;">
                </div>

                <div style="display: flex; gap: 12px; margin-bottom: 12px;">
                    <div style="flex: 1;">
                        <label style="font-size: 13px; color: var(--text-secondary);">团队规模（人）</label>
                        <input type="number" id="costTeamSize" value="5"
                               style="width: 100%; padding: 8px 12px; border: 1px solid var(--border); border-radius: 6px; margin-top: 4px;">
                    </div>
                    <div style="flex: 1;">
                        <label style="font-size: 13px; color: var(--text-secondary);">周期（月）</label>
                        <input type="number" id="costDuration" value="6"
                               style="width: 100%; padding: 8px 12px; border: 1px solid var(--border); border-radius: 6px; margin-top: 4px;">
                    </div>
                </div>

                <div style="display: flex; gap: 12px; margin-bottom: 16px;">
                    <div style="flex: 1;">
                        <label style="font-size: 13px; color: var(--text-secondary);">复杂度</label>
                        <select id="costComplexity" style="width: 100%; padding: 8px 12px; border: 1px solid var(--border); border-radius: 6px; margin-top: 4px;">
                            <option value="简单">简单</option>
                            <option value="中等" selected>中等</option>
                            <option value="复杂">复杂</option>
                            <option value="极复杂">极复杂</option>
                        </select>
                    </div>
                    <div style="flex: 1;">
                        <label style="font-size: 13px; color: var(--text-secondary);">云服务规格</label>
                        <select id="costCloudTier" style="width: 100%; padding: 8px 12px; border: 1px solid var(--border); border-radius: 6px; margin-top: 4px;">
                            <option value="小型">小型</option>
                            <option value="中型" selected>中型</option>
                            <option value="大型">大型</option>
                        </select>
                    </div>
                </div>

                <button onclick="calculateCost()"
                        style="width: 100%; padding: 12px; background: var(--primary); color: white; border: none; border-radius: 8px; cursor: pointer; font-weight: 600;">
                    开始估算
                </button>

                <div id="costResult" style="margin-top: 16px; display: none;">
                </div>
            `;

            document.body.appendChild(panel);

            // 添加遮罩
            const overlay = document.createElement('div');
            overlay.id = 'costOverlay';
            overlay.style.cssText = `
                position: fixed;
                top: 0;
                left: 0;
                right: 0;
                bottom: 0;
                background: rgba(0,0,0,0.5);
                z-index: 999;
            `;
            overlay.onclick = hideCostPanel;
            document.body.appendChild(overlay);

            costPanelVisible = true;
        }

        function hideCostPanel() {
            const panel = document.getElementById('costPanel');
            const overlay = document.getElementById('costOverlay');
            if (panel) panel.remove();
            if (overlay) overlay.remove();
            costPanelVisible = false;
        }
        
        async function calculateCost() {
            const data = {
                project_name: document.getElementById('costProjectName').value,
                team_size: parseInt(document.getElementById('costTeamSize').value),
                duration_months: parseInt(document.getElementById('costDuration').value),
                complexity: document.getElementById('costComplexity').value,
                cloud_tier: document.getElementById('costCloudTier').value
            };
            
            const response = await fetch('/api/cost/estimate', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify(data)
            });
            
            const result = await response.json();
            
            if (result.status === 'ok') {
                const est = result.estimation;
                document.getElementById('costResult').style.display = 'block';
                document.getElementById('costResult').innerHTML = `
                    <div style="background: var(--bg-main); padding: 16px; border-radius: 8px;">
                        <div style="display: flex; justify-content: space-between; margin-bottom: 8px;">
                            <span style="color: var(--text-secondary);">人天数</span>
                            <span style="font-weight: 600;">${est.man_days}人天</span>
                        </div>
                        <div style="display: flex; justify-content: space-between; margin-bottom: 8px;">
                            <span style="color: var(--text-secondary);">人力成本</span>
                            <span style="font-weight: 600;">${est.labor_cost.toLocaleString()}元</span>
                        </div>
                        <div style="display: flex; justify-content: space-between; margin-bottom: 8px;">
                            <span style="color: var(--text-secondary);">云服务成本</span>
                            <span style="font-weight: 600;">${est.cloud_cost.toLocaleString()}元</span>
                        </div>
                        <div style="display: flex; justify-content: space-between; margin-bottom: 8px;">
                            <span style="color: var(--text-secondary);">风险储备</span>
                            <span style="font-weight: 600;">${est.risk_reserve.toLocaleString()}元</span>
                        </div>
                        <div style="border-top: 1px solid var(--border); margin: 12px 0; padding-top: 12px;">
                            <div style="display: flex; justify-content: space-between;">
                                <span style="font-weight: 600;">总成本</span>
                                <span style="font-size: 18px; font-weight: 700; color: var(--primary);">¥${est.total_cost.toLocaleString()}</span>
                            </div>
                        </div>
                    </div>
                `;
            }
        }
        
        // ========== 流式输出测试 ==========
        async function testStreamOutput() {
            const container = document.getElementById('chatMessages');
            
            // 创建测试消息
            const msgDiv = document.createElement('div');
            msgDiv.className = 'message';
            msgDiv.innerHTML = `
                <div class="message-avatar" style="background: #9C27B0;">🌿</div>
                <div class="message-content">
                    <div class="message-header">
                        <span class="message-speaker" style="color: #9C27B0;">🌿 南乔</span>
                        <span class="message-time">流式测试</span>
                    </div>
                    <div class="message-text" id="stream-test-text"><span class="typing-cursor">▊</span></div>
                </div>
            `;
            container.appendChild(msgDiv);
            container.scrollTop = container.scrollHeight;
            
            const textDiv = document.getElementById('stream-test-text');
            let fullText = '';
            
            try {
                const response = await fetch('/api/stream/test');
                const reader = response.body.getReader();
                const decoder = new TextDecoder();
                
                while (true) {
                    const {done, value} = await reader.read();
                    if (done) break;
                    
                    const chunk = decoder.decode(value);
                    const lines = chunk.split(`
`);
                    
                    for (const line of lines) {
                        if (line.startsWith('data: ')) {
                            try {
                                const data = JSON.parse(line.substring(6));
                                if (data.text) {
                                    fullText += data.text;
                                    textDiv.innerHTML = fullText + '<span class="typing-cursor">▊</span>';
                                    container.scrollTop = container.scrollHeight;
                                }
                                if (data.done) {
                                    textDiv.innerHTML = fullText;
                                }
                            } catch (e) {}
                        }
                    }
                }
            } catch (e) {
                console.error('流式输出测试失败:', e);
                textDiv.innerHTML = '流式输出测试失败: ' + e.message;
            }
        }
        
        // ========== 真实流式输出 ==========
        async function streamAgentResponse(agentId, systemPrompt, userMessage) {
            const container = document.getElementById('chatMessages');
            const agent = AGENTS[agentId] || {name: agentId, emoji: '🤖', color: '#909399'};
            
            // 创建消息
            const msgDiv = document.createElement('div');
            msgDiv.className = 'message';
            msgDiv.innerHTML = `
                <div class="message-avatar" style="background: ${agent.color};">${agent.emoji}</div>
                <div class="message-content">
                    <div class="message-header">
                        <span class="message-speaker" style="color: ${agent.color};">${agent.emoji} ${agent.name}</span>
                    </div>
                    <div class="message-text" id="stream-${agentId}"><span class="typing-cursor">▊</span></div>
                </div>
            `;
            container.appendChild(msgDiv);
            container.scrollTop = container.scrollHeight;
            
            const textDiv = document.getElementById(`stream-${agentId}`);
            let fullText = '';
            
            try {
                const response = await fetch('/api/stream', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({
                        system_prompt: systemPrompt,
                        user_message: userMessage
                    })
                });
                
                const reader = response.body.getReader();
                const decoder = new TextDecoder();
                
                while (true) {
                    const {done, value} = await reader.read();
                    if (done) break;
                    
                    const chunk = decoder.decode(value);
                    const lines = chunk.split(`
`);
                    
                    for (const line of lines) {
                        if (line.startsWith('data: ')) {
                            try {
                                const data = JSON.parse(line.substring(6));
                                if (data.text) {
                                    fullText += data.text;
                                    textDiv.innerHTML = fullText + '<span class="typing-cursor">▊</span>';
                                    container.scrollTop = container.scrollHeight;
                                }
                                if (data.error) {
                                    textDiv.innerHTML = `<span style="color: red;">错误: ${data.error}</span>`;
                                }
                                if (data.done) {
                                    textDiv.innerHTML = fullText;
                                }
                            } catch (e) {}
                        }
                    }
                }
            } catch (e) {
                console.error('流式输出失败:', e);
                textDiv.innerHTML = '流式输出失败: ' + e.message;
            }
            
            return fullText;
        }
        
        // 导出Word
        async function exportWord() {
            if (conversations.length === 0) {
                alert('暂无对话可导出');
                return;
            }

            try {
                const response = await fetch('/api/export/word', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'}
                });

                if (response.ok) {
                    const blob = await response.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `讨论报告_${new Date().toISOString().slice(0,10)}.docx`;
                    a.click();
                    URL.revokeObjectURL(url);
                } else {
                    alert('导出失败，请稍后重试');
                }
            } catch (e) {
                console.error('导出Word失败:', e);
                alert('导出失败：' + e.message);
            }
        }

        // 导出Excel
        async function exportExcel() {
            if (conversations.length === 0) {
                alert('暂无对话可导出');
                return;
            }

            try {
                const response = await fetch('/api/export/excel', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'}
                });

                if (response.ok) {
                    const blob = await response.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `讨论报告_${new Date().toISOString().slice(0,10)}.xlsx`;
                    a.click();
                    URL.revokeObjectURL(url);
                } else {
                    alert('导出失败，请稍后重试');
                }
            } catch (e) {
                console.error('导出Excel失败:', e);
                alert('导出失败：' + e.message);
            }
        }

        // 导出Markdown
        async function exportMarkdown() {
            if (conversations.length === 0) {
                alert('暂无对话可导出');
                return;
            }

            try {
                const response = await fetch('/api/export/markdown', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'}
                });

                const result = await response.json();

                if (result.status === 'ok') {
                    const blob = new Blob([result.content], {type: 'text/markdown'});
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `讨论报告_${new Date().toISOString().slice(0,10)}.md`;
                    a.click();
                    URL.revokeObjectURL(url);
                } else {
                    alert('导出失败：' + result.message);
                }
            } catch (e) {
                console.error('导出Markdown失败:', e);
                alert('导出失败：' + e.message);
            }
        }

        // 原有的导出函数保留（向后兼容）
        function exportChat() {
            exportMarkdown();
        }
        
        // ========== 会议纪要预览与下载 ==========
        async function previewMinutes() {
            const modal = document.getElementById('minutesModal');
            const content = document.getElementById('minutesContent');
            modal.style.display = 'block';
            content.innerHTML = '<div style="text-align: center; padding: 40px; color: #999;">正在加载会议纪要...</div>';
            
            try {
                const response = await fetch('/api/minutes/preview');
                const result = await response.json();
                
                if (result.status === 'ok') {
                    // 渲染Markdown为HTML
                    content.innerHTML = renderMarkdown(result.content);
                } else {
                    content.innerHTML = `<div style="text-align: center; padding: 40px; color: #F56C6C;">${result.message || '暂无会议纪要'}</div>`;
                }
            } catch (e) {
                console.error('加载会议纪要失败:', e);
                content.innerHTML = '<div style="text-align: center; padding: 40px; color: #F56C6C;">加载失败：' + e.message + '</div>';
            }
        }
        
        function closeMinutesModal() {
            document.getElementById('minutesModal').style.display = 'none';
        }
        
        async function downloadMinutes() {
            try {
                const response = await fetch('/api/minutes/download?format=md');
                if (response.ok) {
                    const blob = await response.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `会议纪要_${new Date().toISOString().slice(0,10)}.md`;
                    a.click();
                    URL.revokeObjectURL(url);
                } else {
                    alert('下载失败：暂无会议纪要');
                }
            } catch (e) {
                console.error('下载失败:', e);
                alert('下载失败：' + e.message);
            }
        }
        
        async function downloadMinutesWord() {
            try {
                const response = await fetch('/api/minutes/download?format=word');
                if (response.ok) {
                    const blob = await response.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `会议纪要_${new Date().toISOString().slice(0,10)}.docx`;
                    a.click();
                    URL.revokeObjectURL(url);
                } else {
                    alert('下载失败：暂无会议纪要');
                }
            } catch (e) {
                console.error('下载失败:', e);
                alert('下载失败：' + e.message);
            }
        }
        
        async function downloadMinutesPDF() {
            try {
                const response = await fetch('/api/minutes/download?format=pdf');
                if (response.ok) {
                    const blob = await response.blob();
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `会议纪要_${new Date().toISOString().slice(0,10)}.pdf`;
                    a.click();
                    URL.revokeObjectURL(url);
                } else {
                    alert('下载失败：暂无会议纪要');
                }
            } catch (e) {
                console.error('下载失败:', e);
                alert('下载失败：' + e.message);
            }
        }
        
        // 简单的Markdown渲染器
        function renderMarkdown(md) {
            return md
                .replace(/^# (.*)$/gm, '<h1 style="color:#C93832;border-bottom:2px solid #C93832;padding-bottom:10px;">$1</h1>')
                .replace(/^## (.*)$/gm, '<h2 style="color:#006EBD;margin-top:20px;">$1</h2>')
                .replace(/^### (.*)$/gm, '<h3 style="color:#595959;margin-top:15px;">$1</h3>')
                .replace(/\*\*(.*?)\*\*/g, '<strong>$1</strong>')
                .replace(/\*(.*?)\*/g, '<em>$1</em>')
                .replace(/^- (.*)$/gm, '<li style="margin-left:20px;">$1</li>')
                .replace(/^\|(.*)\|$/gm, function(match) {
                    const cells = match.split('|').filter(c => c.trim());
                    if (cells.some(c => c.includes('---'))) {
                        return ''; // 跳过分隔行
                    }
                    return '<tr>' + cells.map(c => `<td style="border:1px solid #ddd;padding:8px;">${c.trim()}</td>`).join('') + '</tr>';
                })
                .replace(/(<tr>.*<\/tr>)+/g, '<table style="width:100%;border-collapse:collapse;margin:10px 0;">$&</table>')
                .replace(/^---$/gm, '<hr style="border:none;border-top:1px solid #ddd;margin:20px 0;">')
                .replace(/\\n\\n/g, '</p><p style="margin:10px 0;">')
                .replace(/^(?!<[ht])/gm, '<p style="margin:10px 0;">');
        }

        // 键盘快捷键
        document.addEventListener('keydown', (e) => {
            if (e.ctrlKey && e.key === 'e') {
                e.preventDefault();
                exportChat();
            }
        });

        // 事件委托 - 处理消息按钮点击
        document.getElementById('chatMessages').addEventListener('click', function(e) {
            const target = e.target;

            // 认可/存疑按钮
            if (target.classList.contains('action-btn') && target.dataset.idx) {
                const idx = parseInt(target.dataset.idx);
                const type = target.dataset.type;

                if (type === 'like' || type === 'dislike') {
                    messageLikes[idx] = type === 'like' ? 'liked' : 'disliked';

                    // 更新按钮样式
                    const msgDiv = target.closest('.message');
                    if (msgDiv) {
                        const buttons = msgDiv.querySelectorAll('.action-btn[data-type]');
                        buttons.forEach(btn => {
                            const btnType = btn.dataset.type;
                            if (btnType === 'like') {
                                btn.className = `action-btn ${type === 'like' ? 'liked' : ''}`;
                            } else if (btnType === 'dislike') {
                                btn.className = `action-btn ${type === 'dislike' ? 'disliked' : ''}`;
                            }
                        });
                    }
                }

                // 追问按钮
                if (target.dataset.action === 'ask') {
                    const speaker = target.dataset.speaker;
                    const agent = AGENTS[speaker];
                    if (agent) {
                        const input = document.getElementById('taskInput');
                        input.value = `@${agent.name} 请详细解释一下 `;
                        input.focus();
                    }
                }
            }

            // 复制按钮
            if (target.classList.contains('message-action') && target.dataset.text) {
                const text = decodeURIComponent(target.dataset.text);
                navigator.clipboard.writeText(text).then(() => {
                    const originalText = target.textContent;
                    target.textContent = '✓ 已复制';
                    setTimeout(() => {
                        target.textContent = originalText;
                    }, 1500);
                });
            }
        });

        // 彩蛋反应 - 持久化到conversations
        function showEasterEgg(keyword, response) {
            // 添加用户消息
            conversations.push({
                speaker: 'user',
                speaker_name: '用户',
                content: keyword,
                timestamp: new Date().toISOString(),
                is_challenging: false
            });

            // 所有Agent一起反应
            const reactions = [
                { id: 'caiwei', text: '😭 为什么要加班！我们强烈建议工作生活平衡！' },
                { id: 'zhijin', text: '🧵 同意！工作生活平衡很重要！' },
                { id: 'yuheng', text: '⚖️ 合理安排时间，效率更高！' },
                { id: 'zhutai', text: '🏗️ 成本不仅是金钱，还有健康！' },
                { id: 'nanqiao', text: '🌿 少帅，注意休息哦~' }
            ];

            reactions.forEach((r, idx) => {
                setTimeout(() => {
                    const agent = AGENTS[r.id];
                    conversations.push({
                        speaker: r.id,
                        speaker_name: agent.name,
                        content: r.text,
                        timestamp: new Date().toISOString(),
                        is_challenging: false
                    });
                    renderMessages();
                }, idx * 500);
            });
        }

        // 启动
        init();
    </script>
    
    <!-- 会议纪要预览弹窗 -->
    <div id="minutesModal" class="modal" style="display: none; position: fixed; top: 0; left: 0; width: 100%; height: 100%; background: rgba(0,0,0,0.5); z-index: 1000;">
        <div class="modal-content" style="position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); background: white; width: 80%; max-width: 900px; max-height: 80vh; border-radius: 12px; overflow: hidden; display: flex; flex-direction: column;">
            <div class="modal-header" style="padding: 16px 20px; background: linear-gradient(135deg, #C93832, #006EBD); color: white; display: flex; justify-content: space-between; align-items: center;">
                <h3 style="margin: 0; font-size: 18px;">📋 会议纪要</h3>
                <button onclick="closeMinutesModal()" style="background: none; border: none; color: white; font-size: 24px; cursor: pointer;">&times;</button>
            </div>
            <div class="modal-body" id="minutesContent" style="padding: 20px; overflow-y: auto; flex: 1;">
                <div style="text-align: center; padding: 40px; color: #999;">正在加载会议纪要...</div>
            </div>
            <div class="modal-footer" style="padding: 12px 20px; background: #f5f5f5; display: flex; justify-content: flex-end; gap: 10px;">
                <button onclick="downloadMinutesWord()" style="padding: 8px 16px; background: #165DFF; color: white; border: none; border-radius: 6px; cursor: pointer;">📥 下载Word</button>
                <button onclick="downloadMinutesPDF()" style="padding: 8px 16px; background: #67C23A; color: white; border: none; border-radius: 6px; cursor: pointer;">📄 下载PDF</button>
                <button onclick="closeMinutesModal()" style="padding: 8px 16px; background: #666; color: white; border: none; border-radius: 6px; cursor: pointer;">关闭</button>
            </div>
        </div>
    </div>
</body>
</html>
'''


# ==================== API路由 ====================
@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/dispatch')
def dispatch_board():
    """任务调度看板"""
    try:
        with open('/root/.openclaw/workspace/03_输出成果/呈彩产出/task_dispatch_board.html', 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return "<h1>调度看板页面未找到</h1><p>请确认文件已生成</p>", 404

@app.route('/test')
def test_page():
    return '''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>测试 - 九星智囊团</title>
    <style>
        body { font-family: sans-serif; margin: 20px; background: #f5f5f5; }
        .sidebar { width: 300px; background: white; border-radius: 8px; padding: 16px; }
        .sidebar-header { font-weight: bold; margin-bottom: 12px; color: #666; }
        .agent-card { display: flex; align-items: center; padding: 10px; margin-bottom: 6px; background: #f5f7fa; border-radius: 8px; border-left: 3px solid #165DFF; }
        .agent-avatar { width: 36px; height: 36px; border-radius: 8px; display: flex; align-items: center; justify-content: center; font-size: 16px; margin-right: 10px; color: white; }
        .agent-name { font-size: 13px; font-weight: 600; }
        .agent-role { font-size: 11px; color: #909399; }
        #status { margin-top: 20px; padding: 10px; background: #e8f3ff; border-radius: 4px; }
    </style>
</head>
<body>
    <h1>🧭 指南针工程 - 成员测试</h1>
    
    <div class="sidebar">
        <div class="sidebar-header">⭐ 九星智囊团</div>
        <div id="agentList" style="min-height:100px;">
            <div style="text-align:center;color:#909399;">加载中...</div>
        </div>
    </div>
    
    <div id="status"></div>

    <script>
        const AGENTS = {
            'fuyao': { name: '扶摇', role: '总指挥', emoji: '🌀', color: '#165DFF' },
            'nanqiao': { name: '南乔', role: '主控Agent', emoji: '🌿', color: '#9C27B0' },
            'yuheng': { name: '玉衡', role: '项目经理', emoji: '⚖️', color: '#F56C6C' },
            'caiwei': { name: '采薇', role: '需求分析专家', emoji: '🌸', color: '#409EFF' },
            'zhijin': { name: '织锦', role: '架构设计师', emoji: '🧵', color: '#67C23A' },
            'zhutai': { name: '筑台', role: '售前工程师', emoji: '🏗️', color: '#E6A23C' },
            'gongchi': { name: '工尺', role: '详细设计师', emoji: '📐', color: '#607D8B' },
            'chengcai': { name: '呈彩', role: '方案设计师', emoji: '🎨', color: '#FF9800' },
            'zhegui': { name: '折桂', role: '资源管家', emoji: '📚', color: '#00BCD4' }
        };

        function renderAgentList() {
            const container = document.getElementById('agentList');
            const countSpan = document.getElementById('agentCount');
            
            container.innerHTML = '';
            
            const order = ['fuyao', 'nanqiao', 'yuheng', 'caiwei', 'zhijin', 'zhutai', 'gongchi', 'chengcai', 'zhegui'];
            
            order.forEach(id => {
                const agent = AGENTS[id];
                if (!agent) return;
                
                const card = document.createElement('div');
                card.className = 'agent-card';
                card.style.borderLeftColor = agent.color;
                
                card.innerHTML = `
                    <div class="agent-avatar" style="background: ${agent.color}">${agent.emoji}</div>
                    <div>
                        <div class="agent-name">${agent.name}</div>
                        <div class="agent-role">${agent.role}</div>
                    </div>
                `;
                container.appendChild(card);
            });
            
            countSpan.textContent = order.length + '位专家';
            document.getElementById('status').innerHTML = '✅ 成功渲染 ' + order.length + ' 位专家';
        }

        // 页面加载后立即执行
        renderAgentList();
    </script>
</body>
</html>'''


@app.route('/api/status')
def api_status():
    return jsonify({'api_connected': bool(QIANFAN_API_KEY)})


@app.route('/api/conversation')
def api_conversation():
    return jsonify({
        'conversations': [
            {'turn_id': t.turn_id, 'speaker': t.speaker, 'speaker_name': t.speaker_name,
             'content': t.content, 'timestamp': t.timestamp, 'msg_type': t.msg_type,
             'is_challenging': t.is_challenging, 'reply_to': t.reply_to}
            for t in memory.history
        ],
        'agentStatus': agent_status,
        'isCompleted': discussion_completed,
        'tokensUsed': total_tokens_used,
        'tokensRemaining': max(0, TOKEN_BUDGET - total_tokens_used),
        'turnCount': len(memory.history)
    })


@app.route('/api/task', methods=['POST'])
def api_task():
    """
    任务处理接口 - V15深度融合版
    
    核心改进：
    1. 使用南乔意图分析器识别项目名称、产出物、复杂度
    2. 多产出物工作流（拓扑排序）
    3. 讨论结束后自动生成会议纪要
    """
    global agent_status, discussion_completed

    data = request.json
    task = data.get('task', '')

    # 重置状态
    discussion_completed = False
    memory.clear()
    agent_status = {aid: 'idle' for aid in AGENTS.keys()}

    # ========== V15深度融合：意图分析 ==========
    intent_result = None
    outputs = []
    project_name = "待定"
    complexity = "low"
    participants = []
    discussion_rounds = 1
    
    if load_v15_modules():
        try:
            intent_analyzer = NanqiaoIntentAnalyzer()
            intent_result = intent_analyzer.analyze(task)
            
            # 使用V15的分析结果
            project_name = intent_result.project_name
            outputs = intent_result.outputs
            complexity = intent_result.complexity
            participants = intent_result.participants
            discussion_rounds = intent_result.discussion_rounds
            
            print(f"[V15] 意图分析成功")
            print(f"[V15] 项目：{project_name}")
            print(f"[V15] 产出物：{[o['name'] for o in outputs]}")
            print(f"[V15] 复杂度：{complexity}")
            print(f"[V15] 参与者：{participants}")
            print(f"[V15] 讨论轮次：{discussion_rounds}")
        except Exception as e:
            print(f"[V15] 意图分析失败，回退到V14逻辑: {e}")
            import traceback
            traceback.print_exc()
    
    # 如果V15分析失败，使用V14逻辑
    if not outputs:
        schedule_results = intelligent_scheduler.process_multi(task)
        outputs = []
        for r in schedule_results:
            outputs.append({
                'name': r['output_template'],
                'code': r['task_code'],
                'owner': AGENT_NAMES.get(r['schedule']['lead_agent'], '采薇')
            })
        participants = schedule_results[0]['schedule']['participants'] if schedule_results else ['caiwei', 'yuheng']
        complexity = schedule_results[0].get('complexity', 'low')
        discussion_rounds = 1
    else:
        # V15识别出多个产出物 → 使用V14多任务识别逻辑
        # 这样可以保证每个产出物都有正确的task_code
        print(f"[V15] 识别到{len(outputs)}个产出物，使用V14多任务调度")
        schedule_results = intelligent_scheduler.process_multi(task)
        print(f"[V14] 识别到{len(schedule_results)}个任务")
        
        # 更新outputs为V14结果
        outputs = []
        for r in schedule_results:
            outputs.append({
                'name': r['output_template'],
                'code': r['task_code'],
                'owner': AGENT_NAMES.get(r['schedule']['lead_agent'], '采薇')
            })
        
        # 更新复杂度
        if len(schedule_results) > 1:
            complexity = 'medium' if len(schedule_results) <= 3 else 'high'


    print(f"[DEBUG] 用户输入: {task}")
    print(f"[DEBUG] 产出物数: {len(outputs)}")

    memory.current_task = task
    memory.start_time = datetime.now()
    memory.add_turn('user', '少帅', task)

    # 动态轮次上限
    num_tasks = len(schedule_results)
    if complexity == 'high':
        dynamic_turn_limit = min(20 + num_tasks * 3, 30)
    elif complexity == 'medium':
        dynamic_turn_limit = min(15 + num_tasks * 2, 25)
    else:
        dynamic_turn_limit = min(10 + num_tasks * 1, 20)

    # 调度提示
    if len(schedule_results) > 1:
        task_names = [r['task_name'] for r in schedule_results]
        schedule_msg = responder.generate_system_message('schedule', {
            'multi': True,
            'task_list': '、'.join(task_names),
            'tasks': task_names
        })
    else:
        participants_cn = [AGENT_NAMES.get(aid, aid) if aid in AGENTS else aid for aid in (participants if isinstance(participants[0], str) else [p['owner'] for p in participants])]
        schedule_msg = responder.generate_system_message('schedule', {
            'multi': False,
            'task_name': outputs[0]['name'] if outputs else '任务',
            'lead_agent': participants_cn[0] if participants_cn else '采薇',
            'participants': '、'.join(participants_cn),
            'est_time': schedule_results[0].get('estimated_time', '3-5天') if schedule_results else '3-5天',
            'complexity': complexity
        })
    memory.add_turn('nanqiao', '南乔', schedule_msg, msg_type='system')

    total_tasks = len(schedule_results)

    def run_multi_task_discussion():
        global discussion_completed, agent_status
        nonlocal outputs, project_name, complexity, discussion_rounds, participants
        
        # 初始化默认值
        all_participants = []
        
        try:
            for task_idx, schedule_result in enumerate(schedule_results):
                task_name = schedule_result['task_name']
                est_time = schedule_result.get('estimated_time', '3-5天')
                discussion_flow = schedule_result['schedule']['discussion_flow']
                task_complexity = schedule_result['complexity']
                lead_agent = schedule_result['schedule']['lead_agent']
                participants = schedule_result['schedule'].get('participants', discussion_flow)
                
                # 收集所有参与人员
                all_participants = list(set(all_participants + participants))
                
                if total_tasks > 1:
                    # 获取任务参与人员
                    participant_names = [AGENTS[p].name for p in participants if p in AGENTS]
                    lead_agent_obj = AGENTS.get(lead_agent)
                    lead_agent_name = lead_agent_obj.name if lead_agent_obj else '负责人'
                    
                    task_start_msg = f"📋 **任务{task_idx + 1}/{total_tasks}：{task_name}**\n\n"
                    task_start_msg += f"👤 负责人：{lead_agent_name}\n"
                    task_start_msg += f"👥 参与人员：{', '.join(participant_names)}\n"
                    task_start_msg += f"⏱️ 预估工期：{est_time}\n\n"
                    task_start_msg += f"正在启动任务讨论..."
                    
                    memory.add_turn('nanqiao', '南乔', task_start_msg, msg_type='system')
                    time.sleep(0.5)

                # ========== V15深度融合：使用MultiRoundDiscussion ==========
                if load_v15_modules() and MultiRoundDiscussion:
                    print(f"[V15] 使用MultiRoundDiscussion执行讨论")
                    
                    # 创建多轮讨论执行器
                    discussion_engine = MultiRoundDiscussion()
                    
                    # 构建结构化意图
                    structured_intent = {
                        'project_name': project_name,
                        'task_name': task_name,
                        'structured_intent': f"为「{project_name}」进行{task_name}",
                        'outputs': outputs,
                        'complexity': task_complexity
                    }
                    
                    # 执行多轮讨论
                    result = discussion_engine.run(
                        task=task,
                        structured_intent=structured_intent,
                        participants=discussion_flow,
                        max_rounds=discussion_rounds
                    )
                    
                    print(f"[V15] 讨论完成，共识: {result['consensus_reached']}")
                    
                    # 将讨论结果转换为memory记录
                    for round_result in result['rounds']:
                        for msg in round_result.messages:
                            agent_key = msg.speaker
                            if agent_key not in AGENTS:
                                continue
                            
                            agent_status[agent_key] = 'challenge' if msg.is_challenge else 'speaking'
                            memory.add_turn(
                                agent_key, 
                                msg.speaker_name, 
                                msg.content,
                                is_challenging=msg.is_challenge,
                                reply_to=msg.reply_to
                            )
                            time.sleep(0.3)  # 优化：1.5→0.3秒
                            agent_status[agent_key] = 'idle'
                    
                    # 如果有关键决策，添加到总结
                    if result['key_decisions']:
                        decisions_text = "\n".join([f"- {d}" for d in result['key_decisions']])
                        memory.add_turn('nanqiao', '南乔', 
                            f"📋 **关键决策**\n{decisions_text}", 
                            msg_type='system')
                
                else:
                    # 降级：使用V14的讨论逻辑
                    print("[V14] 使用传统讨论逻辑")
                    
                    # ========== V15.1 并行调度优化 ==========
                    if PARALLEL_MODE_ENABLED and round_num == 0:
                        # 第一轮使用并行调度
                        print("[V15.1] 并行调度模式启动")
                        context = memory.get_context()
                        
                        # 并行调用所有Agent
                        parallel_results = dispatch_agents_parallel(
                            discussion_flow, task, context
                        )
                        
                        # 处理并行结果
                        for agent_key in discussion_flow:
                            if agent_key not in AGENTS:
                                continue
                            
                            if agent_key in parallel_results:
                                response, from_cache = parallel_results[agent_key]
                                is_challenge = any(kw in response for kw in 
                                    ['质疑', '反对', '不同意', '不可行', '成本过高', '挑战'])
                                
                                # 添加工期评估（仅主Agent）
                                if agent_key == lead_agent:
                                    est_msg = responder.generate_estimate_message(
                                        AGENTS[agent_key].name, task_name, est_time, task_complexity
                                    )
                                    response += f"\n\n{est_msg}"
                                
                                agent_status[agent_key] = 'challenge' if is_challenge else 'speaking'
                                memory.add_turn(agent_key, AGENTS[agent_key].name, response,
                                              is_challenging=is_challenge)
                                agent_status[agent_key] = 'idle'
                            else:
                                # 降级：使用fallback响应
                                if PARALLEL_FALLBACK_ENABLED:
                                    response, is_challenge, reply_to = responder._fallback(
                                        AGENTS[agent_key], task, context
                                    )
                                    memory.add_turn(agent_key, AGENTS[agent_key].name, response,
                                                  is_challenging=is_challenge)
                        
                        print(f"[V15.1] 并行调度完成，成功: {len(parallel_results)}/{len(discussion_flow)}")
                    
                    # ========== 原有串行逻辑（后续轮次或降级时使用）==========
                    for round_num in range(1 if PARALLEL_MODE_ENABLED else 0, discussion_rounds):
                        
                        for agent_key in discussion_flow:
                            if len(memory.history) >= dynamic_turn_limit:
                                break
                            if agent_key not in AGENTS:
                                continue

                            time.sleep(0.5)
                            agent_status[agent_key] = 'speaking'
                            
                            context_prefix = ""
                            if round_num > 0:
                                context_prefix = f"【第{round_num + 1}轮讨论】请深化分析或提出质疑。\n\n"
                            
                            # 质疑回应提示（如果有针对该Agent的质疑）
                            if CHALLENGE_ENHANCEMENT_ENABLED:
                                try:
                                    cm = get_challenge_manager()
                                    response_prompt = cm.get_response_prompt(agent_key)
                                    if response_prompt:
                                        context_prefix += response_prompt + "\n\n"
                                        print(f"[质疑增强] {AGENTS[agent_key].name}有待回应质疑")
                                except Exception as e:
                                    print(f"[质疑增强] 获取回应提示失败: {e}")
                            
                            try:
                                response, is_challenge, reply_to = responder.generate(
                                    AGENTS[agent_key], context_prefix + task, memory
                                )
                            except Exception as e:
                                response, is_challenge, reply_to = responder._fallback(AGENTS[agent_key], task, memory.get_context())

                            # 质疑增强处理
                            if is_challenge and CHALLENGE_ENHANCEMENT_ENABLED:
                                try:
                                    cm = get_challenge_manager()
                                    # 转换历史格式
                                    history_dicts = [{
                                        'speaker': t.speaker,
                                        'speaker_name': t.speaker_name,
                                        'content': t.content
                                    } for t in memory.history]
                                    # 检测质疑详情
                                    challenge = cm.process_message(
                                        response, agent_key, AGENTS[agent_key].name,
                                        history_dicts, AGENTS, round_num + 1
                                    )
                                    if challenge:
                                        print(f"[质疑增强] 检测到{challenge.intensity.name}级质疑: {challenge.category.value}")
                                        # 如果质疑者接受了（说"调整"、"接受"等），更新状态
                                        if '接受' in response or '调整' in response or '同意' in response:
                                            cm.chain.add_response(
                                                challenge.challenge_id, 
                                                challenge.target_agent,
                                                challenge.target_agent_name,
                                                response,
                                                accepts=True
                                            )
                                except Exception as e:
                                    print(f"[质疑增强] 处理失败: {e}")

                            if agent_key == lead_agent and round_num == 0:
                                est_msg = responder.generate_estimate_message(
                                    AGENTS[agent_key].name, task_name, est_time, task_complexity
                                )
                                response += f"\n\n{est_msg}"

                            agent_status[agent_key] = 'challenge' if is_challenge else 'speaking'
                            memory.add_turn(agent_key, AGENTS[agent_key].name, response,
                                          is_challenging=is_challenge, reply_to=reply_to)
                            time.sleep(0.1)  # 优化：2→0.1秒
                            agent_status[agent_key] = 'idle'
                        
                        if len(memory.history) >= dynamic_turn_limit:
                            break

                # 风险辩论
                if task_complexity == 'high' and len(memory.history) < dynamic_turn_limit:
                    risk_msg = responder.generate_system_message('risk_debate', {})
                    memory.add_turn('nanqiao', '南乔', risk_msg, msg_type='system')

                    for agent_key in AGENTS.keys():
                        if len(memory.history) >= dynamic_turn_limit:
                            break
                        if agent_key == 'nanqiao' or agent_key in discussion_flow:
                            continue
                        time.sleep(0.3)
                        agent_status[agent_key] = 'speaking'
                        try:
                            response, is_challenge, reply_to = responder.generate(AGENTS[agent_key], task, memory)
                        except Exception as e:
                            response, is_challenge, reply_to = responder._fallback(AGENTS[agent_key], task, memory.get_context())
                        memory.add_turn(agent_key, AGENTS[agent_key].name, response,
                                      is_challenging=True, reply_to=reply_to)
                        time.sleep(0.1)  # 优化：1.5→0.1秒
                        agent_status[agent_key] = 'idle'

                    # 收集当前任务的讨论内容
                    task_messages = [t for t in memory.history if t.msg_type != 'system']
                    task_summary = f"✅ **任务{task_idx + 1}完成：{task_name}**\n\n"
                    task_summary += f"📊 共识度：{memory.get_consensus_level()}%\n"
                    task_summary += f"💬 讨论轮次：{len([t for t in task_messages if t.speaker != 'user' and t.speaker != 'nanqiao'])}条发言\n"
                    
                    # 获取关键质疑
                    challenges = [t for t in memory.history if t.is_challenging]
                    if challenges:
                        task_summary += f"⚠️ 关键质疑：{len(challenges)}条\n"
                    
                    # 只有还有下一个任务时才显示提示
                    if task_idx < total_tasks - 1:
                        task_summary += f"\n➡️ 即将开始下一个任务..."
                    
                    memory.add_turn('nanqiao', '南乔', task_summary, msg_type='system')

            # 最终总结 - 基于实际讨论内容动态生成
            time.sleep(0.5)
            agent_status['nanqiao'] = 'speaking'
            consensus = memory.get_consensus_level()

            # ========== 提取讨论摘要 ==========
            discussion_summary = ""
            key_decisions = []
            key_challenges = []
            
            # 提取关键发言（每个Agent取最后一条主要发言）
            agent_key_points = {}
            for turn in memory.history:
                if turn.speaker == 'user' or turn.speaker == 'nanqiao':
                    continue
                if turn.speaker not in agent_key_points:
                    agent_key_points[turn.speaker] = {
                        'name': turn.speaker_name,
                        'content': turn.content[:200]
                    }
            
            # 构建讨论摘要
            if agent_key_points:
                discussion_summary = "\n".join([
                    f"- {v['name']}：{v['content'][:100]}..."
                    for v in agent_key_points.values()
                ])
            
            # 提取关键决策（包含"决定"、"确认"、"同意"等关键词）
            decision_keywords = ['决定', '确认', '同意', '采用', '选择', '确定']
            for turn in memory.history:
                if turn.speaker == 'user' or turn.speaker == 'nanqiao':
                    continue
                for kw in decision_keywords:
                    if kw in turn.content:
                        key_decisions.append({
                            'agent': turn.speaker_name,
                            'decision': turn.content[:100]
                        })
                        break
            
            # 获取质疑链路总结
            challenge_summary = ""
            pending_challenges = []
            if CHALLENGE_ENHANCEMENT_ENABLED:
                try:
                    cm = get_challenge_manager()
                    challenge_summary = cm.get_challenge_summary()
                    pending_challenges = [c for c in cm.get_challenges() if c['status'] == 'pending']
                except Exception as e:
                    print(f"[质疑增强] 获取质疑总结失败: {e}")
            
            # 提取质疑
            for turn in memory.history:
                if turn.is_challenging:
                    key_challenges.append({
                        'challenger': turn.speaker_name,
                        'content': turn.content[:80]
                    })

            # ========== 动态生成总结（精简版） ==========
            print(f"[DEBUG] 开始生成总结，schedule_results长度: {len(schedule_results)}")
            outputs = [r.get('output_template', '产出物') for r in schedule_results]
            print(f"[DEBUG] outputs: {outputs}")
            
            # 计算讨论时长
            discussion_duration = ""
            if memory.start_time:
                duration_seconds = (datetime.now() - memory.start_time).total_seconds()
                minutes = int(duration_seconds // 60)
                seconds = int(duration_seconds % 60)
                discussion_duration = f"{minutes}分{seconds}秒" if minutes > 0 else f"{seconds}秒"
            
            # ========== 构建完整总结 ==========
            print("[DEBUG] 开始构建总结内容...")
            summary = ""
            
            # 第1部分：项目概况
            summary += f"📋 项目概况\n\n"
            summary += f"项目名称：{project_name}\n"
            summary += f"任务数量：{total_tasks}个\n"
            summary += f"参与人员：{len(set([t.speaker_name for t in memory.history if t.speaker not in ['user', 'nanqiao']]))}位专家\n"
            summary += f"讨论时长：{discussion_duration}\n"
            summary += f"整体共识度：{consensus}%\n\n"
            
            # 第2部分：关键决策
            summary += "📌 关键决策\n\n"
            if key_decisions:
                for i, d in enumerate(key_decisions[:5], 1):
                    # 不截断，显示完整决策内容
                    summary += f"{i}. 【{d['agent']}】{d['decision']}\n\n"
            else:
                summary += "暂无明确决策记录。\n\n"
            
            # 第3部分：质疑链路
            summary += "⚠️ 质疑链路\n\n"
            if challenge_summary and len(challenge_summary) > 20:
                # 不截断，显示完整质疑总结
                summary += challenge_summary + "\n\n"
            elif key_challenges:
                summary += "主要质疑：\n"
                for i, c in enumerate(key_challenges[:3], 1):
                    # 不截断，显示完整质疑内容
                    summary += f"{i}. 【{c['challenger']}】{c['content']}\n\n"
            else:
                summary += "本次讨论无重大质疑。\n\n"
            pending_cnt = len(pending_challenges) if isinstance(pending_challenges, list) else 0
            if pending_cnt > 0:
                summary += f"⏳ 待回应质疑：{pending_cnt}条\n\n"
            
            # 第4部分：下一步行动（明确责任人、产出物、时间）
            summary += "📝 下一步行动\n\n"
            
            # 根据任务生成具体的行动项
            action_idx = 1
            for idx, result in enumerate(schedule_results):
                task_name = result['task_name']
                lead_agent = result['schedule'].get('lead_agent', 'caiwei')
                lead_agent_obj = AGENTS.get(lead_agent)
                lead_name = lead_agent_obj.name if lead_agent_obj else '负责人'
                output = result.get('output_template', '产出物')
                est_time = result.get('estimated_time', '3-5天')
                
                summary += f"{action_idx}. {lead_name} 负责「{output}」\n"
                summary += f"   - 产出物：{output}\n"
                summary += f"   - 完成时间：{est_time}内\n\n"
                action_idx += 1
            
            # 添加通用行动项
            pending_count = len(pending_challenges) if isinstance(pending_challenges, list) else 0
            if pending_count > 0:
                summary += f"{action_idx}. 相关责任人 回应待处理质疑\n"
                summary += f"   - 待回应质疑：{pending_count}条\n"
                summary += f"   - 完成时间：下轮讨论前\n\n"
                action_idx += 1
            
            summary += "\n🎉 论证完成！\n"

            print(f"[DEBUG] 总结构建完成，长度: {len(summary)}")
            print(f"[DEBUG] 总结内容预览: {summary[:200]}...")
            
            memory.add_turn('nanqiao', '南乔', summary, msg_type='conclusion')
            agent_status['nanqiao'] = 'idle'
            print("[DEBUG] 总结已添加到memory")

        except Exception as e:
            print(f"[ERROR] 讨论线程异常: {e}")
            import traceback
            traceback.print_exc()
            # 发送错误消息到前端
            error_msg = f"❌ 讨论过程出现异常：{str(e)}"
            memory.add_turn('nanqiao', '南乔', error_msg, msg_type='system')
        finally:
            discussion_completed = True
            print(f"[DEBUG] 讨论完成，共识度: {memory.get_consensus_level()}%")
            
            # ========== V15自动生成会议纪要 ==========
            if load_v15_modules() and intent_result and MeetingMinutesGenerator:
                try:
                    print("[V15] 正在生成会议纪要...")
                    minutes_gen = MeetingMinutesGenerator()
                    
                    # 构建完整的讨论数据（包含轮次信息）
                    discussion_data = {
                        'rounds': [],
                        'consensus_reached': discussion_completed,
                        'key_decisions': [],
                        'project_name': project_name,
                        'complexity': complexity,
                        'participants': all_participants
                    }
                    
                    # 按轮次组织消息
                    current_round = 1
                    round_messages = []
                    
                    for turn in memory.history:
                        if turn.speaker == 'user' or turn.speaker == 'nanqiao':
                            continue
                        
                        round_messages.append({
                            'round_num': current_round,
                            'speaker': turn.speaker,
                            'speaker_name': turn.speaker_name,
                            'content': turn.content[:500]
                        })
                        
                        # 每6条消息为一轮（简单估算）
                        if len(round_messages) % 6 == 0:
                            discussion_data['rounds'].append({
                                'round_num': current_round,
                                'messages': round_messages[-6:]
                            })
                            current_round += 1
                    
                    # 添加最后一轮（如果有剩余）
                    if len(round_messages) % 6 != 0:
                        discussion_data['rounds'].append({
                            'round_num': current_round,
                            'messages': round_messages[-(len(round_messages) % 6):]
                        })
                    
                    # 生成会议纪要
                    minutes = minutes_gen.generate(intent_result.__dict__, discussion_data)
                    minutes_file = minutes_gen.export_to_file(minutes)
                    print(f"[V15] 会议纪要已生成: {minutes_file}")
                    
                    # 添加到对话
                    memory.add_turn('nanqiao', '南乔', 
                        f"\n\n📄 **会议纪要已自动生成**\n文件路径：{minutes_file}\n\n可点击上方「导出」按钮下载完整纪要文档。", 
                        msg_type='system')
                except Exception as e:
                    print(f"[V15] 会议纪要生成失败: {e}")
                    import traceback
                    traceback.print_exc()

    thread = threading.Thread(target=run_multi_task_discussion)
    thread.start()

    return jsonify({
        'status': 'ok',
        'v15_mode': intent_result is not None,
        'project_name': project_name,
        'outputs': outputs,
        'complexity': complexity,
        'discussion_rounds': discussion_rounds,
        'schedule': schedule_results[0] if len(schedule_results) == 1 else {'multi': len(schedule_results)},
        'tasks': [r['task_name'] for r in schedule_results],
        'estimated_times': [r.get('estimated_time', '3-5天') for r in schedule_results]
    })


@app.route('/api/clear', methods=['POST'])
def api_clear():
    global agent_status, discussion_completed
    memory.clear()
    agent_status = {aid: 'idle' for aid in AGENTS.keys()}
    discussion_completed = False
    return jsonify({'status': 'ok'})


@app.route('/api/export/word', methods=['POST'])
def api_export_word():
    """导出Word文档"""
    try:
        # 转换对话历史
        turns = []
        for t in memory.history:
            turn = DiscussionTurn(
                turn_id=t.turn_id,
                speaker=t.speaker,
                speaker_name=t.speaker_name,
                speaker_role=AGENTS.get(t.speaker, AGENTS['nanqiao']).role if t.speaker in AGENTS else '用户',
                content=t.content,
                timestamp=t.timestamp,
                is_challenging=t.is_challenging,
                reply_to=t.reply_to
            )
            turns.append(turn)

        # 创建摘要
        participants = list(set([t.speaker_name for t in memory.history if t.speaker != 'user']))
        summary = DiscussionSummary(
            task=memory.current_task or '未命名任务',
            start_time=memory.start_time.isoformat() if memory.start_time else '',
            end_time=datetime.now().isoformat(),
            total_turns=memory.turn_count,
            consensus_level=memory.get_consensus_level(),
            participants=participants,
            key_points=[],  # 可以从讨论中提取
            risks=[],
            decisions=[]
        )

        # 导出
        word_path = export_api.export_word(turns, summary)
        return send_file(word_path, as_attachment=True, download_name=f'讨论报告_{datetime.now().strftime("%Y%m%d")}.docx')
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/export/excel', methods=['POST'])
def api_export_excel():
    """导出Excel表格"""
    try:
        # 转换对话历史
        turns = []
        for t in memory.history:
            turn = DiscussionTurn(
                turn_id=t.turn_id,
                speaker=t.speaker,
                speaker_name=t.speaker_name,
                speaker_role=AGENTS.get(t.speaker, AGENTS['nanqiao']).role if t.speaker in AGENTS else '用户',
                content=t.content,
                timestamp=t.timestamp,
                is_challenging=t.is_challenging,
                reply_to=t.reply_to
            )
            turns.append(turn)

        # 创建摘要
        participants = list(set([t.speaker_name for t in memory.history if t.speaker != 'user']))
        summary = DiscussionSummary(
            task=memory.current_task or '未命名任务',
            start_time=memory.start_time.isoformat() if memory.start_time else '',
            end_time=datetime.now().isoformat(),
            total_turns=memory.turn_count,
            consensus_level=memory.get_consensus_level(),
            participants=participants,
            key_points=[],
            risks=[],
            decisions=[]
        )

        # 导出
        excel_path = export_api.export_excel(turns, summary)
        return send_file(excel_path, as_attachment=True, download_name=f'讨论报告_{datetime.now().strftime("%Y%m%d")}.xlsx')
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/export/markdown', methods=['POST'])
def api_export_markdown():
    """导出Markdown文档"""
    try:
        # 转换对话历史
        turns = []
        for t in memory.history:
            turn = DiscussionTurn(
                turn_id=t.turn_id,
                speaker=t.speaker,
                speaker_name=t.speaker_name,
                speaker_role=AGENTS.get(t.speaker, AGENTS['nanqiao']).role if t.speaker in AGENTS else '用户',
                content=t.content,
                timestamp=t.timestamp,
                is_challenging=t.is_challenging,
                reply_to=t.reply_to
            )
            turns.append(turn)

        # 创建摘要
        participants = list(set([t.speaker_name for t in memory.history if t.speaker != 'user']))
        summary = DiscussionSummary(
            task=memory.current_task or '未命名任务',
            start_time=memory.start_time.isoformat() if memory.start_time else '',
            end_time=datetime.now().isoformat(),
            total_turns=memory.turn_count,
            consensus_level=memory.get_consensus_level(),
            participants=participants,
            key_points=[],
            risks=[],
            decisions=[]
        )

        # 导出
        md_content = export_api.export_markdown(turns, summary)
        return jsonify({'status': 'ok', 'content': md_content})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


# ==================== 文档解析API ====================
@app.route('/api/parse/document', methods=['POST'])
def api_parse_document():
    """解析上传的文档"""
    try:
        if 'document' not in request.files:
            return jsonify({'status': 'error', 'message': '没有上传文件'})

        file = request.files['document']
        if file.filename == '':
            return jsonify({'status': 'error', 'message': '文件名为空'})

        # 保存临时文件
        import tempfile
        import os
        import re
        from pathlib import Path

        # 获取文件扩展名
        filename = file.filename
        ext = os.path.splitext(filename)[1].lower()

        # 支持的格式
        if ext not in ['.pdf', '.doc', '.docx', '.txt', '.png', '.jpg', '.jpeg']:
            return jsonify({'status': 'error', 'message': f'不支持的文件格式: {ext}'})

        # 保存到临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        try:
            # 解析文档
            if ext == '.txt':
                # 直接读取文本文件
                with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()

                # 提取标题
                lines = content.strip().split('\n')
                title = lines[0][:100] if lines else filename

                # 统计字数
                chinese = len(re.findall(r'[\u4e00-\u9fff]', content))
                english = len(re.findall(r'\b[a-zA-Z]+\b', content))
                word_count = chinese + english

                # 分割章节
                sections = []
                current_section = None
                current_content = []

                for line in lines:
                    title_match = re.match(r'^(\d+\.?\s+.+)|(^#{1,6}\s+.+)', line.strip())
                    if title_match:
                        if current_section:
                            current_section['content'] = '\n'.join(current_content).strip()
                            sections.append(current_section)

                        t = title_match.group(0).lstrip('#').strip()
                        level = len(re.match(r'^#+', line).group()) if line.startswith('#') else 2
                        current_section = {'level': level, 'title': t, 'content': '', 'position': len(sections)}
                        current_content = []
                    else:
                        current_content.append(line)

                if current_section:
                    current_section['content'] = '\n'.join(current_content).strip()
                    sections.append(current_section)

                return jsonify({
                    'status': 'ok',
                    'file_path': tmp_path,
                    'file_type': 'text',
                    'title': title,
                    'content': content,
                    'sections': sections,
                    'tables': [],
                    'word_count': word_count,
                    'parse_time': datetime.now().isoformat()
                })
            else:
                # 使用文档解析器
                from document_parser import DocumentParserAPI
                parser = DocumentParserAPI()
                result = parser.parse(tmp_path)

                return jsonify({
                    'status': 'ok',
                    'file_path': result.file_path,
                    'file_type': result.file_type,
                    'title': result.title,
                    'content': result.content,
                    'sections': result.sections,
                    'tables': result.tables,
                    'word_count': result.word_count,
                    'parse_time': result.parse_time
                })
        finally:
            # 清理临时文件
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)})


# ==================== 知识库API ====================
@app.route('/api/knowledge/list', methods=['GET'])
def api_knowledge_list():
    """获取知识库文件列表"""
    try:
        files = knowledge_base.list_available_knowledge()
        return jsonify({
            'status': 'ok',
            'count': len(files),
            'files': files
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/knowledge/load/<path:relative_path>', methods=['GET'])
def api_knowledge_load(relative_path):
    """加载指定知识文件"""
    try:
        content = knowledge_base.load_file(relative_path)
        if content:
            return jsonify({
                'status': 'ok',
                'path': relative_path,
                'content': content,
                'length': len(content)
            })
        else:
            return jsonify({'status': 'error', 'message': '文件不存在或为空'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/knowledge/task/<task_type>', methods=['GET'])
def api_knowledge_for_task(task_type):
    """获取任务相关知识"""
    try:
        knowledge = knowledge_base.get_knowledge_for_task(task_type)
        return jsonify({
            'status': 'ok',
            'task_type': task_type,
            'knowledge': knowledge
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


# ==================== 场景模板API ====================
@app.route('/api/scenes/list', methods=['GET'])
def api_scenes_list():
    """获取场景模板列表"""
    try:
        templates = scene_api.list_templates()
        return jsonify({
            'status': 'ok',
            'count': len(templates),
            'templates': templates
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/scenes/<scene_id>', methods=['GET'])
def api_scene_get(scene_id):
    """获取指定场景模板"""
    try:
        template = scene_api.get_template(scene_id)
        if template:
            return jsonify({
                'status': 'ok',
                'template': scene_api.to_dict(template)
            })
        else:
            return jsonify({'status': 'error', 'message': '模板不存在'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/scenes/fill', methods=['POST'])
def api_scene_fill():
    """填充场景模板参数"""
    try:
        data = request.get_json()
        scene_id = data.get('scene_id')
        params = data.get('params', {})

        result = scene_api.fill_template(scene_id, params)

        if 'error' in result:
            return jsonify({'status': 'error', 'message': result['error']})

        return jsonify({
            'status': 'ok',
            'result': result
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


# ==================== 成本估算API ====================
@app.route('/api/cost/estimate', methods=['POST'])
def api_cost_estimate():
    """成本估算"""
    try:
        data = request.get_json()

        estimation = cost_estimator.estimate(
            project_name=data.get('project_name', '未命名项目'),
            team_size=data.get('team_size', 5),
            duration_months=data.get('duration_months', 6),
            complexity=data.get('complexity', '中等'),
            tech_stack=data.get('tech_stack', ['Java', 'Vue', 'MySQL']),
            cloud_service=data.get('cloud_service', True),
            cloud_tier=data.get('cloud_tier', '中型')
        )

        return jsonify({
            'status': 'ok',
            'estimation': cost_estimator.to_dict(estimation),
            'report': cost_estimator.generate_report(estimation)
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


# ==================== 流式输出API ====================
@app.route('/api/stream', methods=['POST'])
def api_stream():
    """流式输出API - 实时逐字显示"""
    # 在生成器外部获取请求数据，避免request context问题
    data = request.get_json()
    system_prompt = data.get('system_prompt', '')
    user_message = data.get('user_message', '')
    temperature = data.get('temperature', 0.7)

    def generate():
        try:
            for chunk in call_qianfan_stream(system_prompt, user_message, temperature):
                if chunk.startswith('error:'):
                    yield f"data: {json.dumps({'error': chunk[6:]})}\n\n"
                    break
                else:
                    yield f"data: {json.dumps({'text': chunk})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

        yield f"data: {json.dumps({'done': True})}\n\n"

    return Response(generate(), mimetype='text/event-stream')


@app.route('/api/stream/test', methods=['GET'])
def api_stream_test():
    """流式输出测试"""
    def generate():
        test_text = "这是一个流式输出测试。文字将逐字显示，让用户实时看到内容，而不是等待全部生成完毕。这种体验更加流畅自然。"
        for char in test_text:
            yield f"data: {json.dumps({'text': char})}\n\n"
            time.sleep(0.05)
        yield f"data: {json.dumps({'done': True})}\n\n"

    return Response(generate(), mimetype='text/event-stream')


# ==================== V15增强API ====================
@app.route('/api/v15/analyze', methods=['POST'])
def api_v15_analyze():
    """V15意图分析API"""
    if not load_v15_modules():
        return jsonify({'error': 'V15模块未加载'}), 500
    
    try:
        data = request.json
        task = data.get('task', '')
        
        analyzer = NanqiaoIntentAnalyzer()
        intent = analyzer.analyze(task)
        
        return jsonify({
            'success': True,
            'project_name': intent.project_name,
            'task_type': intent.task_type,
            'task_name': intent.task_name,
            'outputs': intent.outputs,
            'complexity': intent.complexity,
            'participants': [AGENT_NAMES.get(p, p) for p in intent.participants],
            'discussion_rounds': intent.discussion_rounds,
            'structured_intent': intent.structured_intent
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v15/minutes', methods=['POST'])
def api_v15_minutes():
    """V15会议纪要API"""
    if not load_v15_modules():
        return jsonify({'error': 'V15模块未加载'}), 500
    
    try:
        data = request.json
        intent_data = data.get('intent', {})
        discussion_data = data.get('discussion', {})
        
        gen = MeetingMinutesGenerator()
        minutes = gen.generate(intent_data, discussion_data)
        filepath = gen.export_to_file(minutes)
        
        return jsonify({
            'success': True,
            'minutes': minutes,
            'file': filepath
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ========== 质疑增强API ==========
@app.route('/api/challenges/list', methods=['GET'])
def api_challenges_list():
    """获取所有质疑"""
    try:
        manager = get_challenge_manager()
        challenges = manager.get_challenges()
        consensus = manager.get_consensus()
        summary = manager.get_challenge_summary()
        
        return jsonify({
            'success': True,
            'enabled': CHALLENGE_ENHANCEMENT_ENABLED,
            'challenges': challenges,
            'consensus': consensus,
            'summary': summary
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/challenges/consensus', methods=['GET'])
def api_challenges_consensus():
    """获取共识度详情"""
    try:
        manager = get_challenge_manager()
        consensus = manager.get_consensus()
        
        return jsonify({
            'success': True,
            'consensus': consensus
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/challenges/reset', methods=['POST'])
def api_challenges_reset():
    """重置质疑状态"""
    try:
        manager = get_challenge_manager()
        manager.reset()
        
        return jsonify({
            'success': True,
            'message': '质疑状态已重置'
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ========== 会议纪要预览与下载API ==========
import glob

@app.route('/api/minutes/preview', methods=['GET'])
def api_minutes_preview():
    """预览会议纪要"""
    try:
        # 查找最新的会议纪要文件
        minutes_dir = '/root/.openclaw/workspace/03_输出成果/会议纪要'
        if not os.path.exists(minutes_dir):
            return jsonify({'status': 'error', 'message': '暂无会议纪要，请先进行讨论'})
        
        md_files = glob.glob(os.path.join(minutes_dir, '*.md'))
        if not md_files:
            return jsonify({'status': 'error', 'message': '暂无会议纪要，请先进行讨论'})
        
        # 获取最新的文件
        latest_file = max(md_files, key=os.path.getmtime)
        
        with open(latest_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return jsonify({
            'status': 'ok',
            'content': content,
            'file': latest_file,
            'mtime': datetime.fromtimestamp(os.path.getmtime(latest_file)).isoformat()
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


@app.route('/api/minutes/download', methods=['GET'])
def api_minutes_download():
    """下载会议纪要"""
    try:
        format_type = request.args.get('format', 'md')
        
        # 查找最新的会议纪要文件
        minutes_dir = '/root/.openclaw/workspace/03_输出成果/会议纪要'
        if not os.path.exists(minutes_dir):
            return jsonify({'status': 'error', 'message': '暂无会议纪要'}), 404
        
        md_files = glob.glob(os.path.join(minutes_dir, '*.md'))
        if not md_files:
            return jsonify({'status': 'error', 'message': '暂无会议纪要'}), 404
        
        # 获取最新的文件
        latest_file = max(md_files, key=os.path.getmtime)
        
        with open(latest_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        if format_type == 'md':
            # 返回Markdown
            return Response(
                content,
                mimetype='text/markdown',
                headers={'Content-Disposition': f'attachment; filename=会议纪要_{datetime.now().strftime("%Y%m%d")}.md'}
            )
        
        elif format_type == 'word':
            # 转换为Word
            try:
                from docx import Document
                from docx.shared import Pt, Inches, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                return jsonify({'status': 'error', 'message': 'Word导出模块未安装'}), 500
            
            doc = Document()
            
            # 解析Markdown并转换为Word
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    # 标题1
                    p = doc.add_heading(line[2:], level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('## '):
                    # 标题2
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # 标题3
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('|'):
                    # 表格（简化处理）
                    cells = [c.strip() for c in line.split('|') if c.strip()]
                    if cells and not all(c.replace('-', '').strip() == '' for c in cells):
                        # 创建表格行
                        if not hasattr(api_minutes_download, 'table_cells'):
                            api_minutes_download.table_cells = []
                        api_minutes_download.table_cells.append(cells)
                elif line.startswith('- '):
                    # 列表
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip() == '---':
                    # 分隔线
                    doc.add_paragraph('─' * 50)
                elif line.strip():
                    # 普通段落
                    doc.add_paragraph(line)
            
            # 处理表格
            if hasattr(api_minutes_download, 'table_cells') and api_minutes_download.table_cells:
                table = doc.add_table(rows=len(api_minutes_download.table_cells), cols=len(api_minutes_download.table_cells[0]))
                table.style = 'Table Grid'
                for i, row_cells in enumerate(api_minutes_download.table_cells):
                    for j, cell_text in enumerate(row_cells):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell_text
                api_minutes_download.table_cells = []
            
            # 保存到临时文件
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
            
            return send_file(tmp_path, as_attachment=True, download_name=f'会议纪要_{datetime.now().strftime("%Y%m%d")}.docx')
        
        elif format_type == 'pdf':
            # PDF导出（需要安装pdfkit或reportlab）
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.units import inch
                import tempfile
                
                # 创建PDF
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                    tmp_path = tmp.name
                
                doc = SimpleDocTemplate(tmp_path, pagesize=A4)
                styles = getSampleStyleSheet()
                story = []
                
                # 简化处理Markdown
                for line in content.split('\n'):
                    if line.startswith('# '):
                        story.append(Paragraph(line[2:], styles['Title']))
                    elif line.startswith('## '):
                        story.append(Paragraph(line[3:], styles['Heading2']))
                    elif line.startswith('### '):
                        story.append(Paragraph(line[4:], styles['Heading3']))
                    elif line.startswith('- '):
                        story.append(Paragraph('• ' + line[2:], styles['Normal']))
                    elif line.strip():
                        story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 0.1 * inch))
                
                doc.build(story)
                
                return send_file(tmp_path, as_attachment=True, download_name=f'会议纪要_{datetime.now().strftime("%Y%m%d")}.pdf')
            except ImportError:
                return jsonify({'status': 'error', 'message': 'PDF导出模块未安装，请安装reportlab'}), 500
        
        else:
            return jsonify({'status': 'error', 'message': '不支持的格式'}), 400
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'status': 'error', 'message': str(e)}), 500


if __name__ == '__main__':
    print("=" * 60)
    print("🧭 指南针工程 - 智能协作平台 V14 + V15.2产物生成")
    print("=" * 60)
    print("新增功能：文档上传 | 知识库集成 | 场景模板 | 成本估算")
    print("流式输出：实时逐字显示 | 打字机效果 | 历史保存")
    print("知识库：43个知识文件（电信行业专家级）")
    print("场景模板：6个预设场景（一键启动）")
    print("成本估算：参数化模型（人天+云服务+风险储备）")
    print("任务类型：20种细粒度任务（全生命周期覆盖）")
    print("")
    print("🌿 V15增强：意图分析 | 会议纪要自动生成")
    print("🚀 V15.2新增：产物生成 | 质量审核 | 模板管理")
    print("   - 讨论结束后点击「生成产物」按钮")
    print("   - 自动生成可交付产物文档")
    print("   - 支持质量检查和下载")
    print("")
    print("访问地址：http://120.48.169.242:5001")
    print("=" * 60)
    app.run(host='0.0.0.0', port=5001, debug=False)
