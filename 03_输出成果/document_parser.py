#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档解析模块
- PDF文档解析
- Word文档解析
- 图片OCR解析
- 多格式统一输出

Author: 南乔
Date: 2026-03-14
"""

import os
import re
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import json

# PDF解析
try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# Word解析
try:
    from docx import Document
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False

# OCR
try:
    import pytesseract
    from PIL import Image
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False


# ==================== 解析结果数据结构 ====================
@dataclass
class ParsedDocument:
    """解析后的文档"""
    file_path: str
    file_type: str
    title: str
    content: str
    sections: List[Dict]  # 章节列表
    tables: List[Dict]    # 表格列表
    images: List[str]     # 图片路径
    metadata: Dict        # 元数据
    parse_time: str
    word_count: int


@dataclass
class DocumentSection:
    """文档章节"""
    level: int       # 标题级别 1-6
    title: str       # 章节标题
    content: str     # 章节内容
    position: int    # 位置索引


# ==================== 文档解析器基类 ====================
class DocumentParser:
    """文档解析器基类"""
    
    def __init__(self):
        self.supported_formats = []
    
    def parse(self, file_path: str) -> ParsedDocument:
        raise NotImplementedError
    
    def _extract_title(self, content: str) -> str:
        """从内容中提取标题"""
        lines = content.strip().split('\n')
        for line in lines[:5]:  # 只看前5行
            line = line.strip()
            if line and len(line) < 100:  # 标题通常较短
                return line
        return "未命名文档"
    
    def _count_words(self, content: str) -> int:
        """统计字数"""
        # 中文字符
        chinese = len(re.findall(r'[\u4e00-\u9fff]', content))
        # 英文单词
        english = len(re.findall(r'\b[a-zA-Z]+\b', content))
        return chinese + english
    
    def _split_sections(self, content: str) -> List[Dict]:
        """分割章节"""
        sections = []
        lines = content.split('\n')
        
        current_section = None
        current_content = []
        
        for line in lines:
            # 检测标题行（数字开头或#开头）
            title_match = re.match(r'^(\d+\.?\s+.+)|(^#{1,6}\s+.+)', line.strip())
            
            if title_match:
                # 保存上一个章节
                if current_section:
                    current_section['content'] = '\n'.join(current_content).strip()
                    sections.append(current_section)
                
                # 开始新章节
                title = title_match.group(0).lstrip('#').strip()
                level = len(re.match(r'^#+', line).group()) if line.startswith('#') else 2
                
                current_section = {
                    'level': level,
                    'title': title,
                    'content': '',
                    'position': len(sections)
                }
                current_content = []
            else:
                current_content.append(line)
        
        # 保存最后一个章节
        if current_section:
            current_section['content'] = '\n'.join(current_content).strip()
            sections.append(current_section)
        
        return sections


# ==================== PDF解析器 ====================
class PDFParser(DocumentParser):
    """PDF文档解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.pdf']
    
    def parse(self, file_path: str) -> ParsedDocument:
        if not PDF_SUPPORT:
            raise ImportError("pdfplumber未安装，无法解析PDF")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        content_parts = []
        tables = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 提取文本
                text = page.extract_text()
                if text:
                    content_parts.append(text)
                
                # 提取表格
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        tables.append({
                            'page': page_num + 1,
                            'data': table
                        })
        
        content = '\n\n'.join(content_parts)
        title = self._extract_title(content)
        sections = self._split_sections(content)
        
        return ParsedDocument(
            file_path=file_path,
            file_type='pdf',
            title=title,
            content=content,
            sections=sections,
            tables=tables,
            images=[],
            metadata={'pages': len(tables) if tables else 1},
            parse_time=str(os.path.getmtime(file_path)),
            word_count=self._count_words(content)
        )


# ==================== Word解析器 ====================
class WordParser(DocumentParser):
    """Word文档解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.docx', '.doc']
    
    def parse(self, file_path: str) -> ParsedDocument:
        if not WORD_SUPPORT:
            raise ImportError("python-docx未安装，无法解析Word")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        doc = Document(file_path)
        
        content_parts = []
        sections = []
        tables = []
        
        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)
                
                # 检测标题
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 2
                    sections.append({
                        'level': level,
                        'title': text,
                        'content': '',
                        'position': len(sections)
                    })
        
        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                tables.append({
                    'table_idx': table_idx + 1,
                    'data': table_data
                })
        
        content = '\n\n'.join(content_parts)
        title = self._extract_title(content)
        
        # 如果没有从样式中提取到章节，使用文本分析
        if not sections:
            sections = self._split_sections(content)
        
        return ParsedDocument(
            file_path=file_path,
            file_type='word',
            title=title,
            content=content,
            sections=sections,
            tables=tables,
            images=[],
            metadata={'paragraphs': len(doc.paragraphs), 'tables': len(doc.tables)},
            parse_time=str(os.path.getmtime(file_path)),
            word_count=self._count_words(content)
        )


# ==================== 图片OCR解析器 ====================
class ImageParser(DocumentParser):
    """图片OCR解析器"""
    
    def __init__(self, lang: str = 'chi_sim+eng'):
        super().__init__()
        self.supported_formats = ['.png', '.jpg', '.jpeg', '.gif', '.bmp']
        self.lang = lang
    
    def parse(self, file_path: str) -> ParsedDocument:
        if not OCR_SUPPORT:
            raise ImportError("pytesseract/PIL未安装，无法进行OCR")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 打开图片
        image = Image.open(file_path)
        
        # OCR识别
        content = pytesseract.image_to_string(image, lang=self.lang)
        
        title = self._extract_title(content)
        sections = self._split_sections(content)
        
        return ParsedDocument(
            file_path=file_path,
            file_type='image',
            title=title,
            content=content,
            sections=sections,
            tables=[],
            images=[file_path],
            metadata={'size': image.size, 'mode': image.mode},
            parse_time=str(os.path.getmtime(file_path)),
            word_count=self._count_words(content)
        )


# ==================== 统一解析接口 ====================
class DocumentParserAPI:
    """文档解析统一API"""
    
    def __init__(self):
        self.parsers = {
            'pdf': PDFParser() if PDF_SUPPORT else None,
            'word': WordParser() if WORD_SUPPORT else None,
            'image': ImageParser() if OCR_SUPPORT else None
        }
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档（自动识别格式）
        
        Args:
            file_path: 文件路径
        
        Returns:
            ParsedDocument: 解析结果
        """
        ext = Path(file_path).suffix.lower()
        
        # PDF
        if ext == '.pdf':
            if not self.parsers['pdf']:
                raise ImportError("PDF解析器不可用")
            return self.parsers['pdf'].parse(file_path)
        
        # Word
        if ext in ['.docx', '.doc']:
            if not self.parsers['word']:
                raise ImportError("Word解析器不可用")
            return self.parsers['word'].parse(file_path)
        
        # 图片
        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
            if not self.parsers['image']:
                raise ImportError("图片OCR解析器不可用")
            return self.parsers['image'].parse(file_path)
        
        raise ValueError(f"不支持的文件格式: {ext}")
    
    def parse_to_dict(self, file_path: str) -> Dict:
        """解析文档并返回字典"""
        result = self.parse(file_path)
        return {
            'file_path': result.file_path,
            'file_type': result.file_type,
            'title': result.title,
            'content': result.content,
            'sections': result.sections,
            'tables': result.tables,
            'images': result.images,
            'metadata': result.metadata,
            'word_count': result.word_count
        }
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的格式列表"""
        formats = []
        if PDF_SUPPORT:
            formats.extend(['.pdf'])
        if WORD_SUPPORT:
            formats.extend(['.docx', '.doc'])
        if OCR_SUPPORT:
            formats.extend(['.png', '.jpg', '.jpeg', '.gif', '.bmp'])
        return formats


# ==================== 测试 ====================
if __name__ == '__main__':
    print("=" * 60)
    print("文档解析模块测试")
    print("=" * 60)
    
    api = DocumentParserAPI()
    
    print(f"\n支持的格式: {api.get_supported_formats()}")
    
    # 测试Word文档
    print("\n测试Word文档解析...")
    test_word = "/root/.openclaw/workspace/03_输出成果/导出文档/讨论报告_20260314_100738.docx"
    if os.path.exists(test_word):
        result = api.parse_to_dict(test_word)
        print(f"  文件类型: {result['file_type']}")
        print(f"  标题: {result['title']}")
        print(f"  字数: {result['word_count']}")
        print(f"  章节数: {len(result['sections'])}")
        print(f"  表格数: {len(result['tables'])}")
        print(f"  内容预览: {result['content'][:100]}...")
    else:
        print(f"  测试文件不存在: {test_word}")
    
    # 测试PDF（如果有）
    print("\n测试PDF解析支持...")
    if PDF_SUPPORT:
        print("  ✅ PDF解析器可用")
    else:
        print("  ⚠️ PDF解析器不可用（需安装pdfplumber）")
    
    # 测试OCR
    print("\n测试OCR解析支持...")
    if OCR_SUPPORT:
        print("  ✅ OCR解析器可用")
    else:
        print("  ⚠️ OCR解析器不可用（需安装pytesseract和PIL）")
    
    print("\n✅ 文档解析模块测试完成")
